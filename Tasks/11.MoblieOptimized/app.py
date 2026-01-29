import streamlit as st
import pdfplumber
import docx
import re
import spacy
from spacy.matcher import PhraseMatcher
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
import io
import requests

# ==================== PAGE CONFIG ====================
st.set_page_config(
    page_title="SkillGapAI Mobile",
    layout="centered",  # Use centered for mobile
    initial_sidebar_state="collapsed"
)

# ==================== MOBILE-OPTIMIZED CSS ====================
st.markdown("""
<style>
/* Mobile-first responsive design */
@media (max-width: 768px) {
    /* Reduce header size */
    h1 { font-size: 1.8rem !important; }
    h2 { font-size: 1.4rem !important; }
    h3 { font-size: 1.2rem !important; }
    
    /* Stack columns vertically */
    .row-widget.stHorizontal { flex-direction: column !important; }
    
    /* Full-width buttons */
    div[data-testid="stButton"] > button {
        width: 100% !important;
        margin: 5px 0 !important;
    }
    
    /* Reduce padding */
    .block-container { padding: 1rem !important; }
    
    /* Timeline container */
    .timeline-container {
        padding: 10px !important;
    }
}

/* Timeline Progress Bar */
.timeline-container {
    background: linear-gradient(90deg, #470047 0%, #7d1a7d 100%);
    padding: 20px;
    border-radius: 10px;
    margin-bottom: 20px;
}

.timeline-step {
    display: flex;
    align-items: center;
    margin: 10px 0;
    position: relative;
}

.timeline-circle {
    width: 40px;
    height: 40px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-weight: bold;
    z-index: 1;
    font-size: 16px;
}

.timeline-circle.completed {
    background-color: #28a745;
    color: white;
}

.timeline-circle.active {
    background-color: #ffc107;
    color: black;
    box-shadow: 0 0 15px rgba(255, 193, 7, 0.8);
    animation: pulse 2s infinite;
}

.timeline-circle.pending {
    background-color: #6c757d;
    color: white;
}

@keyframes pulse {
    0%, 100% { box-shadow: 0 0 10px rgba(255, 193, 7, 0.5); }
    50% { box-shadow: 0 0 20px rgba(255, 193, 7, 1); }
}

.timeline-label {
    margin-left: 15px;
    color: white;
    font-size: 14px;
    font-weight: 500;
}

.timeline-line {
    position: absolute;
    left: 19px;
    top: 40px;
    width: 2px;
    height: 100%;
    background-color: rgba(255, 255, 255, 0.3);
}

/* Next/Continue Button */
.next-button {
    background: linear-gradient(135deg, #28a745 0%, #20c997 100%);
    color: white;
    font-size: 18px;
    font-weight: bold;
    padding: 15px 30px;
    border-radius: 10px;
    border: none;
    width: 100%;
    cursor: pointer;
    transition: all 0.3s ease;
    box-shadow: 0 4px 8px rgba(40, 167, 69, 0.3);
}

.next-button:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 12px rgba(40, 167, 69, 0.5);
}

/* Analyze Button - Enhanced */
div[data-testid="stButton"] > button {
    background: linear-gradient(135deg, #470047 0%, #7d1a7d 100%);
    color: white;
    height: 55px;
    font-size: 18px;
    border-radius: 10px;
    border: none;
    transition: all 0.3s ease;
    box-shadow: 0 4px 6px rgba(71, 0, 71, 0.3);
}

div[data-testid="stButton"] > button:hover {
    background: linear-gradient(135deg, #5a0a5a 0%, #9d4edd 100%);
    transform: translateY(-2px);
    box-shadow: 0 6px 12px rgba(71, 0, 71, 0.5);
}

/* Section headers */
.section-header {
    background-color: #470047;
    padding: 15px;
    border-radius: 10px;
    color: white;
    margin: 20px 0 10px 0;
}

/* Card style for mobile */
.info-card {
    background-color: #f8f9fa;
    padding: 15px;
    border-radius: 8px;
    border-left: 4px solid #470047;
    margin: 10px 0;
}

/* Mobile metrics */
.mobile-metric {
    text-align: center;
    padding: 15px;
    background: white;
    border-radius: 8px;
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    margin: 10px 0;
}

/* Sticky navigation */
.sticky-nav {
    position: sticky;
    top: 0;
    z-index: 999;
    background: white;
    padding: 10px 0;
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
}
</style>
""", unsafe_allow_html=True)

# ==================== SESSION STATE INITIALIZATION ====================
if 'current_step' not in st.session_state:
    st.session_state.current_step = 1
if 'resume_file' not in st.session_state:
    st.session_state.resume_file = None
if 'jd_file' not in st.session_state:
    st.session_state.jd_file = None
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = None
if 'jd_text' not in st.session_state:
    st.session_state.jd_text = None
if 'cleaned_resume' not in st.session_state:
    st.session_state.cleaned_resume = ""
if 'cleaned_jd' not in st.session_state:
    st.session_state.cleaned_jd = ""
if 'resume_skills' not in st.session_state:
    st.session_state.resume_skills = None
if 'jd_skills' not in st.session_state:
    st.session_state.jd_skills = None
if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False

# ==================== TIMELINE STEPS ====================
STEPS = {
    1: {"name": "Upload", "icon": "📤", "desc": "Upload Documents"},
    2: {"name": "Extract", "icon": "📄", "desc": "Text Extraction"},
    3: {"name": "Analyze", "icon": "🔍", "desc": "Skill Analysis"},
    4: {"name": "Results", "icon": "📊", "desc": "View Results"},
    5: {"name": "Reports", "icon": "📥", "desc": "Download Reports"}
}

# ==================== HELPER FUNCTIONS ====================
def extract_text(file):
    """Extract text from PDF, DOCX, or TXT files"""
    try:
        if file.type == "application/pdf":
            text = ""
            file.seek(0)
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            return text.strip()
        
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            file.seek(0)
            document = docx.Document(file)
            paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
            table_text = []
            for table in document.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))
            all_text = "\n".join(paragraphs)
            if table_text:
                all_text += "\n" + "\n".join(table_text)
            return all_text.strip()
        
        elif file.type == "text/plain":
            file.seek(0)
            text = file.read().decode("utf-8")
            return text.strip()
        else:
            return ""
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return ""

def clean(text):
    """Clean and normalize text"""
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r'(\d+)\s*(?:-|–|to)\s*(\d+)', r'\1_\2', text)
    text = re.sub(r'(\d+)\s*\+', r'\1_plus', text)
    text = re.sub(r'\n+', ' ', text)
    text = re.sub(r'\t+', ' ', text)
    text = re.sub(r'[^a-z0-9_\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

@st.cache_resource
def load_spacy_model():
    try:
        return spacy.load("en_core_web_sm")
    except:
        st.warning("Installing spaCy model...")
        import os
        os.system("python -m spacy download en_core_web_sm")
        return spacy.load("en_core_web_sm")

# Master skills
MASTER_TECHNICAL_SKILLS = {
    "python", "java", "javascript", "typescript", "react", "angular", "vue.js",
    "node.js", "django", "flask", "fastapi", "sql", "mongodb", "postgresql",
    "mysql", "aws", "azure", "docker", "kubernetes", "git", "html", "css",
    "machine learning", "deep learning", "tensorflow", "pytorch", "pandas",
    "numpy", "data analysis", "excel", "tableau", "power bi"
}

MASTER_SOFT_SKILLS = {
    "communication", "teamwork", "leadership", "problem solving", "critical thinking",
    "time management", "adaptability", "creativity", "collaboration", "presentation skills"
}

def extract_technical_skills(text, nlp):
    """Extract technical skills using spaCy"""
    if not text or not nlp:
        return []
    matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
    patterns = [nlp.make_doc(skill) for skill in MASTER_TECHNICAL_SKILLS]
    matcher.add("TECH_SKILLS", patterns)
    doc = nlp(text[:100000])  # Limit length for mobile
    matches = matcher(doc)
    skills_found = set()
    for _, start, end in matches:
        skills_found.add(doc[start:end].text.lower())
    return sorted(skills_found)

def extract_soft_skills(text):
    """Extract soft skills using text matching"""
    if not text:
        return []
    text = text.lower()
    found_skills = set()
    for skill in MASTER_SOFT_SKILLS:
        if skill in text:
            found_skills.add(skill)
    return sorted(found_skills)

# ==================== TIMELINE UI ====================
def render_timeline():
    """Render progress timeline"""
    timeline_html = '<div class="timeline-container">'
    
    for step_num in range(1, 6):
        step = STEPS[step_num]
        
        # Determine status
        if step_num < st.session_state.current_step:
            status = "completed"
            icon = "✓"
        elif step_num == st.session_state.current_step:
            status = "active"
            icon = step["icon"]
        else:
            status = "pending"
            icon = step["icon"]
        
        # Build HTML
        timeline_html += f'''
        <div class="timeline-step">
            <div class="timeline-circle {status}">{icon}</div>
            <div class="timeline-label">{step["desc"]}</div>
            {"<div class='timeline-line'></div>" if step_num < 5 else ""}
        </div>
        '''
    
    timeline_html += '</div>'
    st.markdown(timeline_html, unsafe_allow_html=True)

# ==================== HEADER ====================
st.markdown("""
<h1 style="text-align:center; color:#470047;">📱 SkillGapAI Mobile</h1>
<p style="text-align:center; color:#6B7280;">AI-Powered Skill Gap Analysis</p>
<hr>
""", unsafe_allow_html=True)

# ==================== RENDER TIMELINE ====================
render_timeline()

# ==================== NAVIGATION BUTTONS ====================
def next_step():
    if st.session_state.current_step < 5:
        st.session_state.current_step += 1
        st.rerun()

def prev_step():
    if st.session_state.current_step > 1:
        st.session_state.current_step -= 1
        st.rerun()

# ==================== STEP 1: UPLOAD ====================
if st.session_state.current_step == 1:
    st.markdown('<div class="section-header"><h2>📤 Step 1: Upload Documents</h2></div>', unsafe_allow_html=True)
    
    st.info("📱 **Mobile Tip:** Hold your device vertically for the best experience")
    
    # Resume upload
    st.markdown("### 📄 Upload Resume")
    resume_file = st.file_uploader(
        "Choose your resume",
        type=["pdf", "docx", "txt"],
        key="resume_uploader",
        help="PDF, DOCX, or TXT format"
    )
    
    if resume_file:
        st.session_state.resume_file = resume_file
        st.success(f"✅ Resume uploaded: {resume_file.name}")
    
    st.markdown("---")
    
    # JD upload
    st.markdown("### 📋 Upload Job Description")
    jd_file = st.file_uploader(
        "Choose job description",
        type=["pdf", "docx", "txt"],
        key="jd_uploader",
        help="PDF, DOCX, or TXT format"
    )
    
    if jd_file:
        st.session_state.jd_file = jd_file
        st.success(f"✅ Job Description uploaded: {jd_file.name}")
    
    st.markdown("---")
    
    # Continue button
    if st.session_state.resume_file and st.session_state.jd_file:
        if st.button("Continue to Extraction ➡️", use_container_width=True, type="primary"):
            next_step()
    else:
        st.warning("⚠️ Please upload both files to continue")

# ==================== STEP 2: EXTRACT ====================
elif st.session_state.current_step == 2:
    st.markdown('<div class="section-header"><h2>📄 Step 2: Text Extraction</h2></div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅️ Back", use_container_width=True):
            prev_step()
    
    # Extract text
    with st.spinner("🔄 Extracting text from documents..."):
        if not st.session_state.resume_text:
            st.session_state.resume_text = extract_text(st.session_state.resume_file)
        if not st.session_state.jd_text:
            st.session_state.jd_text = extract_text(st.session_state.jd_file)
    
    # Show preview
    if st.session_state.resume_text:
        st.markdown("### Resume Preview")
        with st.expander("View extracted text", expanded=False):
            st.text_area("", st.session_state.resume_text[:500] + "...", height=150, disabled=True)
        st.caption(f"📊 {len(st.session_state.resume_text)} characters extracted")
    
    st.markdown("---")
    
    if st.session_state.jd_text:
        st.markdown("### Job Description Preview")
        with st.expander("View extracted text", expanded=False):
            st.text_area("", st.session_state.jd_text[:500] + "...", height=150, disabled=True, key="jd_preview")
        st.caption(f"📊 {len(st.session_state.jd_text)} characters extracted")
    
    # Clean text
    with st.spinner("🧹 Cleaning text..."):
        st.session_state.cleaned_resume = clean(st.session_state.resume_text)
        st.session_state.cleaned_jd = clean(st.session_state.jd_text)
    
    st.success("✅ Text extraction complete!")
    
    if st.button("Continue to Analysis ➡️", use_container_width=True, type="primary"):
        next_step()

# ==================== STEP 3: ANALYZE ====================
elif st.session_state.current_step == 3:
    st.markdown('<div class="section-header"><h2>🔍 Step 3: Skill Analysis</h2></div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅️ Back", use_container_width=True):
            prev_step()
    
    # Load NLP model
    nlp = load_spacy_model()
    
    if not st.session_state.resume_skills:
        with st.spinner("🤖 Analyzing resume skills..."):
            resume_technical = extract_technical_skills(st.session_state.cleaned_resume, nlp)
            resume_soft = extract_soft_skills(st.session_state.cleaned_resume)
            st.session_state.resume_skills = {
                "technical": resume_technical,
                "soft": resume_soft
            }
    
    if not st.session_state.jd_skills:
        with st.spinner("🤖 Analyzing job description skills..."):
            jd_technical = extract_technical_skills(st.session_state.cleaned_jd, nlp)
            jd_soft = extract_soft_skills(st.session_state.cleaned_jd)
            st.session_state.jd_skills = {
                "technical": jd_technical,
                "soft": jd_soft
            }
    
    # Display extracted skills
    st.markdown("### 📊 Extracted Skills")
    
    # Resume skills
    st.markdown("#### Your Resume")
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Technical Skills", len(st.session_state.resume_skills["technical"]))
    with col2:
        st.metric("Soft Skills", len(st.session_state.resume_skills["soft"]))
    
    with st.expander("View Resume Skills"):
        if st.session_state.resume_skills["technical"]:
            st.markdown("**Technical:**")
            st.write(", ".join(st.session_state.resume_skills["technical"][:10]) + 
                    (f" ...+{len(st.session_state.resume_skills['technical'])-10} more" 
                     if len(st.session_state.resume_skills['technical']) > 10 else ""))
        if st.session_state.resume_skills["soft"]:
            st.markdown("**Soft:**")
            st.write(", ".join(st.session_state.resume_skills["soft"][:10]) +
                    (f" ...+{len(st.session_state.resume_skills['soft'])-10} more" 
                     if len(st.session_state.resume_skills['soft']) > 10 else ""))
    
    st.markdown("---")
    
    # JD skills
    st.markdown("#### Job Requirements")
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Technical Skills", len(st.session_state.jd_skills["technical"]))
    with col2:
        st.metric("Soft Skills", len(st.session_state.jd_skills["soft"]))
    
    with st.expander("View Job Requirements"):
        if st.session_state.jd_skills["technical"]:
            st.markdown("**Technical:**")
            st.write(", ".join(st.session_state.jd_skills["technical"][:10]) +
                    (f" ...+{len(st.session_state.jd_skills['technical'])-10} more" 
                     if len(st.session_state.jd_skills['technical']) > 10 else ""))
        if st.session_state.jd_skills["soft"]:
            st.markdown("**Soft:**")
            st.write(", ".join(st.session_state.jd_skills["soft"][:10]) +
                    (f" ...+{len(st.session_state.jd_skills['soft'])-10} more" 
                     if len(st.session_state.jd_skills['soft']) > 10 else ""))
    
    st.success("✅ Skill analysis complete!")
    
    if st.button("View Results ➡️", use_container_width=True, type="primary"):
        st.session_state.analysis_complete = True
        next_step()

# ==================== STEP 4: RESULTS ====================
elif st.session_state.current_step == 4:
    st.markdown('<div class="section-header"><h2>📊 Step 4: Results</h2></div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅️ Back", use_container_width=True):
            prev_step()
    
    # Calculate match
    resume_all = set(st.session_state.resume_skills["technical"] + st.session_state.resume_skills["soft"])
    jd_all = set(st.session_state.jd_skills["technical"] + st.session_state.jd_skills["soft"])
    
    matched = resume_all & jd_all
    missing = jd_all - resume_all
    
    match_percent = round((len(matched) / len(jd_all) * 100), 1) if len(jd_all) > 0 else 0
    
    # Overall match card
    st.markdown(f"""
    <div style="background: linear-gradient(135deg, #470047 0%, #7d1a7d 100%); 
                padding: 30px; border-radius: 15px; text-align: center; color: white; margin: 20px 0;">
        <h1 style="font-size: 3rem; margin: 0;">{match_percent}%</h1>
        <p style="font-size: 1.2rem; margin: 10px 0 0 0;">Overall Match</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Metrics
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("✅ Matched", len(matched))
    with col2:
        st.metric("📋 Required", len(jd_all))
    with col3:
        st.metric("❌ Missing", len(missing))
    
    st.markdown("---")
    
    # Skills breakdown
    st.markdown("### 📋 Skill Breakdown")
    
    # Matched skills
    with st.expander(f"✅ Matched Skills ({len(matched)})", expanded=True):
        if matched:
            for skill in list(matched)[:10]:
                st.success(f"✓ {skill}")
            if len(matched) > 10:
                st.caption(f"...and {len(matched)-10} more")
        else:
            st.info("No perfectly matched skills")
    
    # Missing skills
    with st.expander(f"❌ Missing Skills ({len(missing)})", expanded=True):
        if missing:
            st.markdown("**🔴 Priority skills to learn:**")
            for i, skill in enumerate(list(missing)[:10], 1):
                st.error(f"{i}. {skill.title()}")
            if len(missing) > 10:
                st.caption(f"...and {len(missing)-10} more")
        else:
            st.success("You have all required skills!")
    
    # Recommendations
    st.markdown("### 💡 Recommendations")
    if match_percent >= 70:
        st.success("🎉 Excellent match! You're well-qualified for this role.")
    elif match_percent >= 50:
        st.warning("⚠️ Good match, but focus on developing the missing skills.")
    else:
        st.error("❗ Significant skill gaps. Consider upskilling before applying.")
    
    if st.button("Download Reports ➡️", use_container_width=True, type="primary"):
        next_step()

# ==================== STEP 5: REPORTS ====================
elif st.session_state.current_step == 5:
    st.markdown('<div class="section-header"><h2>📥 Step 5: Download Reports</h2></div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅️ Back to Results", use_container_width=True):
            prev_step()
    
    st.info("📱 **Mobile Tip:** Reports will be downloaded to your device's Downloads folder")
    
    # Calculate final stats
    resume_all = set(st.session_state.resume_skills["technical"] + st.session_state.resume_skills["soft"])
    jd_all = set(st.session_state.jd_skills["technical"] + st.session_state.jd_skills["soft"])
    matched = resume_all & jd_all
    missing = jd_all - resume_all
    match_percent = round((len(matched) / len(jd_all) * 100), 1) if len(jd_all) > 0 else 0
    
    # CSV Report
    st.markdown("### 📊 CSV Report")
    st.caption("Lightweight format for spreadsheets")
    
    tech_matched = [s for s in matched if s in st.session_state.jd_skills["technical"]]
    tech_missing = [s for s in missing if s in st.session_state.jd_skills["technical"]]
    soft_matched = [s for s in matched if s in st.session_state.jd_skills["soft"]]
    soft_missing = [s for s in missing if s in st.session_state.jd_skills["soft"]]
    
    max_len = max(len(tech_matched), len(tech_missing), len(soft_matched), len(soft_missing), 1)
    tech_matched += [''] * (max_len - len(tech_matched))
    tech_missing += [''] * (max_len - len(tech_missing))
    soft_matched += [''] * (max_len - len(soft_matched))
    soft_missing += [''] * (max_len - len(soft_missing))
    
    df = pd.DataFrame({
        "Tech Matched": tech_matched,
        "Tech Missing": tech_missing,
        "Soft Matched": soft_matched,
        "Soft Missing": soft_missing
    })
    
    csv = df.to_csv(index=False)
    st.download_button(
        "📥 Download CSV",
        csv.encode("utf-8"),
        "skill_gap_report_mobile.csv",
        "text/csv",
        use_container_width=True
    )
    
    st.markdown("---")
    
    # Summary card
    st.markdown(f"""
    <div class="info-card">
        <h3>📊 Analysis Summary</h3>
        <p><strong>Match Rate:</strong> {match_percent}%</p>
        <p><strong>Matched Skills:</strong> {len(matched)}</p>
        <p><strong>Missing Skills:</strong> {len(missing)}</p>
        <p><strong>Total Required:</strong> {len(jd_all)}</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Start over button
    st.markdown("---")
    if st.button("🔄 Analyze Another Job", use_container_width=True, type="secondary"):
        # Reset session state
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

# ==================== FOOTER ====================
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #6B7280; padding: 20px;">
    <p>📱 SkillGapAI Mobile v1.0</p>
    <p style="font-size: 0.9rem;">Swipe-friendly • Touch-optimized • Fast Analysis</p>
</div>
""", unsafe_allow_html=True)