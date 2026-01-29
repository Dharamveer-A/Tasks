import streamlit as st
import pdfplumber
import docx
import re
import spacy
from spacy.matcher import PhraseMatcher
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
import io
import requests
from datetime import datetime, timedelta

# ==================== PAGE CONFIG ====================
st.set_page_config(
    page_title="SkillGapAI Mobile",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ==================== SESSION STATE ====================
if 'step' not in st.session_state:
    st.session_state.step = 1
if 'resume_file' not in st.session_state:
    st.session_state.resume_file = None
if 'jd_file' not in st.session_state:
    st.session_state.jd_file = None
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = None
if 'jd_text' not in st.session_state:
    st.session_state.jd_text = None
if 'cleaned_resume' not in st.session_state:
    st.session_state.cleaned_resume = ""
if 'cleaned_jd' not in st.session_state:
    st.session_state.cleaned_jd = ""
if 'resume_skills' not in st.session_state:
    st.session_state.resume_skills = None
if 'jd_skills' not in st.session_state:
    st.session_state.jd_skills = None
if 'resume_experience' not in st.session_state:
    st.session_state.resume_experience = None
if 'jd_experience' not in st.session_state:
    st.session_state.jd_experience = None
if 'github_skills' not in st.session_state:
    st.session_state.github_skills = []
if 'github_analysis' not in st.session_state:
    st.session_state.github_analysis = None
if 'similarity_matrix' not in st.session_state:
    st.session_state.similarity_matrix = None
if 'skill_match_result' not in st.session_state:
    st.session_state.skill_match_result = None
if 'match_counts' not in st.session_state:
    st.session_state.match_counts = None

# ==================== MOBILE CSS WITH REDDISH PASTEL THEME ====================
st.markdown("""
<style>
/* Reddish Pastel Theme - Mobile Optimized */
@media (max-width: 768px) {
    h1 { font-size: 1.8rem !important; }
    h2 { font-size: 1.4rem !important; }
    h3 { font-size: 1.2rem !important; }
    .block-container { padding: 1rem !important; }
    div[data-testid="stButton"] > button {
        width: 100% !important;
        font-size: 16px !important;
        padding: 12px !important;
    }
}

/* Headers */
h1, h2, h3 {
    color: #D77A7D !important;
}

/* Timeline Container */
.timeline-container {
    background: linear-gradient(135deg, #E8999C 0%, #F4C7C3 100%);
    padding: 20px;
    border-radius: 15px;
    margin: 20px 0;
    box-shadow: 0 4px 12px rgba(232, 153, 156, 0.3);
}

.timeline-steps {
    display: flex;
    justify-content: space-between;
    align-items: center;
    flex-wrap: wrap;
}

.timeline-step {
    flex: 1;
    min-width: 80px;
    text-align: center;
    position: relative;
    padding: 10px;
}

.step-circle {
    width: 50px;
    height: 50px;
    border-radius: 50%;
    margin: 0 auto 10px;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 20px;
    font-weight: bold;
    transition: all 0.3s ease;
}

.step-circle.completed {
    background: #E8999C;
    color: white;
    box-shadow: 0 2px 8px rgba(232, 153, 156, 0.4);
}

.step-circle.active {
    background: #FFB4A2;
    color: white;
    box-shadow: 0 4px 12px rgba(255, 180, 162, 0.6);
    animation: pulse 2s infinite;
}

.step-circle.pending {
    background: rgba(255, 255, 255, 0.2);
    color: white;
}

@keyframes pulse {
    0%, 100% { transform: scale(1); box-shadow: 0 4px 12px rgba(255, 180, 162, 0.4); }
    50% { transform: scale(1.05); box-shadow: 0 6px 16px rgba(255, 180, 162, 0.8); }
}

.step-label {
    color: white;
    font-size: 12px;
    font-weight: 500;
}

.step-line {
    position: absolute;
    top: 35px;
    left: 50%;
    width: 100%;
    height: 2px;
    background: rgba(255, 255, 255, 0.3);
    z-index: 0;
}

/* Enhanced button styling with reddish theme */
div[data-testid="stButton"] > button {
    background: linear-gradient(135deg, #E8999C 0%, #F4C7C3 100%);
    color: white;
    border-radius: 10px;
    border: none;
    transition: all 0.3s ease;
    box-shadow: 0 4px 6px rgba(232, 153, 156, 0.3);
    font-weight: 600;
}

div[data-testid="stButton"] > button:hover {
    background: linear-gradient(135deg, #D77A7D 0%, #FFB4A2 100%);
    transform: translateY(-2px);
    box-shadow: 0 6px 12px rgba(232, 153, 156, 0.5);
}

.stDownloadButton>button {
    background-color: #F4C7C3 !important;
    color: #5A3A3A !important;
    border: 2px solid #E8999C !important;
    border-radius: 8px !important;
}

.stDownloadButton>button:hover {
    background-color: #E8999C !important;
    color: white !important;
}

[data-testid="stMetricValue"] {
    color: #D77A7D !important;
    font-weight: 700 !important;
}

.stProgress > div > div > div > div {
    background-color: #E8999C !important;
}

.streamlit-expanderHeader {
    background-color: #FFE5E5 !important;
    color: #5A3A3A !important;
    border-radius: 8px !important;
}

.stTextArea textarea {
    border: 2px solid #F4C7C3 !important;
    border-radius: 8px !important;
}

.stFileUploader {
    background-color: #FFE5E5 !important;
    border-radius: 8px !important;
    padding: 1rem !important;
}

.stTextInput input {
    border: 2px solid #F4C7C3 !important;
    border-radius: 8px !important;
}

.stTextInput input:focus {
    border-color: #E8999C !important;
    box-shadow: 0 0 0 0.2rem rgba(232, 153, 156, 0.25) !important;
}

.progress-text {
    text-align: center;
    color: white;
    font-size: 14px;
    margin-top: 10px;
    font-weight: 500;
}
</style>
""", unsafe_allow_html=True)

# ==================== TIMELINE COMPONENT ====================
def render_timeline():
    steps = [
        {"num": 1, "name": "Upload", "icon": "📤"},
        {"num": 2, "name": "Extract", "icon": "📄"},
        {"num": 3, "name": "Clean", "icon": "🧹"},
        {"num": 4, "name": "Skills", "icon": "🔍"},
        {"num": 5, "name": "Analysis", "icon": "📊"},
        {"num": 6, "name": "Reports", "icon": "📥"}
    ]
    
    current = st.session_state.step
    
    timeline_html = '<div class="timeline-container"><div class="timeline-steps">'
    
    for i, step in enumerate(steps):
        if step["num"] < current:
            status = "completed"
            icon = "✓"
        elif step["num"] == current:
            status = "active"
            icon = step["icon"]
        else:
            status = "pending"
            icon = step["icon"]
        
        timeline_html += f'''
        <div class="timeline-step">
            {f'<div class="step-line"></div>' if i < len(steps)-1 else ''}
            <div class="step-circle {status}">{icon}</div>
            <div class="step-label">{step["name"]}</div>
        </div>
        '''
    
    timeline_html += '</div>'
    timeline_html += f'<div class="progress-text">Step {current} of 6: {steps[current-1]["name"]}</div>'
    timeline_html += '</div>'
    
    st.markdown(timeline_html, unsafe_allow_html=True)

# ==================== NAVIGATION FUNCTIONS ====================
def next_step():
    if st.session_state.step < 6:
        st.session_state.step += 1
        st.rerun()

def prev_step():
    if st.session_state.step > 1:
        st.session_state.step -= 1
        st.rerun()

def go_to_step(step_num):
    st.session_state.step = step_num
    st.rerun()

# ==================== UTILITY FUNCTIONS ====================
def extract_text(file):
    try:
        if file.type == "application/pdf":
            text = ""
            file.seek(0)
            try:
                with pdfplumber.open(file) as pdf:
                    total_pages = len(pdf.pages)
                    st.info(f"Processing PDF: {total_pages} page(s) detected")
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                    if not text.strip():
                        st.error("No text could be extracted from the PDF.")
                        return ""
            except Exception as pdf_error:
                st.error(f"PDF Processing Error: {str(pdf_error)}")
                return ""
            return text.strip()

        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            file.seek(0)
            try:
                document = docx.Document(file)
                paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
                table_text = []
                for table in document.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(" | ".join(row_text))
                all_text = "\n".join(paragraphs)
                if table_text:
                    all_text += "\n" + "\n".join(table_text)
                if not all_text.strip():
                    st.error("No text could be extracted from the DOCX file.")
                    return ""
                return all_text.strip()
            except Exception as docx_error:
                st.error(f"DOCX Processing Error: {str(docx_error)}")
                return ""

        elif file.type == "text/plain":
            file.seek(0)
            try:
                text = file.read().decode("utf-8")
                if not text.strip():
                    st.error("The text file appears to be empty.")
                    return ""
                return text.strip()
            except UnicodeDecodeError:
                try:
                    file.seek(0)
                    text = file.read().decode("latin-1")
                    st.warning("File encoding detected as Latin-1 instead of UTF-8")
                    return text.strip()
                except:
                    st.error("Unable to read the text file.")
                    return ""
        else:
            st.error(f"Unsupported file type: {file.type}")
            return ""
    except Exception as e:
        st.error(f"Error processing {file.name}: {str(e)}")
        return ""

def clean(text):
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r'(\d+)\s*(?:-|–|to)\s*(\d+)', r'\1_\2', text)
    text = re.sub(r'(\d+)\s*\+', r'\1_plus', text)
    text = re.sub(r'\n+', ' ', text)
    text = re.sub(r'\t+', ' ', text)
    text = re.sub(r'[^a-z0-9_\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def extract_experience(text):
    if not text:
        return None
    text = text.lower()
    range_match = re.search(r'(\d+)\s*(?:-|–|to)\s*(\d+)\s*(?:years?|yrs?)', text)
    if range_match:
        return {"min_exp": int(range_match.group(1)), "max_exp": int(range_match.group(2))}
    plus_match = re.search(r'(\d+)\s*\+\s*(?:years?|yrs?)', text)
    if plus_match:
        return {"min_exp": int(plus_match.group(1)), "max_exp": None}
    with_years_match = re.search(r'(?:with|over|around|approximately)\s+(\d+)\s+(?:years?|yrs?)', text)
    if with_years_match:
        years = int(with_years_match.group(1))
        return {"min_exp": years, "max_exp": years}
    years_of_match = re.search(r'(\d+)\s+(?:years?|yrs?)\s+(?:of\s+)?(?:experience|specializing|in|working)', text)
    if years_of_match:
        years = int(years_of_match.group(1))
        return {"min_exp": years, "max_exp": years}
    general_match = re.search(r'(\d+)\s+(?:years?|yrs?)', text)
    if general_match:
        years = int(general_match.group(1))
        return {"min_exp": years, "max_exp": years}
    return None

@st.cache_resource
def load_spacy_model():
    try:
        return spacy.load("en_core_web_trf")
    except OSError:
        st.error("spaCy model 'en_core_web_trf' not found.")
        return None

master_technical_skills = {
    "python", "java", "c", "c++", "c#", "go", "rust", "kotlin", "swift",
    "javascript", "typescript", "php", "ruby", "r", "matlab", "scala",
    "perl", "bash", "shell scripting", "sql",
    "html", "css", "sass", "bootstrap", "tailwind css",
    "react", "angular", "vue.js", "next.js", "nuxt.js",
    "node.js", "express.js", "django", "flask", "fastapi",
    "spring", "spring boot", "laravel", "asp.net",
    "mysql", "postgresql", "sqlite", "oracle", "sql server",
    "mongodb", "cassandra", "dynamodb", "redis", "firebase",
    "elasticsearch", "neo4j",
    "numpy", "pandas", "scipy", "matplotlib", "seaborn",
    "plotly", "power bi", "tableau", "excel", "statistics",
    "data analysis", "data visualization",
    "machine learning", "deep learning", "artificial intelligence",
    "tensorflow", "keras", "pytorch", "scikit-learn", "xgboost",
    "opencv", "nlp", "computer vision", "speech recognition",
    "transformers", "huggingface", "langchain", "llm",
    "hadoop", "spark", "pyspark", "kafka", "hive", "pig",
    "hbase", "flink", "airflow", "databricks",
    "aws", "azure", "google cloud", "gcp",
    "docker", "kubernetes", "terraform", "ansible",
    "jenkins", "github actions", "ci/cd",
    "linux", "unix",
    "rest api", "graphql", "grpc", "soap",
    "microservices", "jwt", "oauth",
    "unit testing", "integration testing", "system testing",
    "pytest", "unittest", "junit", "selenium", "cypress",
    "postman",
    "android", "ios", "flutter", "react native",
    "swiftui", "kotlin multiplatform",
    "network security", "application security", "penetration testing",
    "ethical hacking", "cryptography", "owasp",
    "siem", "firewall",
    "windows", "linux", "macos",
    "git", "github", "gitlab", "bitbucket",
    "jira", "confluence", "trello",
    "data structures", "algorithms", "object oriented programming",
    "design patterns", "system design",
    "software development lifecycle", "agile", "scrum"
}

master_soft_skills = {
    "communication", "verbal communication", "written communication",
    "public speaking", "presentation skills", "active listening",
    "business communication", "storytelling",
    "teamwork", "collaboration", "interpersonal skills",
    "relationship building", "empathy", "emotional intelligence",
    "conflict resolution", "negotiation",
    "leadership", "people management", "team leadership",
    "decision making", "delegation", "mentoring", "coaching",
    "influencing", "strategic thinking",
    "problem solving", "critical thinking", "analytical thinking",
    "logical reasoning", "creative thinking", "innovation",
    "root cause analysis", "troubleshooting",
    "time management", "prioritization", "multitasking",
    "work ethic", "self discipline", "accountability",
    "goal setting", "organizational skills",
    "adaptability", "flexibility", "resilience",
    "learning agility", "continuous learning",
    "open mindedness", "growth mindset",
    "professionalism", "integrity", "ethical behavior",
    "reliability", "punctuality", "confidentiality",
    "creativity", "idea generation", "design thinking",
    "innovation mindset", "curiosity",
    "stress management", "self awareness", "self motivation",
    "confidence", "positive attitude", "emotional control",
    "customer focus", "customer service",
    "client management", "stakeholder management",
    "user empathy", "service mindset",
    "cross functional collaboration", "cultural awareness",
    "diversity and inclusion", "remote collaboration",
    "team alignment",
    "conflict management", "crisis management",
    "handling pressure", "decision making under pressure",
    "independent work", "collaborative work",
    "attention to detail", "quality focus",
    "result oriented", "ownership",
    "honesty", "trustworthiness", "respect",
    "fairness", "social responsibility"
}

def extract_soft_skills(text, skill_set):
    if not text:
        return []
    text = text.lower()
    found_skills = set()
    for skill in skill_set:
        if skill in text:
            found_skills.add(skill)
    return sorted(found_skills)

def extract_technical_skills(text, skill_set, nlp):
    if not text:
        return []
    max_length = 1000000
    if len(text) > max_length:
        chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
        all_skills = set()
        for chunk in chunks:
            matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
            patterns = [nlp.make_doc(skill) for skill in skill_set]
            matcher.add("TECH_SKILLS", patterns)
            doc = nlp(chunk)
            matches = matcher(doc)
            for _, start, end in matches:
                all_skills.add(doc[start:end].text.lower())
        return sorted(list(all_skills))
    else:
        matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
        patterns = [nlp.make_doc(skill) for skill in skill_set]
        matcher.add("TECH_SKILLS", patterns)
        doc = nlp(text)
        matches = matcher(doc)
        skills_found = set()
        for _, start, end in matches:
            skills_found.add(doc[start:end].text.lower())
        return sorted(skills_found)

def extract_github_username(url):
    match = re.search(r'github\.com/([^/]+)/?', url)
    return match.group(1) if match else None

def detect_github_in_resume(text):
    if not text:
        return None
    url_pattern = r'(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9](?:[a-zA-Z0-9-]{0,38}[a-zA-Z0-9])?)'
    url_match = re.search(url_pattern, text, re.IGNORECASE)
    if url_match:
        return f"https://github.com/{url_match.group(1)}"
    username_pattern = r'(?:github|gh)[:\s]+([a-zA-Z0-9](?:[a-zA-Z0-9-]{0,38}[a-zA-Z0-9])?)'
    username_match = re.search(username_pattern, text, re.IGNORECASE)
    if username_match:
        return f"https://github.com/{username_match.group(1)}"
    return None

def fetch_github_repos(username):
    api_url = f"https://api.github.com/users/{username}/repos"
    try:
        response = requests.get(api_url, timeout=10)
        if response.status_code == 200:
            return response.json()
    except:
        pass
    return []

def extract_github_skills(repos, skill_set):
    found_skills = set()
    for repo in repos:
        text = f"{repo.get('name','')} {repo.get('description','')}".lower()
        for skill in skill_set:
            if skill in text:
                found_skills.add(skill)
        if repo.get("language"):
            lang = repo["language"].lower()
            if lang in skill_set:
                found_skills.add(lang)
    return sorted(found_skills)

def analyze_github_profile(repos):
    if not repos:
        return None
    total_repos = len(repos)
    languages_used = {}
    topics_used = set()
    has_readme_projects = 0
    recent_activity = 0
    for repo in repos:
        if repo.get("language"):
            lang = repo["language"]
            languages_used[lang] = languages_used.get(lang, 0) + 1
        if repo.get("topics"):
            topics_used.update(repo["topics"])
        if repo.get("description"):
            has_readme_projects += 1
        if repo.get("updated_at"):
            try:
                updated = datetime.strptime(repo["updated_at"], "%Y-%m-%dT%H:%M:%SZ")
                if updated > datetime.now() - timedelta(days=180):
                    recent_activity += 1
            except:
                pass
    insights = {
        "total_repos": total_repos,
        "languages": languages_used,
        "top_language": max(languages_used.items(), key=lambda x: x[1])[0] if languages_used else None,
        "topics": list(topics_used),
        "documented_projects": has_readme_projects,
        "active_repos": recent_activity,
        "activity_rate": round((recent_activity / total_repos) * 100, 1) if total_repos > 0 else 0
    }
    suggestions = []
    if has_readme_projects < total_repos * 0.5:
        suggestions.append("Add detailed README files to more projects")
    if recent_activity < 3:
        suggestions.append("Increase recent GitHub activity")
    if len(languages_used) < 3:
        suggestions.append("Diversify your tech stack")
    if not topics_used:
        suggestions.append("Add topics/tags to repositories")
    if total_repos < 5:
        suggestions.append("Build more projects")
    return {"insights": insights, "suggestions": suggestions}

def donut_chart(tech_count, soft_count, title):
    labels = ["Technical Skills", "Soft Skills"]
    sizes = [tech_count, soft_count]
    if sum(sizes) == 0:
        sizes = [1, 1]
    fig = plt.figure(figsize=(4, 4))
    ax = fig.add_axes([0.15, 0.15, 0.7, 0.7])
    wedges, _, _ = ax.pie(sizes, startangle=90, autopct="%1.0f%%", radius=1,
                           wedgeprops=dict(width=0.4, edgecolor="white"),
                           colors=['#E8999C', '#FFB4A2'])
    ax.set(aspect="equal")
    ax.set_title(title, pad=10, fontsize=12, fontweight='bold', color='#5A3A3A')
    ax.legend(wedges, labels, loc="lower center", bbox_to_anchor=(0.5, -0.25), ncol=2, frameon=False)
    return fig

def compute_metrics(resume_skills, jd_skills):
    resume_tech = set(resume_skills["technical"])
    resume_soft = set(resume_skills["soft"])
    jd_tech = set(jd_skills["technical"])
    jd_soft = set(jd_skills["soft"])
    matched_skills = (resume_tech & jd_tech) | (resume_soft & jd_soft)
    total_jd_skills = len(jd_tech) + len(jd_soft)
    match_percentage = 0
    if total_jd_skills > 0:
        match_percentage = round((len(matched_skills) / total_jd_skills) * 100, 1)
    return {
        "resume_tech": len(resume_tech),
        "resume_soft": len(resume_soft),
        "jd_tech": len(jd_tech),
        "jd_soft": len(jd_soft),
        "total_resume": len(resume_tech) + len(resume_soft),
        "total_jd": total_jd_skills,
        "match_percent": match_percentage
    }

def categorize_skills(skills):
    categories = {
        "Programming Languages": ["python", "java", "c", "c++", "c#", "go", "rust", "kotlin", "swift",
                                 "javascript", "typescript", "php", "ruby", "r", "matlab", "scala", "perl"],
        "Web Technologies": ["html", "css", "sass", "bootstrap", "tailwind css", "react", "angular",
                            "vue.js", "next.js", "nuxt.js", "node.js", "express.js", "django", "flask",
                            "fastapi", "spring", "spring boot", "laravel", "asp.net"],
        "Databases": ["mysql", "postgresql", "sqlite", "oracle", "sql server", "mongodb", "cassandra",
                     "dynamodb", "redis", "firebase", "elasticsearch", "neo4j", "sql"],
        "Data Science & ML": ["numpy", "pandas", "scipy", "matplotlib", "seaborn", "plotly", "power bi",
                             "tableau", "excel", "machine learning", "deep learning", "tensorflow", "keras",
                             "pytorch", "scikit-learn", "xgboost", "opencv", "nlp", "computer vision",
                             "transformers", "huggingface", "langchain", "llm", "artificial intelligence",
                             "statistics", "data analysis", "data visualization"],
        "Cloud & DevOps": ["aws", "azure", "google cloud", "gcp", "docker", "kubernetes", "terraform",
                          "ansible", "jenkins", "github actions", "ci/cd", "linux", "unix"],
        "Big Data": ["hadoop", "spark", "pyspark", "kafka", "hive", "pig", "hbase", "flink",
                    "airflow", "databricks"],
        "Mobile Dev": ["android", "ios", "flutter", "react native", "swiftui", "kotlin multiplatform"],
        "Testing & QA": ["unit testing", "integration testing", "system testing", "pytest", "unittest",
                        "junit", "selenium", "cypress", "postman"],
        "APIs & Backend": ["rest api", "graphql", "grpc", "soap", "microservices", "jwt", "oauth"],
        "Security": ["network security", "application security", "penetration testing",
                    "ethical hacking", "cryptography", "owasp", "siem", "firewall"],
        "Leadership & Management": ["leadership", "people management", "team leadership", "decision making",
                                   "delegation", "mentoring", "coaching", "strategic thinking"],
        "Communication": ["communication", "verbal communication", "written communication", "public speaking",
                         "presentation skills", "active listening", "storytelling", "business communication"],
        "Problem Solving": ["problem solving", "critical thinking", "analytical thinking", "logical reasoning",
                           "creative thinking", "innovation", "troubleshooting", "root cause analysis"],
        "Work Management": ["time management", "prioritization", "multitasking", "work ethic",
                           "self discipline", "accountability", "goal setting", "organizational skills"],
        "Adaptability": ["adaptability", "flexibility", "resilience", "learning agility",
                        "continuous learning", "open mindedness", "growth mindset"],
        "Collaboration": ["teamwork", "collaboration", "interpersonal skills", "relationship building",
                         "cross functional collaboration", "cultural awareness", "diversity and inclusion",
                         "remote collaboration", "team alignment"],
        "Soft Skills": ["empathy", "emotional intelligence", "conflict resolution", "negotiation",
                       "professionalism", "integrity", "customer focus", "customer service"]
    }
    categorized = {}
    uncategorized = []
    for skill in skills:
        skill_lower = skill.lower()
        found = False
        for category, category_skills in categories.items():
            if skill_lower in category_skills:
                if category not in categorized:
                    categorized[category] = []
                categorized[category].append(skill)
                found = True
                break
        if not found:
            uncategorized.append(skill)
    if uncategorized:
        categorized["Other Skills"] = uncategorized
    return categorized

def build_similarity_matrix(resume_skills, jd_skills):
    if not resume_skills or not jd_skills:
        return np.array([])
    matrix = np.zeros((len(resume_skills), len(jd_skills)))
    for i, r_skill in enumerate(resume_skills):
        for j, j_skill in enumerate(jd_skills):
            if r_skill == j_skill:
                matrix[i][j] = 1.0
            elif r_skill in j_skill or j_skill in r_skill:
                matrix[i][j] = 0.5
            else:
                matrix[i][j] = 0.0
    return matrix

def plot_category_match_heatmap(resume_skills, jd_skills):
    jd_categorized = categorize_skills(jd_skills)
    resume_categorized = categorize_skills(resume_skills)
    common_categories = set(jd_categorized.keys()) & set(resume_categorized.keys())
    if not common_categories:
        st.warning("No common skill categories found")
        return None
    categories = sorted(common_categories)
    match_data = []
    for category in categories:
        jd_cat_skills = jd_categorized[category]
        resume_cat_skills = set(resume_categorized[category])
        matched = 0
        partial = 0
        missing = 0
        for jd_skill in jd_cat_skills:
            if jd_skill in resume_cat_skills:
                matched += 1
            else:
                found_partial = False
                for res_skill in resume_cat_skills:
                    if jd_skill in res_skill or res_skill in jd_skill:
                        partial += 1
                        found_partial = True
                        break
                if not found_partial:
                    missing += 1
        match_data.append({'Category': category, 'Matched': matched, 'Partial': partial, 'Missing': missing})
    fig = go.Figure()
    categories_list = [d['Category'] for d in match_data]
    fig.add_trace(go.Bar(name='Matched', x=categories_list, y=[d['Matched'] for d in match_data],
                         marker_color='#28a745', text=[d['Matched'] for d in match_data], textposition='inside'))
    fig.add_trace(go.Bar(name='Partial Match', x=categories_list, y=[d['Partial'] for d in match_data],
                         marker_color='#ffc107', text=[d['Partial'] for d in match_data], textposition='inside'))
    fig.add_trace(go.Bar(name='Missing', x=categories_list, y=[d['Missing'] for d in match_data],
                         marker_color='#dc3545', text=[d['Missing'] for d in match_data], textposition='inside'))
    fig.update_layout(title="Category-wise Skill Match", barmode='stack', height=500,
                      xaxis_title="Categories", yaxis_title="Skills", hovermode='x unified',
                      legend=dict(orientation="h", y=1.02, x=1, xanchor="right", yanchor="bottom"))
    return fig

def classify_skill_matches(similarity_matrix, resume_skills, jd_skills):
    matched = set()
    partial = set()
    for j, jd_skill in enumerate(jd_skills):
        column = similarity_matrix[:, j]
        if 1.0 in column:
            matched.add(jd_skill)
        elif 0.5 in column:
            partial.add(jd_skill)
    missing = set(jd_skills) - matched - partial
    return {"matched": sorted(list(matched)), "partial": sorted(list(partial)), "missing": sorted(list(missing))}

def calculate_skill_match(resume_skills, jd_skills):
    resume_set = set(resume_skills)
    jd_set = set(jd_skills)
    matched = resume_set & jd_set
    missing = jd_set - resume_set
    partial = set()
    for jd_skill in missing.copy():
        for res_skill in resume_set:
            if jd_skill in res_skill or res_skill in jd_skill:
                partial.add(jd_skill)
    missing = missing - partial
    avg_match = 0
    if len(jd_set) > 0:
        avg_match = round(((len(matched) + 0.5 * len(partial)) / len(jd_set)) * 100, 1)
    return {"matched": len(matched), "partial": len(partial), "missing": len(missing), "avg_match": avg_match}

def generate_word_report(overall_match, matched_skills_count, match_counts, skill_match_result, jd_skills):
    doc = Document()
    title = doc.add_heading('Skill Gap Analysis Report', 0)
    title.alignment = 1
    doc.add_heading('Executive Summary', 1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    summary_data = [
        ('Overall Match Rate', f'{overall_match}%'),
        ('Matched Skills', str(matched_skills_count)),
        ('Partially Matched Skills', str(match_counts["partial"])),
        ('Missing Skills', str(match_counts["missing"]))
    ]
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
    doc.add_paragraph()
    doc.add_heading('Technical Skills Breakdown', 1)
    tech_matched = [s for s in skill_match_result['matched'] if s in jd_skills["technical"]]
    tech_partial = [s for s in skill_match_result['partial'] if s in jd_skills["technical"]]
    tech_missing = [s for s in skill_match_result['missing'] if s in jd_skills["technical"]]
    doc.add_heading('Matched Technical Skills', 2)
    if tech_matched:
        for skill in tech_matched:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No perfectly matched technical skills found.')
    doc.add_heading('Partially Matched Technical Skills', 2)
    if tech_partial:
        for skill in tech_partial:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No partially matched technical skills found.')
    doc.add_heading('Missing Technical Skills', 2)
    if tech_missing:
        for skill in tech_missing:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No missing technical skills!')
    doc.add_heading('Soft Skills Breakdown', 1)
    soft_matched = [s for s in skill_match_result['matched'] if s in jd_skills["soft"]]
    soft_partial = [s for s in skill_match_result['partial'] if s in jd_skills["soft"]]
    soft_missing = [s for s in skill_match_result['missing'] if s in jd_skills["soft"]]
    doc.add_heading('Matched Soft Skills', 2)
    if soft_matched:
        for skill in soft_matched:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No perfectly matched soft skills found.')
    doc.add_heading('Partially Matched Soft Skills', 2)
    if soft_partial:
        for skill in soft_partial:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No partially matched soft skills found.')
    doc.add_heading('Missing Soft Skills', 2)
    if soft_missing:
        for skill in soft_missing:
            doc.add_paragraph(skill, style='List Bullet')
    else:
        doc.add_paragraph('No missing soft skills!')
    doc.add_heading('Recommendations', 1)
    doc.add_paragraph(f'Your current match rate is {overall_match}%. To improve your candidacy, consider focusing on the missing skills listed above.')
    if overall_match >= 70:
        doc.add_paragraph('You have a strong skill match for this position!', style='List Bullet')
    else:
        doc.add_paragraph('Focus on upskilling in the missing areas to increase your match rate.', style='List Bullet')
    doc.add_paragraph()
    footer = doc.add_paragraph(f'Generated on: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
    footer.alignment = 1
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf_report(overall_match, matched_skills_count, match_counts, skill_match_result, jd_skills):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 20)
    pdf.cell(0, 15, "Skill Gap Analysis Report", ln=True, align='C')
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Executive Summary", ln=True)
    pdf.set_font("Arial", '', 11)
    pdf.ln(2)
    pdf.set_fill_color(255, 229, 229)
    pdf.cell(90, 8, "Overall Match Rate:", 1, 0, fill=True)
    pdf.cell(90, 8, f"{overall_match}%", 1, 1)
    pdf.cell(90, 8, "Matched Skills:", 1, 0, fill=True)
    pdf.cell(90, 8, str(matched_skills_count), 1, 1)
    pdf.cell(90, 8, "Partially Matched Skills:", 1, 0, fill=True)
    pdf.cell(90, 8, str(match_counts['partial']), 1, 1)
    pdf.cell(90, 8, "Missing Skills:", 1, 0, fill=True)
    pdf.cell(90, 8, str(match_counts['missing']), 1, 1)
    pdf.ln(10)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 8, "Match Rate Visualization", ln=True)
    pdf.set_font("Arial", '', 10)
    bar_width = 170
    bar_height = 15
    x_start = 20
    y_start = pdf.get_y()
    pdf.set_fill_color(244, 199, 195)
    pdf.rect(x_start, y_start, bar_width, bar_height, 'F')
    filled_width = (overall_match / 100) * bar_width
    pdf.set_fill_color(232, 153, 156)
    pdf.rect(x_start, y_start, filled_width, bar_height, 'F')
    pdf.set_draw_color(215, 122, 125)
    pdf.rect(x_start, y_start, bar_width, bar_height, 'D')
    pdf.set_xy(x_start + bar_width/2 - 10, y_start + 3)
    pdf.set_font("Arial", 'B', 11)
    if overall_match > 30:
        pdf.set_text_color(255, 255, 255)
    else:
        pdf.set_text_color(90, 58, 58)
    pdf.cell(20, 8, f"{overall_match}%", align='C')
    pdf.set_text_color(90, 58, 58)
    pdf.ln(20)
    tech_matched = [s for s in skill_match_result['matched'] if s in jd_skills["technical"]]
    tech_missing = [s for s in skill_match_result['missing'] if s in jd_skills["technical"]]
    pdf.set_font("Arial", 'B', 12)
    pdf.set_fill_color(232, 153, 156)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 8, "Technical Skills - Matched", 0, 1, fill=True)
    pdf.set_text_color(90, 58, 58)
    pdf.set_font("Arial", '', 10)
    if tech_matched:
        for skill in tech_matched[:20]:
            clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, clean_skill, 0, 1)
    else:
        pdf.cell(0, 6, "No technical skills matched.", 0, 1)
    pdf.ln(3)
    pdf.set_font("Arial", 'B', 12)
    pdf.set_fill_color(244, 199, 195)
    pdf.set_text_color(90, 58, 58)
    pdf.cell(0, 8, "Technical Skills - Missing", 0, 1, fill=True)
    pdf.set_text_color(90, 58, 58)
    pdf.set_font("Arial", '', 10)
    if tech_missing:
        for skill in tech_missing[:20]:
            clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, clean_skill, 0, 1)
    else:
        pdf.cell(0, 6, "No missing technical skills!", 0, 1)
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 8, "Recommendations", 0, 1)
    pdf.set_font("Arial", '', 10)
    recommendation = f"Your current match rate is {overall_match}%. "
    if overall_match >= 70:
        recommendation += "You have a strong skill match for this position!"
    elif overall_match >= 50:
        recommendation += "Focus on developing the missing skills."
    else:
        recommendation += "Significant improvement needed. Prioritize missing skills."
    pdf.multi_cell(0, 6, recommendation)
    pdf.ln(3)
    pdf.set_font("Arial", 'I', 9)
    pdf.cell(0, 10, f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 0, 'C')
    return pdf.output(dest="S").encode("latin-1")

# ==================== HEADER ====================
st.markdown("""
<h1 style="text-align:center; color:#E8999C;">📱 SkillGapAI</h1>
<p style="text-align:center; font-size:18px; color:#6B7280;">
AI-Powered Skill Gap Analysis & Career Insights
</p>
<hr>
""", unsafe_allow_html=True)

# ==================== RENDER TIMELINE ====================
render_timeline()

# ==================== STEP 1: UPLOAD ====================
if st.session_state.step == 1:
    st.markdown("""
        <div style="background: linear-gradient(135deg, #E8999C 0%, #F4C7C3 100%);
                    padding:25px;border-radius:15px;
                    box-shadow: 0 4px 15px rgba(232, 153, 156, 0.3);">
            <h2 style="color:white;">Step 1: Upload Documents</h2>
            <p style="color:white;">Upload your resume and job description to begin analysis</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.write("")
    
    left_col, right_col = st.columns(2)
    
    with left_col:
        resume_file = st.file_uploader(
            "Upload Resume (Multi-page supported)",
            type=["pdf", "docx", "txt"],
            help="Upload your resume in PDF, DOCX, or TXT format.",
            key="resume_upload"
        )
        if resume_file:
            st.session_state.resume_file = resume_file
            st.success(f"Resume uploaded: {resume_file.name}")
    
    with right_col:
        jd_file = st.file_uploader(
            "Upload Job Description (Multi-page supported)",
            type=["pdf", "docx", "txt"],
            help="Upload job description in PDF, DOCX, or TXT format.",
            key="jd_upload"
        )
        if jd_file:
            st.session_state.jd_file = jd_file
            st.success(f"Job Description uploaded: {jd_file.name}")
    
    st.write("")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.session_state.resume_file and st.session_state.jd_file:
            if st.button("Continue to Extraction ➡️", use_container_width=True, type="primary"):
                next_step()
        else:
            st.warning("⚠️ Please upload both files to continue")

# ==================== STEP 2: EXTRACT TEXT ====================
elif st.session_state.step == 2:
    st.markdown("""
        <div style="background: linear-gradient(135deg, #E8999C 0%, #F4C7C3 100%);
                    padding:25px;border-radius:15px;
                    box-shadow: 0 4px 15px rgba(232, 153, 156, 0.3);">
            <h2 style="color:white;">Step 2: Text Extraction</h2>
            <p style="color:white;">Extracting and previewing text from documents</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.write("")
    
    if not st.session_state.resume_text:
        with st.spinner("Extracting resume text..."):
            st.session_state.resume_text = extract_text(st.session_state.resume_file)
    
    if not st.session_state.jd_text:
        with st.spinner("Extracting job description text..."):
            st.session_state.jd_text = extract_text(st.session_state.jd_file)
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.session_state.resume_text:
            char_count = len(st.session_state.resume_text)
            word_count = len(st.session_state.resume_text.split())
            st.caption(f"📊 Resume: {char_count} characters, {word_count} words")
            st.text_area("Resume Preview", st.session_state.resume_text[:2000] + ("..." if len(st.session_state.resume_text) > 2000 else ""),
                        height=200, disabled=True, key="resume_prev")
    
    with col2:
        if st.session_state.jd_text:
            char_count = len(st.session_state.jd_text)
            word_count = len(st.session_state.jd_text.split())
            st.caption(f"📊 Job Description: {char_count} characters, {word_count} words")
            st.text_area("JD Preview", st.session_state.jd_text[:2000] + ("..." if len(st.session_state.jd_text) > 2000 else ""),
                        height=200, disabled=True, key="jd_prev")
    
    st.success("✅ Text extraction complete!")
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("⬅️ Back", use_container_width=True):
            prev_step()
    with col3:
        if st.button("Continue to Cleaning ➡️", use_container_width=True, type="primary"):
            next_step()

# ==================== STEP 3: CLEAN TEXT ====================
elif st.session_state.step == 3:
    st.markdown("""
        <div style="background: linear-gradient(135deg, #E8999C 0%, #F4C7C3 100%);
                    padding:25px;border-radius:15px;
                    box-shadow: 0 4px 15px rgba(232, 153, 156, 0.3);">
            <h2 style="color:white;">Step 3: Text Cleaning & Processing</h2>
            <p style="color:white;">Cleaning and normalizing text for NLP analysis</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.write("")
    
    if not st.session_state.cleaned_resume:
        with st.spinner("Cleaning resume text..."):
            st.session_state.cleaned_resume = clean(st.session_state.resume_text)
            st.session_state.resume_experience = extract_experience(st.session_state.resume_text)
    
    if not st.session_state.cleaned_jd:
        with st.spinner("Cleaning job description text..."):
            st.session_state.cleaned_jd = clean(st.session_state.jd_text)
            st.session_state.jd_experience = extract_experience(st.session_state.jd_text)
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.session_state.cleaned_resume:
            st.caption(f"Cleaned Resume: {len(st.session_state.cleaned_resume)} characters")
            st.text_area("Cleaned Resume", st.session_state.cleaned_resume[:2000] + ("..." if len(st.session_state.cleaned_resume) > 2000 else ""),
                        height=200, disabled=True, key="clean_resume")
    
    with col2:
        if st.session_state.cleaned_jd:
            st.caption(f"Cleaned Job Description: {len(st.session_state.cleaned_jd)} characters")
            st.text_area("Cleaned JD", st.session_state.cleaned_jd[:2000] + ("..." if len(st.session_state.cleaned_jd) > 2000 else ""),
                        height=200, disabled=True, key="clean_jd")
    
    st.subheader("Experience Analysis")
    col1, col2 = st.columns(2)
    with col1:
        if st.session_state.resume_experience:
            st.info(f"Resume Experience: {st.session_state.resume_experience['min_exp']}" +
                   (f"-{st.session_state.resume_experience['max_exp']}" if st.session_state.resume_experience['max_exp'] else "+") + " years")
        else:
            st.warning("No experience information found")
    with col2:
        if st.session_state.jd_experience:
            st.info(f"JD Required Experience: {st.session_state.jd_experience['min_exp']}" +
                   (f"-{st.session_state.jd_experience['max_exp']}" if st.session_state.jd_experience['max_exp'] else "+") + " years")
        else:
            st.warning("No experience requirement found")
    
    st.success("✅ Text cleaning complete!")
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("⬅️ Back", use_container_width=True):
            prev_step()
    with col3:
        if st.button("Continue to Skill Extraction ➡️", use_container_width=True, type="primary"):
            next_step()

# ==================== STEP 4: SKILL EXTRACTION ====================
elif st.session_state.step == 4:
    st.markdown("""
        <div style="background: linear-gradient(135deg, #E8999C 0%, #F4C7C3 100%);
                    padding:25px;border-radius:15px;
                    box-shadow: 0 4px 15px rgba(232, 153, 156, 0.3);">
            <h2 style="color:white;">Step 4: Skill Extraction using NLP</h2>
            <p style="color:white;">Extracting technical and soft skills using AI</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.write("")
    
    nlp = load_spacy_model()
    if nlp is None:
        st.stop()
    
    if not st.session_state.resume_skills:
        with st.spinner("Extracting skills using NLP..."):
            resume_technical = extract_technical_skills(st.session_state.cleaned_resume, master_technical_skills, nlp)
            resume_soft = extract_soft_skills(st.session_state.cleaned_resume, master_soft_skills)
            
            github_skills = []
            github_analysis = None
            detected_github_url = detect_github_in_resume(st.session_state.resume_text)
            
            if detected_github_url:
                st.success(f"GitHub profile detected: {detected_github_url}")
                username = extract_github_username(detected_github_url)
                if username:
                    repos = fetch_github_repos(username)
                    if repos:
                        github_skills = extract_github_skills(repos, master_technical_skills)
                        github_analysis = analyze_github_profile(repos)
                        st.session_state.github_skills = github_skills
                        st.session_state.github_analysis = github_analysis
            
            combined_technical_skills = sorted(set(resume_technical) | set(github_skills))
            st.session_state.resume_skills = {"technical": combined_technical_skills, "soft": resume_soft}
    
    if not st.session_state.jd_skills:
        with st.spinner("Extracting JD skills..."):
            jd_technical = extract_technical_skills(st.session_state.cleaned_jd, master_technical_skills, nlp)
            jd_soft = extract_soft_skills(st.session_state.cleaned_jd, master_soft_skills)
            st.session_state.jd_skills = {"technical": jd_technical, "soft": jd_soft}
    
    st.subheader("Extracted Skills")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### Resume Skills")
        with st.expander("Technical Skills", expanded=True):
            if st.session_state.resume_skills["technical"]:
                st.write(", ".join(st.session_state.resume_skills["technical"]))
                st.caption(f"Total: {len(st.session_state.resume_skills['technical'])} skills")
            else:
                st.info("No technical skills found")
        with st.expander("Soft Skills", expanded=True):
            if st.session_state.resume_skills["soft"]:
                st.write(", ".join(st.session_state.resume_skills["soft"]))
                st.caption(f"Total: {len(st.session_state.resume_skills['soft'])} skills")
            else:
                st.info("No soft skills found")
    
    with col2:
        st.markdown("### JD Skills")
        with st.expander("Technical Skills", expanded=True):
            if st.session_state.jd_skills["technical"]:
                st.write(", ".join(st.session_state.jd_skills["technical"]))
                st.caption(f"Total: {len(st.session_state.jd_skills['technical'])} skills")
            else:
                st.info("No technical skills found")
        with st.expander("Soft Skills", expanded=True):
            if st.session_state.jd_skills["soft"]:
                st.write(", ".join(st.session_state.jd_skills["soft"]))
                st.caption(f"Total: {len(st.session_state.jd_skills['soft'])} skills")
            else:
                st.info("No soft skills found")
    
    st.subheader("Skill Distribution Analysis")
    col1, col2 = st.columns(2)
    
    with col1:
        fig1 = donut_chart(len(st.session_state.resume_skills["technical"]),
                          len(st.session_state.resume_skills["soft"]),
                          "Resume Skill Distribution")
        st.pyplot(fig1, use_container_width=False)
        plt.close(fig1)
    
    with col2:
        fig2 = donut_chart(len(st.session_state.jd_skills["technical"]),
                          len(st.session_state.jd_skills["soft"]),
                          "JD Skill Distribution")
        st.pyplot(fig2, use_container_width=False)
        plt.close(fig2)
    
    metrics = compute_metrics(st.session_state.resume_skills, st.session_state.jd_skills)
    
    st.subheader("Skill Extraction Metrics")
    col1, col2, col3, col4 = st.columns(4)
    col1.metric(label="JD Technical Skills", value=metrics["jd_tech"])
    col2.metric(label="JD Soft Skills", value=metrics["jd_soft"])
    col3.metric(label="Total JD Skills", value=metrics["total_jd"])
    col4.metric(label="Basic Match %", value=f'{metrics["match_percent"]}%')
    
    st.success("✅ Skills extracted successfully!")
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("⬅️ Back", use_container_width=True):
            prev_step()
    with col3:
        if st.button("Continue to Analysis ➡️", use_container_width=True, type="primary"):
            next_step()

# ==================== STEP 5: SKILL GAP ANALYSIS ====================
elif st.session_state.step == 5:
    st.markdown("""
        <div style="background: linear-gradient(135deg, #E8999C 0%, #F4C7C3 100%);
                    padding:25px;border-radius:15px;
                    box-shadow: 0 4px 15px rgba(232, 153, 156, 0.3);">
            <h2 style="color:white;">Step 5: Skill Gap Analysis</h2>
            <p style="color:white;">Analyzing skill matches and identifying gaps</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.write("")
    
    resume_all_skills = st.session_state.resume_skills["technical"] + st.session_state.resume_skills["soft"]
    jd_all_skills = st.session_state.jd_skills["technical"] + st.session_state.jd_skills["soft"]
    
    if st.session_state.similarity_matrix is None:
        with st.spinner("Building skill gap analysis..."):
            st.session_state.similarity_matrix = build_similarity_matrix(resume_all_skills, jd_all_skills)
    
    if resume_all_skills and jd_all_skills:
        st.subheader("Category-wise Skill Match Heatmap")
        fig_heatmap = plot_category_match_heatmap(resume_all_skills, jd_all_skills)
        if fig_heatmap:
            st.plotly_chart(fig_heatmap, use_container_width=True)
    
    if st.session_state.skill_match_result is None:
        st.session_state.skill_match_result = classify_skill_matches(st.session_state.similarity_matrix, resume_all_skills, jd_all_skills)
    
    if st.session_state.match_counts is None:
        all_resume_skills = set(st.session_state.resume_skills["technical"]) | set(st.session_state.resume_skills["soft"])
        all_jd_skills = set(st.session_state.jd_skills["technical"]) | set(st.session_state.jd_skills["soft"])
        st.session_state.match_counts = calculate_skill_match(all_resume_skills, all_jd_skills)
    
    st.subheader("Skill Match Overview")
    
    left, right = st.columns([1.3, 1])
    
    with left:
        labels = ["Matched", "Partially Matched", "Missing"]
        sizes = [st.session_state.match_counts["matched"], st.session_state.match_counts["partial"], st.session_state.match_counts["missing"]]
        colors = ["#E8999C", "#FFB4A2", "#F4C7C3"]
        if sum(sizes) == 0:
            sizes = [1, 1, 1]
        fig = plt.figure(figsize=(3.6, 3.6))
        ax = fig.add_axes([0.1, 0.15, 0.8, 0.75])
        wedges, _, autotexts = ax.pie(sizes, startangle=90, autopct="%1.0f%%", radius=1, colors=colors, wedgeprops=dict(width=0.4, edgecolor="white"))
        for autotext in autotexts:
            autotext.set_color('#5A3A3A')
            autotext.set_fontweight('bold')
            autotext.set_fontsize(11)
        ax.set(aspect="equal")
        ax.set_title("Skill Match Distribution", pad=8, fontsize=12, fontweight='bold', color='#5A3A3A')
        ax.legend(wedges, labels, loc="lower center", bbox_to_anchor=(0.5, -0.25), ncol=3, frameon=False)
        st.pyplot(fig, use_container_width=False)
        plt.close(fig)
    
    with right:
        with st.container():
            r1, r2 = st.columns(2)
            r1.metric("Matched Skills", st.session_state.match_counts["matched"])
            r2.metric("Partially Matched", st.session_state.match_counts["partial"])
            r3, r4 = st.columns(2)
            r3.metric("Missing Skills", st.session_state.match_counts["missing"])
            r4.metric("Avg Match %", f'{st.session_state.match_counts["avg_match"]}%')
    
    st.subheader("Skill Gap Details")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("### Matched Skills")
        if st.session_state.skill_match_result["matched"]:
            for skill in st.session_state.skill_match_result["matched"][:10]:
                st.success(f"✓ {skill}")
            if len(st.session_state.skill_match_result["matched"]) > 10:
                st.caption(f"...and {len(st.session_state.skill_match_result['matched'])-10} more")
        else:
            st.info("No perfectly matched skills")
    
    with col2:
        st.markdown("### Partial Matches")
        if st.session_state.skill_match_result["partial"]:
            for skill in st.session_state.skill_match_result["partial"][:10]:
                st.warning(f"≈ {skill}")
            if len(st.session_state.skill_match_result["partial"]) > 10:
                st.caption(f"...and {len(st.session_state.skill_match_result['partial'])-10} more")
        else:
            st.info("No partially matched skills")
    
    with col3:
        st.markdown("### Missing Skills")
        if st.session_state.skill_match_result["missing"]:
            for skill in st.session_state.skill_match_result["missing"][:10]:
                st.error(f"✗ {skill}")
            if len(st.session_state.skill_match_result["missing"]) > 10:
                st.caption(f"...and {len(st.session_state.skill_match_result['missing'])-10} more")
        else:
            st.success("No missing skills!")
    
    st.success("✅ Skill gap analysis complete!")
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("⬅️ Back", use_container_width=True):
            prev_step()
    with col3:
        if st.button("Continue to Reports ➡️", use_container_width=True, type="primary"):
            next_step()

# ==================== STEP 6: DASHBOARD & REPORTS ====================
elif st.session_state.step == 6:
    st.markdown("""
        <div style="background: linear-gradient(135deg, #E8999C 0%, #F4C7C3 100%);
                    padding:25px;border-radius:15px;
                    box-shadow: 0 4px 15px rgba(232, 153, 156, 0.3);">
            <h2 style="color:white;">Step 6: Dashboard & Report Export</h2>
            <p style="color:white;">View comprehensive results and download reports</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.write("")
    
    all_resume_skills = set(st.session_state.resume_skills["technical"]) | set(st.session_state.resume_skills["soft"])
    all_jd_skills = set(st.session_state.jd_skills["technical"]) | set(st.session_state.jd_skills["soft"])
    overall_match = st.session_state.match_counts["avg_match"]
    matched_skills_count = st.session_state.match_counts["matched"]
    missing_skills_count = st.session_state.match_counts["missing"]
    
    c1, c2, c3 = st.columns(3)
    c1.metric("Overall Match", f"{overall_match}%")
    c2.metric("Matched Skills", str(matched_skills_count))
    c3.metric("Missing Skills", str(missing_skills_count))
    
    st.subheader("Categorized Skill Match Overview")
    jd_categorized = categorize_skills(list(all_jd_skills))
    categories_to_plot = []
    resume_category_scores = []
    jd_category_scores = []
    
    for category, cat_skills in jd_categorized.items():
        matched_in_category = sum(1 for skill in cat_skills if skill in all_resume_skills)
        total_in_category = len(cat_skills)
        if total_in_category > 0:
            categories_to_plot.append(category)
            resume_category_scores.append((matched_in_category / total_in_category) * 100)
            jd_category_scores.append(100)
    
    bar_fig = go.Figure()
    bar_fig.add_trace(go.Bar(name='Your Skills', x=categories_to_plot, y=resume_category_scores, marker_color='#E8999C',
                             text=[f'{score:.0f}%' for score in resume_category_scores], textposition='outside'))
    bar_fig.add_trace(go.Bar(name='Job Requirements', x=categories_to_plot, y=jd_category_scores, marker_color='#28a745',
                             text=['100%'] * len(categories_to_plot), textposition='outside'))
    bar_fig.update_layout(title="Skill Match by Category", barmode="group", height=500, xaxis_title="Skill Categories",
                         yaxis_title="Coverage (%)", hovermode='x unified', yaxis=dict(range=[0, 120]),
                         legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1))
    st.plotly_chart(bar_fig, use_container_width=True)
    
    st.subheader("Upskilling Recommendations")
    st.caption("Based on missing and partially matched skills")
    
    if st.session_state.skill_match_result["missing"]:
        st.markdown("**🔴 Priority Skills to Learn (Missing):**")
        for i, skill in enumerate(st.session_state.skill_match_result["missing"][:5], 1):
            st.error(f"{i}. **{skill.title()}** - Not found in your resume")
    
    if st.session_state.skill_match_result["partial"]:
        st.markdown("**🟡 Skills to Strengthen (Partial Match):**")
        for i, skill in enumerate(st.session_state.skill_match_result["partial"][:5], 1):
            st.warning(f"{i}. **{skill.title()}** - Improve proficiency")
    
    if not st.session_state.skill_match_result["missing"] and not st.session_state.skill_match_result["partial"]:
        st.success("Excellent! You have all the required skills!")
    
    st.subheader("Report Download")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        tech_matched = [s for s in st.session_state.skill_match_result['matched'] if s in st.session_state.jd_skills["technical"]]
        tech_partial = [s for s in st.session_state.skill_match_result['partial'] if s in st.session_state.jd_skills["technical"]]
        tech_missing = [s for s in st.session_state.skill_match_result['missing'] if s in st.session_state.jd_skills["technical"]]
        soft_matched = [s for s in st.session_state.skill_match_result['matched'] if s in st.session_state.jd_skills["soft"]]
        soft_partial = [s for s in st.session_state.skill_match_result['partial'] if s in st.session_state.jd_skills["soft"]]
        soft_missing = [s for s in st.session_state.skill_match_result['missing'] if s in st.session_state.jd_skills["soft"]]
        
        max_tech_len = max(len(tech_matched), len(tech_partial), len(tech_missing), 1)
        tech_matched += [''] * (max_tech_len - len(tech_matched))
        tech_partial += [''] * (max_tech_len - len(tech_partial))
        tech_missing += [''] * (max_tech_len - len(tech_missing))
        
        max_soft_len = max(len(soft_matched), len(soft_partial), len(soft_missing), 1)
        soft_matched += [''] * (max_soft_len - len(soft_matched))
        soft_partial += [''] * (max_soft_len - len(soft_partial))
        soft_missing += [''] * (max_soft_len - len(soft_missing))
        
        tech_df = pd.DataFrame({"Technical - Matched": tech_matched, "Technical - Partially Matched": tech_partial, "Technical - Missing": tech_missing})
        soft_df = pd.DataFrame({"Soft Skills - Matched": soft_matched, "Soft Skills - Partially Matched": soft_partial, "Soft Skills - Missing": soft_missing})
        csv_output = tech_df.to_csv(index=False) + "\n" + soft_df.to_csv(index=False)
        
        st.download_button("⬇Download CSV Report", csv_output.encode("utf-8"), "skill_gap_report.csv", "text/csv",
                          help="Download skill comparison", use_container_width=True)
    
    with col2:
        doc_file = generate_word_report(overall_match, matched_skills_count, st.session_state.match_counts,
                                       st.session_state.skill_match_result, st.session_state.jd_skills)
        st.download_button("⬇Download DOCX Report", doc_file, "skill_gap_analysis_report.docx",
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          help="Download formatted Word report", use_container_width=True)
    
    with col3:
        try:
            pdf_bytes = generate_pdf_report(overall_match, matched_skills_count, st.session_state.match_counts,
                                           st.session_state.skill_match_result, st.session_state.jd_skills)
            st.download_button("⬇Download PDF Report", pdf_bytes, "skill_gap_analysis_report.pdf", "application/pdf",
                              help="Download formatted PDF report", use_container_width=True)
        except Exception as e:
            st.error(f"Error generating PDF: {e}")
    
    st.success("✅ Dashboard loaded successfully with export options!")
    
    st.markdown("---")
    st.markdown("""
    ### Analysis Complete!
    
    **Next Steps:**
    1. Review the missing skills identified
    2. Consider upskilling in those areas
    3. Update your resume with newly acquired skills
    4. Download the report for your records
    
    **Pro Tip:** Aim for at least 70% match rate for better job prospects!
    """)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        if st.button("⬅️ Back to Analysis", use_container_width=True):
            prev_step()
    with col3:
        if st.button("🔄 Start New Analysis", use_container_width=True, type="primary"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()

# ==================== FOOTER ====================
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #6B7280; padding: 20px;">
    <p>📱 SkillGapAI Mobile - Step-by-Step Analysis</p>
    <p style="font-size: 0.9rem;">Multi-page support • AI-powered • Comprehensive reports • Reddish Pastel Theme</p>
</div>
""", unsafe_allow_html=True)