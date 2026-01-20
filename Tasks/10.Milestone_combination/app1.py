import streamlit as st
import pdfplumber
import docx
import re
import spacy
from spacy.matcher import PhraseMatcher
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
import io

# Milestone - 1

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="Skill Gap Analysis Dashboard",
    layout="wide"
)

# ---------------- HEADER ----------------
st.markdown(
    """
    <div style="background-color:#4a5bdc;padding:20px;border-radius:10px">
        <h2 style="color:white;">
            Milestone 1: Data Ingestion, Parsing and Cleaning Module 
        </h2>
        <p style="color:white;">
            This module allows uploading resumes and job descriptions,
            extracting text from multiple document formats, and displaying
            clean, normalized content for further processing.
        </p>
    </div>
    """,
    unsafe_allow_html=True
)

st.write("")

# ---------------- INITIALIZE ----------------
resume_text = None
jd_text = None
cleaned_resume = ""
cleaned_jd = ""
resume_experience = None
jd_experience = None

# ---------------- FUNCTIONS ----------------
def extract_text(file):
    if file.type == "application/pdf":
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                if page.extract_text():
                    text += page.extract_text() + "\n"
        return text

    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        document = docx.Document(file)
        return "\n".join([para.text for para in document.paragraphs])

    elif file.type == "text/plain":
        return file.read().decode("utf-8")

    return ""

def clean(text):
    if not text:
        return ""

    text = text.lower()
    text = re.sub(r'(\d+)\s*(?:-|–|to)\s*(\d+)', r'\1_\2', text)
    text = re.sub(r'(\d+)\s*\+', r'\1_plus', text)
    text = re.sub(r'\n|\t', ' ', text)
    text = re.sub(r'[^a-z0-9_\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()

    return text

def extract_experience(text):
    if not text:
        return None

    text = text.lower()

    # experience range: 1-3 years, 2 to 5 yrs
    range_match = re.search(
        r'(\d+)\s*(?:-|–|to)\s*(\d+)\s*(?:years?|yrs?|year experience|years experience)',
        text
    )
    if range_match:
        return {
            "min_exp": int(range_match.group(1)),
            "max_exp": int(range_match.group(2))
        }

    # experience plus: 3+ years
    plus_match = re.search(
        r'(\d+)\s*\+\s*(?:years?|yrs?|year experience|years experience)',
        text
    )
    if plus_match:
        return {
            "min_exp": int(plus_match.group(1)),
            "max_exp": None
        }

    return None


# ---------------- UPLOAD ----------------
st.subheader("Upload Documents")

left_col, right_col = st.columns(2)

with left_col:
    resume_file = st.file_uploader(
        "Upload Resume",
        type=["pdf", "docx", "txt"]
    )

with right_col:
    jd_file = st.file_uploader(
        "Upload Job Description",
        type=["pdf", "docx", "txt"]
    )

if jd_file and resume_file:

    # ---------------- PREVIEW ----------------
    st.subheader("Parsed Document Preview")

    col1, col2 = st.columns(2)

    with col1:
        if resume_file:
            resume_text = extract_text(resume_file)
            st.text_area("Resume Preview", resume_text, height=300)

    with col2:
        if jd_file:
            jd_text = extract_text(jd_file)
            st.text_area("Job Description Preview", jd_text, height=300)

    # ---------------- PROCESS ----------------
    if resume_text:
        cleaned_resume = clean(resume_text)
        resume_experience = extract_experience(resume_text)

    if jd_text:
        cleaned_jd = clean(jd_text)
        jd_experience = extract_experience(jd_text)

    # ---------------- CLEANED OUTPUT ----------------
    st.subheader("Cleaned Files")

    col1, col2 = st.columns(2)

    with col1:
        if cleaned_resume:
            st.text_area("Cleaned Resume", cleaned_resume, height=200)
    with col2:
        if cleaned_jd:
            st.text_area("Cleaned Job Description", cleaned_jd, height=200)

    st.success(
        "Milestone 1 Completed: Documents uploaded, parsed, cleaned and previewed successfully."
    )


#-----------------------------------------------------------------------------------------------------------------------------------

    # Milestone - 2

    st.markdown(
        """
        <div style="background-color:#4a5bdc;padding:20px;border-radius:10px">
            <h2 style="color:white;">
                Milestone 2: Skill Extraction using NLP Module
            </h2>
            <p style="color:white;">
                Module: Skill Extraction using NLP <br>
                • spaCy-based pipelines <br>
                • Technical and soft skills identification <br>
                • Structured skill display
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    # ---------------- Experience ----------------
    st.subheader("Experience")

    col1, col2 = st.columns(2)
    with col1:
        st.write("Resume Experience found:", resume_experience)
    with col2:
        st.write("JD Experience found:", jd_experience)

    # ---------------- NLP Model ----------------
    import spacy
    from spacy.matcher import PhraseMatcher

    nlp = spacy.load("en_core_web_trf")

    # ---------------- Skill Masters ----------------
    master_technical_skills = {
        "python","java","c","c++","javascript","react","spring boot","node.js",
        "mysql","mongodb","postgresql","aws","docker","kubernetes",
        "pandas","numpy","data analysis","machine learning","nlp"
    }

    master_soft_skills = {
        "communication","teamwork","leadership","problem solving",
        "time management","adaptability","critical thinking"
    }

    # ---------------- Skill Extraction Functions ----------------
    def extract_technical_skills(text, skill_set):
        if not text:
            return []
        matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
        patterns = [nlp.make_doc(skill) for skill in skill_set]
        matcher.add("TECH", patterns)
        doc = nlp(text)
        return sorted({doc[start:end].text.lower() for _, start, end in matcher(doc)})

    def extract_soft_skills(text, skill_set):
        if not text:
            return []
        text = text.lower()
        return sorted({skill for skill in skill_set if skill in text})

    # ---------------- Extract Skills ----------------
    resume_skills = {
        "technical": extract_technical_skills(cleaned_resume, master_technical_skills),
        "soft": extract_soft_skills(cleaned_resume, master_soft_skills)
    }

    jd_skills = {
        "technical": extract_technical_skills(cleaned_jd, master_technical_skills),
        "soft": extract_soft_skills(cleaned_jd, master_soft_skills)
    }

    # ---------------- Skill Comparison ----------------
    resume_all = set(resume_skills["technical"]) | set(resume_skills["soft"])
    jd_all = set(jd_skills["technical"]) | set(jd_skills["soft"])

    matched_skills = sorted(resume_all & jd_all)
    missing_skills = sorted(jd_all - resume_all)
    extra_skills = sorted(resume_all - jd_all)

    # ---------------- Extracted Skills UI ----------------
    st.markdown("<h2 style='text-align:center;'>Extracted Skills</h2>", unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 👤 Resume Skills")
        st.markdown("**Technical**")
        for s in resume_skills["technical"]:
            st.success(s)

        st.markdown("**Soft Skills**")
        for s in resume_skills["soft"]:
            st.success(s)

    with col2:
        st.markdown("### 📄 JD Skills")
        st.markdown("**Technical**")
        for s in jd_skills["technical"]:
            st.success(s)

        st.markdown("**Soft Skills**")
        for s in jd_skills["soft"]:
            st.success(s)

    # ---------------- Matched / Missing / Extra ----------------
    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("<div style='background:#e8f5e9;padding:15px;border-radius:10px'><h4>✅ Matched Skills</h4></div>", unsafe_allow_html=True)
        for s in matched_skills:
            st.success(s)

    with col2:
        st.markdown("<div style='background:#ffebee;padding:15px;border-radius:10px'><h4>⚠️ Missing Skills</h4></div>", unsafe_allow_html=True)
        for s in missing_skills:
            st.error(s)

    with col3:
        st.markdown("<div style='background:#e3f2fd;padding:15px;border-radius:10px'><h4>➕ Extra Skills</h4></div>", unsafe_allow_html=True)
        for s in extra_skills:
            st.info(s)

    # ---------------- Metrics ----------------
    match_percent = round((len(matched_skills) / len(jd_all)) * 100, 1) if jd_all else 0

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Resume Skills", len(resume_all))
    m2.metric("JD Skills", len(jd_all))
    m3.metric("Matched Skills", len(matched_skills))
    m4.metric("Match %", f"{match_percent}%")





#------------------------------------------------------------------------------------------------------------------------------------------
    # Milestone - 3

    def build_similarity_matrix(resume_skills, jd_skills):
        matrix = np.zeros((len(resume_skills), len(jd_skills)))

        for i, r_skill in enumerate(resume_skills):
            for j, j_skill in enumerate(jd_skills):
                if r_skill == j_skill:
                    matrix[i][j] = 1.0
                elif r_skill in j_skill or j_skill in r_skill:
                    matrix[i][j] = 0.5
                else:
                    matrix[i][j] = 0.0

        return matrix

    def plot_similarity_heatmap(matrix, resume_labels, jd_labels):
        fig, ax = plt.subplots(figsize=(8, 5))

        im = ax.imshow(matrix, cmap="YlGn")

        ax.set_xticks(range(len(jd_labels)))
        ax.set_yticks(range(len(resume_labels)))

        ax.set_xticklabels(jd_labels, rotation=45, ha="right")
        ax.set_yticklabels(resume_labels)

        ax.set_title("Skill Similarity Matrix")

        cbar = fig.colorbar(im, ax=ax)
        cbar.set_label("Similarity Score")

        st.pyplot(fig)

    resume_all_skills = resume_skills["technical"] + resume_skills["soft"]
    jd_all_skills = jd_skills["technical"] + jd_skills["soft"]

    st.markdown(
        """
        <div style="background-color:#4a5bdc;padding:20px;border-radius:10px">
            <h2 style="color:white;">
                Milestone 3: Skill Gap Analysis and Similarity Matching Module
            </h2>
            <p style="color:white;">
                • Skill similarity matrix visualization <br>
                • Resume vs JD skill comparison <br>
                • Missing skill identification
            </p>
        </div>
        <br>
        """,
        unsafe_allow_html=True
    )

    if resume_all_skills and jd_all_skills:
        similarity_matrix = build_similarity_matrix(
            resume_all_skills,
            jd_all_skills
        )

        plot_similarity_heatmap(
            similarity_matrix,
            resume_all_skills,
            jd_all_skills
        )
    else:
        st.warning("Insufficient skills to build similarity matrix")

    missing_skills = list(set(jd_all_skills) - set(resume_all_skills))

    def classify_skill_matches(similarity_matrix, resume_skills, jd_skills):
        matched = set()
        partial = set()

        for j, jd_skill in enumerate(jd_skills):
            column = similarity_matrix[:, j]

            if 1.0 in column:
                matched.add(jd_skill)
            elif 0.5 in column:
                partial.add(jd_skill)

        missing = set(jd_skills) - matched - partial

        return {
            "matched": list(matched),
            "partial": list(partial),
            "missing": list(missing)
        }

    skill_match_result = classify_skill_matches(
        similarity_matrix,
        resume_all_skills,
        jd_all_skills
    )

    def calculate_skill_match(resume_skills, jd_skills):
        resume_set = set(resume_skills)
        jd_set = set(jd_skills)

        matched = resume_set & jd_set
        missing = jd_set - resume_set

        # Partial = same base word (simple heuristic)
        partial = set()
        for jd_skill in missing.copy():
            for res_skill in resume_set:
                if jd_skill in res_skill or res_skill in jd_skill:
                    partial.add(jd_skill)

        missing = missing - partial

        avg_match = 0
        if len(jd_set) > 0:
            avg_match = round(
                ((len(matched) + 0.5 * len(partial)) / len(jd_set)) * 100,
                1
            )

        return {
            "matched": len(matched),
            "partial": len(partial),
            "missing": len(missing),
            "avg_match": avg_match
        }

    all_resume_skills = (
        set(resume_skills["technical"]) |
        set(resume_skills["soft"])
    )

    all_jd_skills = (
        set(jd_skills["technical"]) |
        set(jd_skills["soft"])
    )

    match_counts = calculate_skill_match(
        all_resume_skills,
        all_jd_skills
    )


    st.subheader("Skill Match Overview")

    left, right = st.columns([1.3, 1])

    # ---------------- LEFT: DONUT ----------------
    with left:
        labels = ["Matched", "Partially Matched", "Missing"]
        sizes = [match_counts["matched"], match_counts["partial"], match_counts["missing"]]
        colors = ["#28a745", "#fd7e14", "#dc3545"]

        if sum(sizes) == 0:
            sizes = [1, 1, 1]

        fig = plt.figure(figsize=(3.6, 3.6))
        ax = fig.add_axes([0.1, 0.15, 0.8, 0.75])

        wedges, _, _ = ax.pie(
            sizes,
            startangle=90,
            autopct="%1.0f%%",
            radius=1,
            colors=colors,
            wedgeprops=dict(width=0.4, edgecolor="white")
        )

        ax.set(aspect="equal")
        ax.set_title("Skill Match Distribution", pad=8)
        ax.legend(
            wedges,
            labels,
            loc="lower center",
            bbox_to_anchor=(0.5, -0.25),
            ncol=3,
            frameon=False
        )

        st.pyplot(fig, use_container_width=False)

    # ---------------- RIGHT: METRICS ----------------
    with right:
        with st.container():
            r1, r2 = st.columns(2)
            r1.metric("Matched Skills", match_counts["matched"])
            r2.metric("Partially Matched", match_counts["partial"])

            r3, r4 = st.columns(2)
            r3.metric("Missing Skills", match_counts["missing"])
            r4.metric("Avg Match %", f'{match_counts["avg_match"]}%')



    st.subheader("Skill Gap Details")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("### Matched Skills")
        st.write(skill_match_result["matched"])

    with col2:
        st.markdown("### Partial Matches")
        st.write(skill_match_result["partial"])

    with col3:
        st.markdown("### Missing Skills")
        st.write(skill_match_result["missing"])

#---------------------------------------------------------------------------------------------------------------------------------------------
    # Milestone - 4


    st.markdown(
        """
        <div style="background-color:#4a5bdc;padding:20px;border-radius:10px">
            <h2 style="color:white;">
                Milestone 4: Dashboard and Report Export Module
            </h2>
            <p style="color:white;">
                Interactive dashboard • Graphs • Multi-format report export
            </p>
        </div>
        <br>
        """,
        unsafe_allow_html=True
    )

    st.write("")

    # ---------------- DATA ----------------
    skills = list(all_jd_skills)
    resume_scores = [
        100 if skill in all_resume_skills else 0
        for skill in skills
    ]

    job_scores = [100 for _ in skills]

    overall_match = match_counts["avg_match"]
    matched_skills = match_counts["matched"]
    missing_skills_count = match_counts["missing"]

    # ---------------- METRICS ----------------
    c1, c2, c3 = st.columns(3)
    c1.metric("Overall Match", f"{overall_match}%")
    c2.metric("Matched Skills", str(matched_skills))
    c3.metric("Missing Skills", str(missing_skills_count))

    # ---------------- BAR CHART ----------------
    bar_fig = go.Figure()
    bar_fig.add_bar(x=skills, y=resume_scores, name="Resume Skills")
    bar_fig.add_bar(x=skills, y=job_scores, name="Job Requirements")
    bar_fig.update_layout(
        title="Skill Match Overview",
        barmode="group",
        height=400
    )

    # ---------------- RADAR CHART ----------------
    radar_skills = skills[:5] if len(skills) >= 5 else skills
    radar_resume = resume_scores[:5] if len(resume_scores) >= 5 else resume_scores
    radar_job = job_scores[:5] if len(job_scores) >= 5 else job_scores
    
    radar_fig = go.Figure()
    radar_fig.add_trace(go.Scatterpolar(
        r=radar_resume,
        theta=radar_skills,
        fill='toself',
        name='Current Profile'
    ))
    radar_fig.add_trace(go.Scatterpolar(
        r=radar_job,
        theta=radar_skills,
        fill='toself',
        name='Job Requirements'
    ))
    radar_fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
        title="Technical Skill Comparison",
        height=400
    )

    left, right = st.columns([2, 1])
    left.plotly_chart(bar_fig, use_container_width=True)
    right.plotly_chart(radar_fig, use_container_width=True)

    # ---------------- SKILL PROGRESS ----------------
    st.subheader("Skill Comparison")
    sample_skills = [("Python", 92), ("Machine Learning", 88), ("SQL", 65)]
    for skill, score in sample_skills:
        st.write(skill)
        st.progress(score / 100)

    # ---------------- UPSKILLING ----------------
    st.subheader("Upskilling Recommendations")
    st.info("☁️ AWS Cloud Services\nComplete AWS Certified Solutions Architect course")
    st.info("📊 Advanced Statistics\nEnroll in Advanced Statistics for Data Science")
    st.info("📈 Project Management\nConsider PMP certification")

    # ---------------- REPORT CONTENT ----------------
    report_text = f"""
SKILL GAP ANALYSIS REPORT

Overall Match: {overall_match}%
Matched Skills: {matched_skills}
Missing Skills: {missing_skills_count}

SKILL DETAILS
"""

    for i in range(len(skills)):
        report_text += f"""
{skills[i]}
Resume Score: {resume_scores[i]}%
Job Requirement: {job_scores[i]}%
"""

    st.subheader("Report Download")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        # ---------------- CSV ----------------
        df = pd.DataFrame({
            "Skill": skills,
            "Resume Score (%)": resume_scores,
            "Job Requirement (%)": job_scores
        })

        st.download_button(
            "⬇️ CSV Report",
            df.to_csv(index=False).encode("utf-8"),
            "skill_gap_report.csv",
            "text/csv"
        )

    with col2:
        # ---------------- TXT ----------------
        st.download_button(
            "⬇️ TXT Report",
            report_text.encode("utf-8"),
            "skill_gap_report.txt",
            "text/plain"
        )

    with col3:
        # ---------------- DOC ----------------
        def generate_doc(text):
            doc = Document()
            doc.add_heading("Skill Gap Analysis Report", level=1)
            for line in text.split("\n"):
                doc.add_paragraph(line)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        doc_file = generate_doc(report_text)

        st.download_button(
            "⬇️ DOCX Report",
            doc_file,
            "skill_gap_report.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    with col4:
        # ---------------- PDF ----------------
        def generate_pdf(text):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in text.split("\n"):
                pdf.multi_cell(0, 8, line)
            return pdf.output(dest="S").encode("latin-1")

        pdf_bytes = generate_pdf(report_text)

        st.download_button(
            "⬇️ PDF Report",
            pdf_bytes,
            "skill_gap_report.pdf",
            "application/pdf"
        )

    st.success("Dashboard loaded successfully ✔️")