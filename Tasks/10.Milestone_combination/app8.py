import streamlit as st
import pdfplumber
import docx
import re
import spacy
from spacy.matcher import PhraseMatcher
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
import io
import requests

# Milestone - 1

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="Skill Gap Analysis Dashboard",
    layout="wide"
)

# ---------------- HEADER ----------------
st.markdown(
    """
    <div style="background-color:#4a5bdc;padding:20px;border-radius:10px">
        <h2 style="color:white;">
            Data Ingestion, Parsing and Cleaning Module 
        </h2>
        <p style="color:white;">
            This module allows uploading resumes and job descriptions,
            extracting text from multiple document formats, and displaying
            clean, normalized content for further processing.
        </p>
    </div>
    """,
    unsafe_allow_html=True
)

st.write("")

# ---------------- INITIALIZE ----------------
resume_text = None
jd_text = None
cleaned_resume = ""
cleaned_jd = ""
resume_experience = None
jd_experience = None

# ---------------- FUNCTIONS ----------------
def extract_text(file):
    """Extract text from PDF, DOCX, or TXT files"""
    try:
        if file.type == "application/pdf":
            text = ""
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    if page.extract_text():
                        text += page.extract_text() + "\n"
            return text

        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            file.seek(0)
            document = docx.Document(file)
            return "\n".join([para.text for para in document.paragraphs])

        elif file.type == "text/plain":
            file.seek(0)
            return file.read().decode("utf-8")
        
        return ""
    except Exception as e:
        st.error(f"Error extracting text from {file.name}: {str(e)}")
        st.info("💡 Try re-uploading the file or use a different format (PDF/TXT)")
        return ""

def clean(text):
    """Clean and normalize text for NLP processing"""
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r'(\d+)\s*(?:-|–|to)\s*(\d+)', r'\1_\2', text)
    text = re.sub(r'(\d+)\s*\+', r'\1_plus', text)
    text = re.sub(r'\n|\t', ' ', text)
    text = re.sub(r'[^a-z0-9_\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def extract_experience(text):
    """Extract years of experience from text"""
    if not text:
        return None
    text = text.lower()
    range_match = re.search(
        r'(\d+)\s*(?:-|–|to)\s*(\d+)\s*(?:years?|yrs?|year experience|years experience)',
        text
    )
    if range_match:
        return {"min_exp": int(range_match.group(1)), "max_exp": int(range_match.group(2))}
    plus_match = re.search(
        r'(\d+)\s*\+\s*(?:years?|yrs?|year experience|years experience)',
        text
    )
    if plus_match:
        return {"min_exp": int(plus_match.group(1)), "max_exp": None}
    return None

# ---------------- UPLOAD ----------------
st.subheader("Upload Documents")

left_col, right_col = st.columns(2)

with left_col:
    resume_file = st.file_uploader(
        "Upload Resume",
        type=["pdf", "docx", "txt"],
        help="Upload your resume in PDF, DOCX, or TXT format"
    )

with right_col:
    jd_file = st.file_uploader(
        "Upload Job Description",
        type=["pdf", "docx", "txt"],
        help="Upload job description in PDF, DOCX, or TXT format"
    )

if jd_file and resume_file:

    # ---------------- PREVIEW ----------------
    st.subheader("Parsed Document Preview")

    col1, col2 = st.columns(2)

    with col1:
        if resume_file:
            resume_text = extract_text(resume_file)
            st.text_area("Resume Preview", resume_text, height=300, key="resume_preview")

    with col2:
        if jd_file:
            jd_text = extract_text(jd_file)
            st.text_area("Job Description Preview", jd_text, height=300, key="jd_preview")

    # ---------------- PROCESS ----------------
    if resume_text:
        cleaned_resume = clean(resume_text)
        resume_experience = extract_experience(resume_text)

    if jd_text:
        cleaned_jd = clean(jd_text)
        jd_experience = extract_experience(jd_text)

    # ---------------- CLEANED OUTPUT ----------------
    st.subheader("Cleaned Files")

    col1, col2 = st.columns(2)

    with col1:
        if cleaned_resume:
            st.text_area("Cleaned Resume", cleaned_resume, height=200, key="cleaned_resume")
    with col2:
        if cleaned_jd:
            st.text_area("Cleaned Job Description", cleaned_jd, height=200, key="cleaned_jd")

    st.success("✅Completed: Documents uploaded, parsed, cleaned and previewed successfully.")

    # -------------------------------------------------------------------------
    # Milestone - 2

    st.markdown(
        """
        <div style="background-color:#4a5bdc;padding:20px;border-radius:10px;margin-top:30px;">
            <h2 style="color:white;">
                Skill Extraction using NLP Module
            </h2>
            <p style="color:white;">
                Module: Skill Extraction using NLP <br>
                • spaCy and BERT-based pipelines  <br>
                • Technical and soft skills identification <br>
                • Structured skill display
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.write("")

    st.subheader("Experience Analysis")

    col1, col2 = st.columns(2)
    with col1:
        if resume_experience:
            st.info(f"Resume Experience: {resume_experience['min_exp']}" + 
                   (f"-{resume_experience['max_exp']}" if resume_experience['max_exp'] else "+") + " years")
        else:
            st.warning("No experience information found in resume")
    with col2:
        if jd_experience:
            st.info(f"JD Required Experience: {jd_experience['min_exp']}" + 
                   (f"-{jd_experience['max_exp']}" if jd_experience['max_exp'] else "+") + " years")
        else:
            st.warning("No experience requirement found in JD")

    @st.cache_resource
    def load_spacy_model():
        try:
            return spacy.load("en_core_web_trf")
        except OSError:
            st.error("spaCy model 'en_core_web_trf' not found. Please install it using: python -m spacy download en_core_web_trf")
            return None
    
    nlp = load_spacy_model()
    
    if nlp is None:
        st.stop()

    master_technical_skills = {
        "python", "java", "c", "c++", "c#", "go", "rust", "kotlin", "swift",
        "javascript", "typescript", "php", "ruby", "r", "matlab", "scala",
        "perl", "bash", "shell scripting", "sql",
        "html", "css", "sass", "bootstrap", "tailwind css",
        "react", "angular", "vue.js", "next.js", "nuxt.js",
        "node.js", "express.js", "django", "flask", "fastapi",
        "spring", "spring boot", "laravel", "asp.net",
        "mysql", "postgresql", "sqlite", "oracle", "sql server",
        "mongodb", "cassandra", "dynamodb", "redis", "firebase",
        "elasticsearch", "neo4j",
        "numpy", "pandas", "scipy", "matplotlib", "seaborn",
        "plotly", "power bi", "tableau", "excel", "statistics",
        "data analysis", "data visualization",
        "machine learning", "deep learning", "artificial intelligence",
        "tensorflow", "keras", "pytorch", "scikit-learn", "xgboost",
        "opencv", "nlp", "computer vision", "speech recognition",
        "transformers", "huggingface", "langchain", "llm",
        "hadoop", "spark", "pyspark", "kafka", "hive", "pig",
        "hbase", "flink", "airflow", "databricks",
        "aws", "azure", "google cloud", "gcp",
        "docker", "kubernetes", "terraform", "ansible",
        "jenkins", "github actions", "ci/cd",
        "linux", "unix",
        "rest api", "graphql", "grpc", "soap",
        "microservices", "jwt", "oauth",
        "unit testing", "integration testing", "system testing",
        "pytest", "unittest", "junit", "selenium", "cypress",
        "postman",
        "android", "ios", "flutter", "react native",
        "swiftui", "kotlin multiplatform",
        "network security", "application security", "penetration testing",
        "ethical hacking", "cryptography", "owasp",
        "siem", "firewall",
        "windows", "linux", "macos",
        "git", "github", "gitlab", "bitbucket",
        "jira", "confluence", "trello",
        "data structures", "algorithms", "object oriented programming",
        "design patterns", "system design",
        "software development lifecycle", "agile", "scrum"
    }

    master_soft_skills = {
        "communication", "verbal communication", "written communication",
        "public speaking", "presentation skills", "active listening",
        "business communication", "storytelling",
        "teamwork", "collaboration", "interpersonal skills",
        "relationship building", "empathy", "emotional intelligence",
        "conflict resolution", "negotiation",
        "leadership", "people management", "team leadership",
        "decision making", "delegation", "mentoring", "coaching",
        "influencing", "strategic thinking",
        "problem solving", "critical thinking", "analytical thinking",
        "logical reasoning", "creative thinking", "innovation",
        "root cause analysis", "troubleshooting",
        "time management", "prioritization", "multitasking",
        "work ethic", "self discipline", "accountability",
        "goal setting", "organizational skills",
        "adaptability", "flexibility", "resilience",
        "learning agility", "continuous learning",
        "open mindedness", "growth mindset",
        "professionalism", "integrity", "ethical behavior",
        "reliability", "punctuality", "confidentiality",
        "creativity", "idea generation", "design thinking",
        "innovation mindset", "curiosity",
        "stress management", "self awareness", "self motivation",
        "confidence", "positive attitude", "emotional control",
        "customer focus", "customer service",
        "client management", "stakeholder management",
        "user empathy", "service mindset",
        "cross functional collaboration", "cultural awareness",
        "diversity and inclusion", "remote collaboration",
        "team alignment",
        "conflict management", "crisis management",
        "handling pressure", "decision making under pressure",
        "independent work", "collaborative work",
        "attention to detail", "quality focus",
        "result oriented", "ownership",
        "honesty", "trustworthiness", "respect",
        "fairness", "social responsibility"
    }

    def extract_soft_skills(text, skill_set):
        """Extract soft skills using simple text matching"""
        if not text:
            return []
        text = text.lower()
        found_skills = set()
        for skill in skill_set:
            if skill in text:
                found_skills.add(skill)
        return sorted(found_skills)

    def extract_technical_skills(text, skill_set):
        """Extract technical skills using spaCy PhraseMatcher"""
        if not text:
            return []
        matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
        patterns = [nlp.make_doc(skill) for skill in skill_set]
        matcher.add("TECH_SKILLS", patterns)
        doc = nlp(text)
        matches = matcher(doc)
        skills_found = set()
        for _, start, end in matches:
            skills_found.add(doc[start:end].text.lower())
        return sorted(skills_found)

    def extract_github_username(url):
        """Extract username from GitHub URL"""
        match = re.search(r'github\.com/([^/]+)/?', url)
        return match.group(1) if match else None

    def fetch_github_repos(username):
        """Fetch public repositories of a GitHub user"""
        api_url = f"https://api.github.com/users/{username}/repos"
        try:
            response = requests.get(api_url, timeout=10)
            if response.status_code == 200:
                return response.json()
        except:
            pass
        return []

    def extract_github_skills(repos, skill_set):
        """Extract skills from repo names, descriptions, and languages"""
        found_skills = set()
        for repo in repos:
            text = f"{repo.get('name','')} {repo.get('description','')}".lower()
            for skill in skill_set:
                if skill in text:
                    found_skills.add(skill)
            if repo.get("language"):
                lang = repo["language"].lower()
                if lang in skill_set:
                    found_skills.add(lang)
        return sorted(found_skills)

    with st.spinner("Extracting skills using NLP..."):
        resume_technical = extract_technical_skills(cleaned_resume, master_technical_skills)
        resume_soft = extract_soft_skills(cleaned_resume, master_soft_skills)
        
        # Initialize GitHub skills
        github_skills = []
        
        # GitHub Profile Analyzer Section
        st.subheader("🐙 GitHub Profile Analyzer (Optional)")
        st.caption("Enhance your skill profile by analyzing your GitHub repositories")
        
        github_url = st.text_input(
            "Enter GitHub Profile URL",
            placeholder="https://github.com/username",
            help="We'll extract technical skills from your public repositories"
        )
        
        if github_url:
            username = extract_github_username(github_url)
            if username:
                with st.spinner(f"Analyzing GitHub profile for @{username}..."):
                    repos = fetch_github_repos(username)
                    if repos:
                        github_skills = extract_github_skills(repos, master_technical_skills)
                        if github_skills:
                            st.success(f"✅ Found {len(github_skills)} technical skills from {len(repos)} repositories!")
                            with st.expander("View GitHub Skills", expanded=True):
                                st.write(", ".join(github_skills))
                        else:
                            st.warning("⚠️ No recognizable technical skills found in your repositories")
                    else:
                        st.error("❌ Unable to fetch repositories. Please check the username.")
            else:
                st.error("❌ Invalid GitHub URL format")
        
        # Merge GitHub skills with resume skills
        combined_technical_skills = sorted(set(resume_technical) | set(github_skills))
        
        resume_skills = {
            "technical": combined_technical_skills,
            "soft": resume_soft
        }
        
        # Show combined stats if GitHub was used
        if github_skills:
            col_a, col_b, col_c = st.columns(3)
            col_a.metric("📄 Resume Skills", len(resume_technical))
            col_b.metric("🐙 GitHub Skills", len(github_skills))
            col_c.metric("🎯 Total Unique", len(combined_technical_skills))

        jd_skills = {
            "technical": extract_technical_skills(cleaned_jd, master_technical_skills),
            "soft": extract_soft_skills(cleaned_jd, master_soft_skills)
        }

    st.subheader("Extracted Skills")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### Resume Skills")
        with st.expander("Technical Skills", expanded=True):
            if resume_skills["technical"]:
                st.write(", ".join(resume_skills["technical"]))
            else:
                st.info("No technical skills found")
        with st.expander("Soft Skills"):
            if resume_skills["soft"]:
                st.write(", ".join(resume_skills["soft"]))
            else:
                st.info("No soft skills found")

    with col2:
        st.markdown("### JD Skills")
        with st.expander("Technical Skills", expanded=True):
            if jd_skills["technical"]:
                st.write(", ".join(jd_skills["technical"]))
            else:
                st.info("No technical skills found")
        with st.expander("Soft Skills"):
            if jd_skills["soft"]:
                st.write(", ".join(jd_skills["soft"]))
            else:
                st.info("No soft skills found")

    def donut_chart(tech_count, soft_count, title):
        """Create donut chart for skill distribution"""
        labels = ["Technical Skills", "Soft Skills"]
        sizes = [tech_count, soft_count]
        if sum(sizes) == 0:
            sizes = [1, 1]
        fig = plt.figure(figsize=(4, 4))
        ax = fig.add_axes([0.15, 0.15, 0.7, 0.7]) 
        wedges, _, _ = ax.pie(
            sizes,
            startangle=90,
            autopct="%1.0f%%",
            radius=1,
            wedgeprops=dict(width=0.4, edgecolor="white"),
            colors=['#4a5bdc', '#fd7e14']
        )
        ax.set(aspect="equal")
        ax.set_title(title, pad=10, fontsize=12, fontweight='bold')
        ax.legend(
            wedges,
            labels,
            loc="lower center",
            bbox_to_anchor=(0.5, -0.25),
            ncol=2,
            frameon=False
        )
        return fig

    st.subheader("Skill Distribution Analysis")

    col1, col2 = st.columns(2)

    with col1:
        fig1 = donut_chart(
            len(resume_skills["technical"]),
            len(resume_skills["soft"]),
            "Resume Skill Distribution"
        )
        st.pyplot(fig1, use_container_width=False)
        plt.close(fig1)

    with col2:
        fig2 = donut_chart(
            len(jd_skills["technical"]),
            len(jd_skills["soft"]),
            "JD Skill Distribution"
        )
        st.pyplot(fig2, use_container_width=False)
        plt.close(fig2)

    def compute_metrics(resume_skills, jd_skills):
        """Compute skill matching metrics"""
        resume_tech = set(resume_skills["technical"])
        resume_soft = set(resume_skills["soft"])
        jd_tech = set(jd_skills["technical"])
        jd_soft = set(jd_skills["soft"])
        matched_skills = (resume_tech & jd_tech) | (resume_soft & jd_soft)
        total_jd_skills = len(jd_tech) + len(jd_soft)
        match_percentage = 0
        if total_jd_skills > 0:
            match_percentage = round((len(matched_skills) / total_jd_skills) * 100, 1)
        return {
            "resume_tech": len(resume_tech),
            "resume_soft": len(resume_soft),
            "jd_tech": len(jd_tech),
            "jd_soft": len(jd_soft),
            "total_resume": len(resume_tech) + len(resume_soft),
            "total_jd": total_jd_skills,
            "match_percent": match_percentage
        }
    
    metrics = compute_metrics(resume_skills, jd_skills)

    st.subheader("Skill Extraction Metrics")

    col1, col2, col3, col4 = st.columns(4)
    col1.metric(label="JD Technical Skills", value=metrics["jd_tech"])
    col2.metric(label="JD Soft Skills", value=metrics["jd_soft"])
    col3.metric(label="Total JD Skills", value=metrics["total_jd"])
    col4.metric(label="Basic Match %", value=f'{metrics["match_percent"]}%')

    st.success("✅Completed: Skills extracted successfully using NLP.")

    # --------------------------------------------------------------------------
    # Milestone - 3

    def categorize_skills(skills):
        """Categorize skills into meaningful groups"""
        categories = {
            "Programming Languages": ["python", "java", "c", "c++", "c#", "go", "rust", "kotlin", "swift",
                                     "javascript", "typescript", "php", "ruby", "r", "matlab", "scala", "perl"],
            "Web Technologies": ["html", "css", "sass", "bootstrap", "tailwind css", "react", "angular", 
                                "vue.js", "next.js", "nuxt.js", "node.js", "express.js", "django", "flask", 
                                "fastapi", "spring", "spring boot", "laravel", "asp.net"],
            "Databases": ["mysql", "postgresql", "sqlite", "oracle", "sql server", "mongodb", "cassandra", 
                         "dynamodb", "redis", "firebase", "elasticsearch", "neo4j", "sql"],
            "Data Science & ML": ["numpy", "pandas", "scipy", "matplotlib", "seaborn", "plotly", "power bi", 
                                 "tableau", "excel", "machine learning", "deep learning", "tensorflow", "keras", 
                                 "pytorch", "scikit-learn", "xgboost", "opencv", "nlp", "computer vision", 
                                 "transformers", "huggingface", "langchain", "llm", "artificial intelligence",
                                 "statistics", "data analysis", "data visualization"],
            "Cloud & DevOps": ["aws", "azure", "google cloud", "gcp", "docker", "kubernetes", "terraform", 
                              "ansible", "jenkins", "github actions", "ci/cd", "linux", "unix"],
            "Big Data": ["hadoop", "spark", "pyspark", "kafka", "hive", "pig", "hbase", "flink", 
                        "airflow", "databricks"],
            "Mobile Dev": ["android", "ios", "flutter", "react native", "swiftui", "kotlin multiplatform"],
            "Testing & QA": ["unit testing", "integration testing", "system testing", "pytest", "unittest", 
                            "junit", "selenium", "cypress", "postman"],
            "APIs & Backend": ["rest api", "graphql", "grpc", "soap", "microservices", "jwt", "oauth"],
            "Security": ["network security", "application security", "penetration testing",
                        "ethical hacking", "cryptography", "owasp", "siem", "firewall"],
            "Leadership & Management": ["leadership", "people management", "team leadership", "decision making", 
                                       "delegation", "mentoring", "coaching", "strategic thinking"],
            "Communication": ["communication", "verbal communication", "written communication", "public speaking", 
                             "presentation skills", "active listening", "storytelling", "business communication"],
            "Problem Solving": ["problem solving", "critical thinking", "analytical thinking", "logical reasoning", 
                               "creative thinking", "innovation", "troubleshooting", "root cause analysis"],
            "Work Management": ["time management", "prioritization", "multitasking", "work ethic", 
                               "self discipline", "accountability", "goal setting", "organizational skills"],
            "Adaptability": ["adaptability", "flexibility", "resilience", "learning agility", 
                            "continuous learning", "open mindedness", "growth mindset"],
            "Collaboration": ["teamwork", "collaboration", "interpersonal skills", "relationship building",
                             "cross functional collaboration", "cultural awareness", "diversity and inclusion",
                             "remote collaboration", "team alignment"],
            "Soft Skills": ["empathy", "emotional intelligence", "conflict resolution", "negotiation",
                           "professionalism", "integrity", "customer focus", "customer service"]
        }
        
        categorized = {}
        uncategorized = []
        
        for skill in skills:
            skill_lower = skill.lower()
            found = False
            for category, category_skills in categories.items():
                if skill_lower in category_skills:
                    if category not in categorized:
                        categorized[category] = []
                    categorized[category].append(skill)
                    found = True
                    break
            if not found:
                uncategorized.append(skill)
        
        if uncategorized:
            categorized["Other Skills"] = uncategorized
        
        return categorized

    def build_similarity_matrix(resume_skills, jd_skills):
        """Build similarity matrix between resume and JD skills"""
        if not resume_skills or not jd_skills:
            return np.array([])
        matrix = np.zeros((len(resume_skills), len(jd_skills)))
        for i, r_skill in enumerate(resume_skills):
            for j, j_skill in enumerate(jd_skills):
                if r_skill == j_skill:
                    matrix[i][j] = 1.0
                elif r_skill in j_skill or j_skill in r_skill:
                    matrix[i][j] = 0.5
                else:
                    matrix[i][j] = 0.0
        return matrix

    def plot_category_match_heatmap(resume_skills, jd_skills):
        """Plot single consolidated heatmap showing category-wise skill match"""
        jd_categorized = categorize_skills(jd_skills)
        resume_categorized = categorize_skills(resume_skills)
        
        common_categories = set(jd_categorized.keys()) & set(resume_categorized.keys())
        
        if not common_categories:
            st.warning("No common skill categories found between resume and JD")
            return None
        
        # Build category-wise match matrix
        categories = sorted(common_categories)
        match_data = []
        
        for category in categories:
            jd_cat_skills = jd_categorized[category]
            resume_cat_skills = set(resume_categorized[category])
            
            matched = 0
            partial = 0
            missing = 0
            
            for jd_skill in jd_cat_skills:
                if jd_skill in resume_cat_skills:
                    matched += 1
                else:
                    # Check for partial match
                    found_partial = False
                    for res_skill in resume_cat_skills:
                        if jd_skill in res_skill or res_skill in jd_skill:
                            partial += 1
                            found_partial = True
                            break
                    if not found_partial:
                        missing += 1
            
            total = len(jd_cat_skills)
            
            match_data.append({
                'Category': category,
                'Matched': matched,
                'Partial': partial,
                'Missing': missing,
                'Total': total
            })
        
        # Create stacked bar chart
        fig = go.Figure()
        
        categories_list = [d['Category'] for d in match_data]
        
        fig.add_trace(go.Bar(
            name='Matched',
            x=categories_list,
            y=[d['Matched'] for d in match_data],
            marker_color='#28a745',
            text=[d['Matched'] for d in match_data],
            textposition='inside',
            hovertemplate='<b>%{x}</b><br>Matched: %{y}<extra></extra>'
        ))
        
        fig.add_trace(go.Bar(
            name='Partial Match',
            x=categories_list,
            y=[d['Partial'] for d in match_data],
            marker_color='#ffc107',
            text=[d['Partial'] for d in match_data],
            textposition='inside',
            hovertemplate='<b>%{x}</b><br>Partial: %{y}<extra></extra>'
        ))
        
        fig.add_trace(go.Bar(
            name='Missing',
            x=categories_list,
            y=[d['Missing'] for d in match_data],
            marker_color='#dc3545',
            text=[d['Missing'] for d in match_data],
            textposition='inside',
            hovertemplate='<b>%{x}</b><br>Missing: %{y}<extra></extra>'
        ))
        
        fig.update_layout(
            title="Category-wise Skill Match Analysis (Stacked)",
            barmode='stack',
            height=500,
            xaxis_title="Skill Categories",
            yaxis_title="Number of Skills",
            hovermode='x unified',
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1
            ),
            showlegend=True
        )
        
        return fig

    resume_all_skills = resume_skills["technical"] + resume_skills["soft"]
    jd_all_skills = jd_skills["technical"] + jd_skills["soft"]

    st.markdown("""
        <div style="background-color:#4a5bdc;padding:20px;border-radius:10px;margin-top:30px;">
            <h2 style="color:white;">Skill Gap Analysis and Similarity Matching Module</h2>
            <p style="color:white;">• Skill similarity matrix visualization <br>• Resume vs JD skill comparison <br>• Missing skill identification</p>
        </div><br>""", unsafe_allow_html=True)

    if resume_all_skills and jd_all_skills:
        similarity_matrix = build_similarity_matrix(resume_all_skills, jd_all_skills)
        st.subheader("📊 Category-wise Skill Match Heatmap")
        fig_heatmap = plot_category_match_heatmap(resume_all_skills, jd_all_skills)
        if fig_heatmap:
            st.plotly_chart(fig_heatmap, use_container_width=True)
    else:
        st.warning("⚠️ Insufficient skills to build skill gap analysis")

    def classify_skill_matches(similarity_matrix, resume_skills, jd_skills):
        """Classify skills as matched, partial, or missing"""
        matched = set()
        partial = set()
        for j, jd_skill in enumerate(jd_skills):
            column = similarity_matrix[:, j]
            if 1.0 in column:
                matched.add(jd_skill)
            elif 0.5 in column:
                partial.add(jd_skill)
        missing = set(jd_skills) - matched - partial
        return {
            "matched": sorted(list(matched)),
            "partial": sorted(list(partial)),
            "missing": sorted(list(missing))
        }

    skill_match_result = classify_skill_matches(similarity_matrix, resume_all_skills, jd_all_skills)

    def calculate_skill_match(resume_skills, jd_skills):
        """Calculate detailed skill match statistics"""
        resume_set = set(resume_skills)
        jd_set = set(jd_skills)
        matched = resume_set & jd_set
        missing = jd_set - resume_set
        partial = set()
        for jd_skill in missing.copy():
            for res_skill in resume_set:
                if jd_skill in res_skill or res_skill in jd_skill:
                    partial.add(jd_skill)
        missing = missing - partial
        avg_match = 0
        if len(jd_set) > 0:
            avg_match = round(((len(matched) + 0.5 * len(partial)) / len(jd_set)) * 100, 1)
        return {
            "matched": len(matched),
            "partial": len(partial),
            "missing": len(missing),
            "avg_match": avg_match
        }

    all_resume_skills = set(resume_skills["technical"]) | set(resume_skills["soft"])
    all_jd_skills = set(jd_skills["technical"]) | set(jd_skills["soft"])
    match_counts = calculate_skill_match(all_resume_skills, all_jd_skills)

    st.subheader("Skill Match Overview")

    left, right = st.columns([1.3, 1])

    with left:
        labels = ["Matched", "Partially Matched", "Missing"]
        sizes = [match_counts["matched"], match_counts["partial"], match_counts["missing"]]
        colors = ["#28a745", "#fd7e14", "#dc3545"]
        if sum(sizes) == 0:
            sizes = [1, 1, 1]
        fig = plt.figure(figsize=(3.6, 3.6))
        ax = fig.add_axes([0.1, 0.15, 0.8, 0.75])
        wedges, _, autotexts = ax.pie(sizes, startangle=90, autopct="%1.0f%%", radius=1, colors=colors, wedgeprops=dict(width=0.4, edgecolor="white"))
        for autotext in autotexts:
            autotext.set_color('black')
            autotext.set_fontweight('bold')
            autotext.set_fontsize(11)
        ax.set(aspect="equal")
        ax.set_title("Skill Match Distribution", pad=8, fontsize=12, fontweight='bold')
        ax.legend(wedges, labels, loc="lower center", bbox_to_anchor=(0.5, -0.25), ncol=3, frameon=False)
        st.pyplot(fig, use_container_width=False)
        plt.close(fig)

    with right:
        with st.container():
            r1, r2 = st.columns(2)
            r1.metric("✅ Matched Skills", match_counts["matched"])
            r2.metric("⚠️ Partially Matched", match_counts["partial"])
            r3, r4 = st.columns(2)
            r3.metric("❌ Missing Skills", match_counts["missing"])
            r4.metric("📊 Avg Match %", f'{match_counts["avg_match"]}%')

    st.subheader("Skill Gap Details")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("### ✅ Matched Skills")
        if skill_match_result["matched"]:
            for skill in skill_match_result["matched"]:
                st.success(f"✓ {skill}")
        else:
            st.info("No perfectly matched skills")

    with col2:
        st.markdown("### ⚠️ Partial Matches")
        if skill_match_result["partial"]:
            for skill in skill_match_result["partial"]:
                st.warning(f"≈ {skill}")
        else:
            st.info("No partially matched skills")

    with col3:
        st.markdown("### ❌ Missing Skills")
        if skill_match_result["missing"]:
            for skill in skill_match_result["missing"]:
                st.error(f"✗ {skill}")
        else:
            st.success("No missing skills!")

    st.success("✅Completed: Skill gap analysis completed successfully.")

    # -----------------------------------------------------------------------------
    # Milestone - 4

    st.markdown("""<div style="background-color:#4a90e2;padding:15px;border-radius:8px;margin-top:30px;">
    <h2 style="color:white;text-align:center;">Dashboard and Report Export Module</h2>
    <p style="color:white;text-align:center;">Interactive dashboard • Graphs • Multi-format report export</p>
    </div>""", unsafe_allow_html=True)

    st.write("")

    skills = list(all_jd_skills)
    resume_scores = [100 if skill in all_resume_skills else 0 for skill in skills]
    job_scores = [100 for _ in skills]
    overall_match = match_counts["avg_match"]
    matched_skills_count = match_counts["matched"]
    missing_skills_count = match_counts["missing"]

    c1, c2, c3 = st.columns(3)
    c1.metric("Overall Match", f"{overall_match}%")
    c2.metric("Matched Skills", str(matched_skills_count))
    c3.metric("Missing Skills", str(missing_skills_count))

    st.subheader("Categorized Skill Match Overview")
    
    jd_categorized = categorize_skills(list(all_jd_skills))
    categories_to_plot = []
    resume_category_scores = []
    jd_category_scores = []
    
    for category, cat_skills in jd_categorized.items():
        matched_in_category = sum(1 for skill in cat_skills if skill in all_resume_skills)
        total_in_category = len(cat_skills)
        if total_in_category > 0:
            categories_to_plot.append(category)
            resume_category_scores.append((matched_in_category / total_in_category) * 100)
            jd_category_scores.append(100)
    
    bar_fig = go.Figure()
    bar_fig.add_trace(go.Bar(name='Your Skills', x=categories_to_plot, y=resume_category_scores, marker_color='#4a5bdc', text=[f'{score:.0f}%' for score in resume_category_scores], textposition='outside'))
    bar_fig.add_trace(go.Bar(name='Job Requirements', x=categories_to_plot, y=jd_category_scores, marker_color='#28a745', text=['100%'] * len(categories_to_plot), textposition='outside'))
    bar_fig.update_layout(title="Skill Match by Category", barmode="group", height=500, xaxis_title="Skill Categories", yaxis_title="Coverage (%)", hovermode='x unified', yaxis=dict(range=[0, 120]), legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1))
    st.plotly_chart(bar_fig, use_container_width=True)

    radar_categories = categories_to_plot[:6] if len(categories_to_plot) > 6 else categories_to_plot
    radar_resume_scores = resume_category_scores[:6] if len(resume_category_scores) > 6 else resume_category_scores
    radar_jd_scores = [100] * len(radar_categories)
    
    radar_fig = go.Figure()
    radar_fig.add_trace(go.Scatterpolar(r=radar_resume_scores, theta=radar_categories, fill='toself', name='Current Profile', line_color='#4a5bdc'))
    radar_fig.add_trace(go.Scatterpolar(r=radar_jd_scores, theta=radar_categories, fill='toself', name='Job Requirements', line_color='#28a745'))
    radar_fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), title="Top Skill Categories Comparison", height=500, showlegend=True)
    st.plotly_chart(radar_fig, use_container_width=True)

    st.subheader("📈 Sample Skill Proficiency")
    st.caption("Note: These are sample scores for demonstration. For actual proficiency, use skill assessment tools.")
    
    sample_skills = [("Python", 92), ("Machine Learning", 88), ("SQL", 65)]
    for skill, score in sample_skills:
        col_skill, col_progress = st.columns([1, 4])
        with col_skill:
            st.write(f"**{skill}**")
        with col_progress:
            st.progress(score / 100)
            st.caption(f"{score}%")

    st.subheader("📊 Category-wise Similarity Score Distribution")
    
    if resume_all_skills and jd_all_skills and len(similarity_matrix) > 0:
        # Categorize skills and calculate average similarity per category
        jd_categorized = categorize_skills(jd_all_skills)
        
        category_similarities = {}
        
        for category, cat_skills in jd_categorized.items():
            cat_scores = []
            for skill in cat_skills:
                if skill in jd_all_skills:
                    j = jd_all_skills.index(skill)
                    column_scores = similarity_matrix[:, j]
                    avg_score = np.mean(column_scores) * 100
                    cat_scores.append(avg_score)
            
            if cat_scores:
                category_similarities[category] = np.mean(cat_scores)
        
        if category_similarities:
            categories = list(category_similarities.keys())
            scores = list(category_similarities.values())
            
            # Create area chart using plotly
            area_fig = go.Figure()
            area_fig.add_trace(go.Scatter(
                x=categories,
                y=scores,
                fill='tozeroy',
                name='Avg Similarity Score',
                line_color='#7c3aed',
                fillcolor='rgba(124, 58, 237, 0.3)',
                mode='lines+markers'
            ))
            
            area_fig.update_layout(
                title="Average Similarity Score by Category",
                xaxis_title="Skill Categories",
                yaxis_title="Similarity Score (%)",
                height=400,
                hovermode='x unified',
                yaxis=dict(range=[0, 100])
            )
            
            st.plotly_chart(area_fig, use_container_width=True)
        else:
            st.info("No category data available for similarity analysis")
    else:
        st.info("Upload documents to see similarity score distribution")

    st.subheader("💡 Upskilling Recommendations")
    st.caption("Based on missing and partially matched skills from job description")
    
    if skill_match_result["missing"]:
        st.markdown("**🔴 Priority Skills to Learn (Missing):**")
        for i, skill in enumerate(skill_match_result["missing"][:5], 1):
            st.error(f"{i}. **{skill.title()}** - Not found in your resume")
    
    if skill_match_result["partial"]:
        st.markdown("**🟡 Skills to Strengthen (Partial Match):**")
        for i, skill in enumerate(skill_match_result["partial"][:5], 1):
            st.warning(f"{i}. **{skill.title()}** - Improve proficiency in this area")
    
    if not skill_match_result["missing"] and not skill_match_result["partial"]:
        st.success("🎉 Excellent! You have all the required skills for this job!")

    st.subheader("📥 Report Download")

    col1, col2, col3 = st.columns(3)

    with col1:
        # Separate technical and soft skills
        tech_matched = [s for s in skill_match_result['matched'] if s in jd_skills["technical"]]
        tech_partial = [s for s in skill_match_result['partial'] if s in jd_skills["technical"]]
        tech_missing = [s for s in skill_match_result['missing'] if s in jd_skills["technical"]]
        
        soft_matched = [s for s in skill_match_result['matched'] if s in jd_skills["soft"]]
        soft_partial = [s for s in skill_match_result['partial'] if s in jd_skills["soft"]]
        soft_missing = [s for s in skill_match_result['missing'] if s in jd_skills["soft"]]
        
        # Pad lists to same length within each category
        max_tech_len = max(len(tech_matched), len(tech_partial), len(tech_missing), 1)
        tech_matched += [''] * (max_tech_len - len(tech_matched))
        tech_partial += [''] * (max_tech_len - len(tech_partial))
        tech_missing += [''] * (max_tech_len - len(tech_missing))
        
        max_soft_len = max(len(soft_matched), len(soft_partial), len(soft_missing), 1)
        soft_matched += [''] * (max_soft_len - len(soft_matched))
        soft_partial += [''] * (max_soft_len - len(soft_partial))
        soft_missing += [''] * (max_soft_len - len(soft_missing))
        
        # Create dataframe with technical skills first, then gap, then soft skills
        tech_df = pd.DataFrame({
            "Technical - Matched": tech_matched,
            "Technical - Partially Matched": tech_partial,
            "Technical - Missing": tech_missing
        })
        
        # Add empty row as separator
        separator_df = pd.DataFrame({
            "Technical - Matched": [''],
            "Technical - Partially Matched": [''],
            "Technical - Missing": ['']
        })
        
        soft_df = pd.DataFrame({
            "Soft Skills - Matched": soft_matched,
            "Soft Skills - Partially Matched": soft_partial,
            "Soft Skills - Missing": soft_missing
        })
        
        # Combine: technical, separator, soft
        csv_output = tech_df.to_csv(index=False)
        csv_output += "\n"  # Add blank line
        csv_output += soft_df.to_csv(index=False)
        
        st.download_button(
            "⬇️ CSV Report",
            csv_output.encode("utf-8"),
            "skill_gap_report.csv",
            "text/csv",
            help="Download skill comparison separated by Technical and Soft Skills"
        )

    with col2:
        def generate_word_report():
            doc = Document()
            title = doc.add_heading('Skill Gap Analysis Report', 0)
            title.alignment = 1
            doc.add_heading('Executive Summary', 1)
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Light Grid Accent 1'
            summary_data = [
                ('Overall Match Rate', f'{overall_match}%'),
                ('Matched Skills', str(matched_skills_count)),
                ('Partially Matched Skills', str(match_counts["partial"])),
                ('Missing Skills', str(missing_skills_count))
            ]
            for i, (label, value) in enumerate(summary_data):
                summary_table.rows[i].cells[0].text = label
                summary_table.rows[i].cells[1].text = value
            doc.add_paragraph()
            
            # Technical Skills Section
            doc.add_heading('Technical Skills Breakdown', 1)
            
            tech_matched = [s for s in skill_match_result['matched'] if s in jd_skills["technical"]]
            tech_partial = [s for s in skill_match_result['partial'] if s in jd_skills["technical"]]
            tech_missing = [s for s in skill_match_result['missing'] if s in jd_skills["technical"]]
            
            doc.add_heading('Matched Technical Skills', 2)
            if tech_matched:
                for skill in tech_matched:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No perfectly matched technical skills found.')
            
            doc.add_heading('Partially Matched Technical Skills', 2)
            if tech_partial:
                for skill in tech_partial:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No partially matched technical skills found.')
            
            doc.add_heading('Missing Technical Skills', 2)
            if tech_missing:
                for skill in tech_missing:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No missing technical skills!')
            
            # Soft Skills Section
            doc.add_heading('Soft Skills Breakdown', 1)
            
            soft_matched = [s for s in skill_match_result['matched'] if s in jd_skills["soft"]]
            soft_partial = [s for s in skill_match_result['partial'] if s in jd_skills["soft"]]
            soft_missing = [s for s in skill_match_result['missing'] if s in jd_skills["soft"]]
            
            doc.add_heading('Matched Soft Skills', 2)
            if soft_matched:
                for skill in soft_matched:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No perfectly matched soft skills found.')
            
            doc.add_heading('Partially Matched Soft Skills', 2)
            if soft_partial:
                for skill in soft_partial:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No partially matched soft skills found.')
            
            doc.add_heading('Missing Soft Skills', 2)
            if soft_missing:
                for skill in soft_missing:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No missing soft skills!')
            
            # Recommendations
            doc.add_heading('Recommendations', 1)
            doc.add_paragraph(
                f'Your current match rate is {overall_match}%. '
                f'To improve your candidacy, consider focusing on the missing skills listed above.'
            )
            
            if overall_match >= 70:
                doc.add_paragraph('You have a strong skill match for this position!', style='List Bullet')
            else:
                doc.add_paragraph('Focus on upskilling in the missing areas to increase your match rate.', style='List Bullet')
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph(f'Generated on: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
            footer.alignment = 1
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
        doc_file = generate_word_report()
        st.download_button("⬇️ DOCX Report", doc_file, "skill_gap_analysis_report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", help="Download formatted Word report with complete analysis")
    
    with col3:
        def generate_pdf_report():
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", 'B', 20)
            pdf.cell(0, 15, "Skill Gap Analysis Report", ln=True, align='C')
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Executive Summary", ln=True)
            pdf.set_font("Arial", '', 11)
            pdf.ln(2)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(90, 8, "Overall Match Rate:", 1, 0, fill=True)
            pdf.cell(90, 8, f"{overall_match}%", 1, 1)
            pdf.cell(90, 8, "Matched Skills:", 1, 0, fill=True)
            pdf.cell(90, 8, str(matched_skills_count), 1, 1)
            pdf.cell(90, 8, "Partially Matched Skills:", 1, 0, fill=True)
            pdf.cell(90, 8, str(match_counts['partial']), 1, 1)
            pdf.cell(90, 8, "Missing Skills:", 1, 0, fill=True)
            pdf.cell(90, 8, str(missing_skills_count), 1, 1)
            pdf.ln(10)
            
            # Match Rate Visualization
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "Match Rate Visualization", ln=True)
            pdf.set_font("Arial", '', 10)
            bar_width = 170
            bar_height = 15
            x_start = 20
            y_start = pdf.get_y()
            pdf.set_fill_color(220, 220, 220)
            pdf.rect(x_start, y_start, bar_width, bar_height, 'F')
            filled_width = (overall_match / 100) * bar_width
            if overall_match >= 70:
                pdf.set_fill_color(40, 167, 69)
            elif overall_match >= 50:
                pdf.set_fill_color(253, 126, 20)
            else:
                pdf.set_fill_color(220, 53, 69)
            pdf.rect(x_start, y_start, filled_width, bar_height, 'F')
            pdf.set_draw_color(100, 100, 100)
            pdf.rect(x_start, y_start, bar_width, bar_height, 'D')
            pdf.set_xy(x_start + bar_width/2 - 10, y_start + 3)
            pdf.set_font("Arial", 'B', 11)
            if overall_match > 30:
                pdf.set_text_color(255, 255, 255)
            else:
                pdf.set_text_color(0, 0, 0)
            pdf.cell(20, 8, f"{overall_match}%", align='C')
            pdf.set_text_color(0, 0, 0)
            pdf.ln(20)
            
            # Skills Distribution Chart
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "Skills Distribution", ln=True)
            pdf.ln(2)
            matched = match_counts['matched']
            partial = match_counts['partial']
            missing = match_counts['missing']
            total_skills = matched + partial + missing
            if total_skills > 0:
                chart_width = 170
                chart_height = 40
                x_chart = 20
                y_chart = pdf.get_y()
                matched_width = (matched / total_skills) * chart_width
                partial_width = (partial / total_skills) * chart_width
                missing_width = (missing / total_skills) * chart_width
                if matched > 0:
                    pdf.set_fill_color(40, 167, 69)
                    pdf.rect(x_chart, y_chart, matched_width, chart_height, 'F')
                if partial > 0:
                    pdf.set_fill_color(253, 126, 20)
                    pdf.rect(x_chart + matched_width, y_chart, partial_width, chart_height, 'F')
                if missing > 0:
                    pdf.set_fill_color(220, 53, 69)
                    pdf.rect(x_chart + matched_width + partial_width, y_chart, missing_width, chart_height, 'F')
                pdf.set_draw_color(100, 100, 100)
                pdf.rect(x_chart, y_chart, chart_width, chart_height, 'D')
                pdf.ln(chart_height + 5)
                pdf.set_font("Arial", '', 9)
                legend_x = 20
                legend_y = pdf.get_y()
                pdf.set_fill_color(40, 167, 69)
                pdf.rect(legend_x, legend_y, 5, 5, 'F')
                pdf.set_xy(legend_x + 7, legend_y - 1)
                pdf.cell(40, 6, f"Matched ({matched})")
                pdf.set_fill_color(253, 126, 20)
                pdf.rect(legend_x + 50, legend_y, 5, 5, 'F')
                pdf.set_xy(legend_x + 57, legend_y - 1)
                pdf.cell(40, 6, f"Partial ({partial})")
                pdf.set_fill_color(220, 53, 69)
                pdf.rect(legend_x + 100, legend_y, 5, 5, 'F')
                pdf.set_xy(legend_x + 107, legend_y - 1)
                pdf.cell(40, 6, f"Missing ({missing})")
                pdf.ln(12)
            
            # Technical Skills Section
            tech_matched = [s for s in skill_match_result['matched'] if s in jd_skills["technical"]]
            tech_partial = [s for s in skill_match_result['partial'] if s in jd_skills["technical"]]
            tech_missing = [s for s in skill_match_result['missing'] if s in jd_skills["technical"]]
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(40, 167, 69)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Technical Skills - Matched", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if tech_matched:
                for skill in tech_matched:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No technical skills matched.", 0, 1)
            pdf.ln(3)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(253, 126, 20)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Technical Skills - Partially Matched", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if tech_partial:
                for skill in tech_partial:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No technical skills partially matched.", 0, 1)
            pdf.ln(3)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(220, 53, 69)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Technical Skills - Missing (Priority)", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if tech_missing:
                for skill in tech_missing:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No missing technical skills!", 0, 1)
            pdf.ln(5)
            
            # Soft Skills Section
            soft_matched = [s for s in skill_match_result['matched'] if s in jd_skills["soft"]]
            soft_partial = [s for s in skill_match_result['partial'] if s in jd_skills["soft"]]
            soft_missing = [s for s in skill_match_result['missing'] if s in jd_skills["soft"]]
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(40, 167, 69)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Soft Skills - Matched", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if soft_matched:
                for skill in soft_matched:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No soft skills matched.", 0, 1)
            pdf.ln(3)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(253, 126, 20)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Soft Skills - Partially Matched", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if soft_partial:
                for skill in soft_partial:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No soft skills partially matched.", 0, 1)
            pdf.ln(3)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(220, 53, 69)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Soft Skills - Missing", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if soft_missing:
                for skill in soft_missing:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No missing soft skills!", 0, 1)
            pdf.ln(5)
            
            # Recommendations
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "Recommendations", 0, 1)
            pdf.set_font("Arial", '', 10)
            recommendation = f"Your current match rate is {overall_match}%. "
            if overall_match >= 70:
                recommendation += "You have a strong skill match for this position! Continue to maintain and update these skills."
            elif overall_match >= 50:
                recommendation += "You have a moderate skill match. Focus on developing the missing skills and strengthening partially matched skills to improve your candidacy."
            else:
                recommendation += "There is significant room for improvement. Prioritize learning the missing skills listed above to increase your match rate to 70% or above."
            pdf.multi_cell(0, 6, recommendation)
            pdf.ln(3)
            
            # Action Items
            pdf.set_font("Arial", 'B', 10)
            pdf.cell(0, 6, "Action Items:", ln=True)
            pdf.set_font("Arial", '', 10)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "1. Focus on acquiring the missing skills through courses and certifications", ln=True)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "2. Strengthen partially matched skills through practical projects", ln=True)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "3. Update your resume once you acquire new skills", ln=True)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "4. Target positions with 70%+ match rate for better success", ln=True)
            pdf.ln(5)
            
            # Footer
            pdf.set_font("Arial", 'I', 9)
            pdf.cell(0, 10, f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 0, 'C')
            
            return pdf.output(dest="S").encode("latin-1")

        try:
            pdf_bytes = generate_pdf_report()
            st.download_button(label="⬇️ PDF Report", data=pdf_bytes, file_name="skill_gap_analysis_report.pdf", mime="application/pdf", help="Download formatted PDF report with visualizations")
        except Exception as e:
            st.error(f"Error generating PDF: {e}")

    st.success("✅Completed: Dashboard loaded successfully with export options.")
    
    st.markdown("---")
    st.markdown("""
    ### 🎯 Analysis Complete!
    
    **Next Steps:**
    1. Review the missing skills identified above
    2. Consider upskilling in those areas
    3. Update your resume with newly acquired skills
    4. Download the report for your records
    
    **Pro Tip:** Aim for at least 70% match rate for better job prospects!
    """)

else:
    st.info("👆 Please upload both Resume and Job Description files to begin the analysis.")
    st.markdown("""
    ### 📋 Supported File Formats:
    - **PDF** (.pdf)
    - **Word Document** (.docx)
    - **Text File** (.txt)
    
    ### 💡 Tips:
    - Ensure your files are properly formatted
    - Include clear skill listings in both documents
    - Experience information should be mentioned as "X years" or "X-Y years"
    """)