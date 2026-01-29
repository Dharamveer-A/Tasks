import streamlit as st
import pdfplumber
import docx
import re
import spacy
from spacy.matcher import PhraseMatcher
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
import io
import requests

st.set_page_config(page_title="SkillGapAI", layout="wide")

st.markdown("""
<h1 style="text-align:center; color:#470047;">SkillGapAI</h1>
<p style="text-align:center; font-size:18px; color:#6B7280;">
AI-Powered Skill Gap Analysis & Career Insights
</p>
<hr>
""", unsafe_allow_html=True)

with st.expander("How to Use - Watch Tutorial", expanded=False):
    st.markdown("""
    ### Quick Tutorial
    Watch this short video to learn how to use the Skill Gap Analysis Dashboard:
    """)
    
    # Option 1: If you have a GIF file
    # Uncomment and replace with your GIF path
    # st.image("path/to/your/tutorial.gif", caption="Step-by-step tutorial", use_container_width=True)
    
    # Option 2: If you have a video file (mp4, webm, etc.)
    # Uncomment and replace with your video path
    d1, d2, d3 = st.columns(3)

    with d2:
        video_file = open('tutorial.mov', 'rb')
        video_bytes = video_file.read()
        st.video(video_bytes)
    
    

    # Option 3: If you have a YouTube video
    # Uncomment and replace with your YouTube video URL
    # st.video("https://www.youtube.com/watch?v=YOUR_VIDEO_ID")
    
    # Option 4: If you have a GIF URL
    # Uncomment and replace with your GIF URL
    # st.markdown("![Tutorial GIF](https://your-url.com/tutorial.gif)")
    
    # Placeholder instructions (remove when you add actual GIF/video)
    st.info("""
    **Step-by-Step Instructions:**
    
    1. **Upload Resume** - Click on the left uploader and select your resume (PDF, DOCX, or TXT)
    2. **Upload Job Description** - Click on the right uploader and select the job description
    3. **Click Analyze** - Press the "Analyze Documents" button
    4. **View Results** - Review skill matches, gaps, and recommendations
    5. **Download Reports** - Export your analysis in CSV, DOCX, or PDF format
    """)

st.markdown("---")

# Milestone - 1


# ---------------- HEADER ----------------
st.markdown(
    """
    <div style="background-color:#470047;padding:20px;border-radius:10px">
        <h2 style="color:white;">
            Data Ingestion, Parsing and Cleaning Module 
        </h2>
        <p style="color:white;">
            This module allows uploading resumes and job descriptions,
            extracting text from multiple document formats (including multi-page documents),
            and displaying clean, normalized content for further processing.
        </p>
    </div>
    """,
    unsafe_allow_html=True
)

st.write("")

# ---------------- TUTORIAL VIDEO SECTION ----------------
st.markdown("---")

# Create an expander for the tutorial




# ---------------- INITIALIZE ----------------
resume_text = None
jd_text = None
cleaned_resume = ""
cleaned_jd = ""
resume_experience = None
jd_experience = None

# ---------------- FUNCTIONS ----------------
def extract_text(file):
    """Extract text from PDF, DOCX, or TXT files - handles multi-page documents"""
    try:
        if file.type == "application/pdf":
            text = ""
            file.seek(0)  # Reset file pointer to beginning
            
            try:
                with pdfplumber.open(file) as pdf:
                    total_pages = len(pdf.pages)
                    st.info(f"Processing PDF: {total_pages} page(s) detected")
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            # Add page separator for better text flow
                            text += page_text + "\n\n"
                        else:
                            st.warning(f"Page {page_num} appears to be empty or contains only images")
                    
                    if not text.strip():
                        st.error("No text could be extracted from the PDF. The file might contain only images or be corrupted.")
                        return ""
            except Exception as pdf_error:
                st.error(f"**PDF Processing Error:** Unable to read the PDF file.")
                st.warning(f"Error details: {str(pdf_error)}")
                st.info("**Possible solutions:**\n- The PDF might be corrupted or password-protected\n- Try opening and re-saving the PDF\n- Convert it to DOCX or TXT format")
                return ""
                    
            return text.strip()

        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            file.seek(0)
            
            try:
                document = docx.Document(file)
                
                # Extract all text from paragraphs
                paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
                
                # Extract text from tables (often resumes have tables)
                table_text = []
                for table in document.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(" | ".join(row_text))
                
                # Combine paragraphs and table content
                all_text = "\n".join(paragraphs)
                if table_text:
                    all_text += "\n" + "\n".join(table_text)
                
                if not all_text.strip():
                    st.error("No text could be extracted from the DOCX file.")
                    st.info("The document might be empty or contain only images.")
                    return ""
                
                total_paragraphs = len(paragraphs) + len(table_text)
                st.info(f"Processing DOCX: {total_paragraphs} paragraph(s)/section(s) detected")
                
                return all_text.strip()
            except Exception as docx_error:
                st.error(f"**DOCX Processing Error:** Unable to read the Word document.")
                st.warning(f"Error details: {str(docx_error)}")
                st.info("**Possible solutions:**\n- The file might be corrupted\n- Try opening and re-saving the document\n- Save as .docx format (not .doc)\n- Convert to PDF or TXT format")
                return ""

        elif file.type == "text/plain":
            file.seek(0)
            
            try:
                text = file.read().decode("utf-8")
                
                if not text.strip():
                    st.error("The text file appears to be empty.")
                    return ""
                
                # Count lines for feedback
                line_count = len([line for line in text.split('\n') if line.strip()])
                st.info(f"Processing TXT: {line_count} line(s) detected")
                
                return text.strip()
            except UnicodeDecodeError:
                try:
                    # Try different encoding
                    file.seek(0)
                    text = file.read().decode("latin-1")
                    st.warning("File encoding detected as Latin-1 instead of UTF-8")
                    return text.strip()
                except Exception as encoding_error:
                    st.error(f"**Encoding Error:** Unable to read the text file.")
                    st.warning(f"Error details: {str(encoding_error)}")
                    st.info("**Possible solutions:**\n- Save the file with UTF-8 encoding\n- Copy content to a new text file\n- Convert to PDF or DOCX format")
                    return ""
            except Exception as txt_error:
                st.error(f"**TXT Processing Error:** Unable to read the text file.")
                st.warning(f"Error details: {str(txt_error)}")
                st.info("The file might be corrupted or in an unsupported format.")
                return ""
        else:
            st.error(f"**Unsupported file type:** {file.type}")
            st.warning("Please upload only PDF, DOCX, or TXT files.")
            return ""
        
    except Exception as e:
        st.error(f"**Unexpected Error** while processing {file.name}")
        st.warning(f"Error details: {str(e)}")
        st.info("**Try:**\n- Re-uploading the file\n- Using a different file format\n- Checking if the file is corrupted")
        return ""

def clean(text):
    """Clean and normalize text for NLP processing - optimized for longer documents"""
    if not text:
        return ""
    
    # Convert to lowercase
    text = text.lower()
    
    # Preserve experience ranges before cleaning
    text = re.sub(r'(\d+)\s*(?:-|–|to)\s*(\d+)', r'\1_\2', text)
    text = re.sub(r'(\d+)\s*\+', r'\1_plus', text)
    
    # Remove extra whitespace and newlines while preserving sentence structure
    text = re.sub(r'\n+', ' ', text)  # Replace multiple newlines with space
    text = re.sub(r'\t+', ' ', text)  # Replace tabs with space
    
    # Remove special characters but keep alphanumeric and underscores
    text = re.sub(r'[^a-z0-9_\s]', ' ', text)
    
    # Normalize multiple spaces to single space
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def extract_experience(text):
    """Extract years of experience from text - improved pattern matching"""
    if not text:
        return None
    text = text.lower()
    
    # Pattern 1: Experience range: 1-3 years, 2 to 5 yrs
    range_match = re.search(
        r'(\d+)\s*(?:-|–|to)\s*(\d+)\s*(?:years?|yrs?)',
        text
    )
    if range_match:
        return {"min_exp": int(range_match.group(1)), "max_exp": int(range_match.group(2))}
    
    # Pattern 2: Experience plus: 3+ years
    plus_match = re.search(
        r'(\d+)\s*\+\s*(?:years?|yrs?)',
        text
    )
    if plus_match:
        return {"min_exp": int(plus_match.group(1)), "max_exp": None}
    
    # Pattern 3: "with X years" or "X years of" or "X years specializing"
    with_years_match = re.search(
        r'(?:with|over|around|approximately)\s+(\d+)\s+(?:years?|yrs?)',
        text
    )
    if with_years_match:
        years = int(with_years_match.group(1))
        return {"min_exp": years, "max_exp": years}
    
    # Pattern 4: "X years of experience" or "X years specializing" or "X years in"
    years_of_match = re.search(
        r'(\d+)\s+(?:years?|yrs?)\s+(?:of\s+)?(?:experience|specializing|in|working)',
        text
    )
    if years_of_match:
        years = int(years_of_match.group(1))
        return {"min_exp": years, "max_exp": years}
    
    # Pattern 5: General "X year" or "X years" pattern (fallback)
    general_match = re.search(
        r'(\d+)\s+(?:years?|yrs?)',
        text
    )
    if general_match:
        years = int(general_match.group(1))
        return {"min_exp": years, "max_exp": years}
    
    return None

# ---------------- UPLOAD ----------------
st.subheader("Upload Documents")

left_col, right_col = st.columns(2)

with left_col:
    resume_file = st.file_uploader(
        "Upload Resume (Multi-page supported)",
        type=["pdf", "docx", "txt"],
        help="Upload your resume in PDF, DOCX, or TXT format. Multi-page documents are fully supported.",
        accept_multiple_files=False
    )
    if resume_file:
        # Additional file validation
        allowed_extensions = ['.pdf', '.docx', '.txt']
        file_extension = '.' + resume_file.name.split('.')[-1].lower()
        
        if file_extension not in allowed_extensions:
            st.error(f"**Invalid file format: {file_extension}**")
            st.warning("Please upload only PDF (.pdf), Word Document (.docx), or Text (.txt) files.")
            resume_file = None
        else:
            # Validate MIME type
            allowed_mime_types = {
                "application/pdf": "PDF",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "DOCX",
                "text/plain": "TXT"
            }
            
            if resume_file.type in allowed_mime_types:
                st.success(f"Resume uploaded: {resume_file.name} ({allowed_mime_types[resume_file.type]})")
            else:
                st.error(f"**Unsupported file type detected!**")
                st.warning(f"File type '{resume_file.type}' is not supported. Please upload PDF, DOCX, or TXT files only.")
                st.info("**Tip:** Make sure your file is saved in the correct format. Some files may have incorrect extensions.")
                resume_file = None

with right_col:
    jd_file = st.file_uploader(
        "Upload Job Description (Multi-page supported)",
        type=["pdf", "docx", "txt"],
        help="Upload job description in PDF, DOCX, or TXT format. Multi-page documents are fully supported.",
        accept_multiple_files=False
    )
    if jd_file:
        # Additional file validation
        allowed_extensions = ['.pdf', '.docx', '.txt']
        file_extension = '.' + jd_file.name.split('.')[-1].lower()
        
        if file_extension not in allowed_extensions:
            st.error(f"**Invalid file format: {file_extension}**")
            st.warning("Please upload only PDF (.pdf), Word Document (.docx), or Text (.txt) files.")
            jd_file = None
        else:
            # Validate MIME type
            allowed_mime_types = {
                "application/pdf": "PDF",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "DOCX",
                "text/plain": "TXT"
            }
            
            if jd_file.type in allowed_mime_types:
                st.success(f"Job Description uploaded: {jd_file.name} ({allowed_mime_types[jd_file.type]})")
            else:
                st.error(f"**Unsupported file type detected!**")
                st.warning(f"File type '{jd_file.type}' is not supported. Please upload PDF, DOCX, or TXT files only.")
                st.info("**Tip:** Make sure your file is saved in the correct format. Some files may have incorrect extensions.")
                jd_file = None

st.write("")

# ---------------- ANALYZE BUTTON ----------------
st.markdown("""
<style>
div[data-testid="stButton"] {
    display: flex;
    justify-content: center;
}

div[data-testid="stButton"] > button {
    background-color: #470047;
    color: white;
    width: 260px;
    height: 55px;
    font-size: 18px;
    border-radius: 10px;
    border: none;
}

div[data-testid="stButton"] > button:hover {
    background-color: #2E002E;
    transform: scale(1.03);
    transition: all 0.2s ease-in-out;
}
</style>
""", unsafe_allow_html=True)


d1, d2, d3, d4, d5 = st.columns(5)

with d3:
    analyze_button = st.button("Analyze Documents")



# analyze_button = st.button("🔍 Analyze Documents", type="secondary", use_container_width=True)

# ---------------- VALIDATION ----------------
if analyze_button:
    # Validation 1: Check if both files are uploaded
    if not resume_file and not jd_file:
        st.error("❌ **Error: No files uploaded!**")
        st.warning("⚠️ Please upload both Resume and Job Description files before analyzing.")
        st.stop()
    
    # Validation 2: Check if only resume is uploaded
    elif resume_file and not jd_file:
        st.error("❌ **Error: Job Description is missing!**")
        st.warning("⚠️ Please upload the Job Description file to proceed with the analysis.")
        st.stop()
    
    # Validation 3: Check if only JD is uploaded
    elif jd_file and not resume_file:
        st.error("❌ **Error: Resume is missing!**")
        st.warning("⚠️ Please upload the Resume file to proceed with the analysis.")
        st.stop()
    
    # Validation 4: Check file types (additional safety check)
    allowed_types = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain"]
    
    if resume_file.type not in allowed_types:
        st.error(f"❌ **Error: Invalid resume file format!**")
        st.warning(f"⚠️ Resume file type '{resume_file.type}' is not supported. Please upload PDF, DOCX, or TXT files only.")
        st.stop()
    
    if jd_file.type not in allowed_types:
        st.error(f"❌ **Error: Invalid job description file format!**")
        st.warning(f"⚠️ Job Description file type '{jd_file.type}' is not supported. Please upload PDF, DOCX, or TXT files only.")
        st.stop()
    
    # If all validations pass, proceed with analysis
    st.success("Both files uploaded successfully! Starting analysis...")
    st.write("")

if analyze_button and jd_file and resume_file:

    # ---------------- PREVIEW ----------------
    st.subheader("Parsed Document Preview")

    col1, col2 = st.columns(2)

    with col1:
        if resume_file:
            with st.spinner("Extracting resume text..."):
                resume_text = extract_text(resume_file)
            if resume_text:
                # Show character count and preview
                char_count = len(resume_text)
                word_count = len(resume_text.split())
                st.caption(f"📊 Resume: {char_count} characters, {word_count} words")
                st.text_area("Resume Preview", resume_text[:2000] + ("..." if len(resume_text) > 2000 else ""), 
                           height=200, key="resume_preview",
                           help=f"Showing first 2000 characters. Total: {char_count} characters")

    with col2:
        if jd_file:
            with st.spinner("Extracting job description text..."):
                jd_text = extract_text(jd_file)
            if jd_text:
                # Show character count and preview
                char_count = len(jd_text)
                word_count = len(jd_text.split())
                st.caption(f"Job Description: {char_count} characters, {word_count} words")
                st.text_area("Job Description Preview", jd_text[:2000] + ("..." if len(jd_text) > 2000 else ""), 
                           height=200, key="jd_preview",
                           help=f"Showing first 2000 characters. Total: {char_count} characters")

    # ---------------- PROCESS ----------------
    if resume_text:
        with st.spinner("Cleaning resume text..."):
            cleaned_resume = clean(resume_text)
            resume_experience = extract_experience(resume_text)

    if jd_text:
        with st.spinner("Cleaning job description text..."):
            cleaned_jd = clean(jd_text)
            jd_experience = extract_experience(jd_text)

    # ---------------- CLEANED OUTPUT ----------------
    st.subheader("Cleaned Files")

    col1, col2 = st.columns(2)

    with col1:
        if cleaned_resume:
            st.caption(f"Cleaned Resume: {len(cleaned_resume)} characters")
            st.text_area("Cleaned Resume", cleaned_resume[:2000] + ("..." if len(cleaned_resume) > 2000 else ""), 
                        height=200, key="cleaned_resume",
                        help=f"Showing first 2000 characters. Total: {len(cleaned_resume)} characters")
    with col2:
        if cleaned_jd:
            st.caption(f"Cleaned Job Description: {len(cleaned_jd)} characters")
            st.text_area("Cleaned Job Description", cleaned_jd[:2000] + ("..." if len(cleaned_jd) > 2000 else ""), 
                        height=200, key="cleaned_jd",
                        help=f"Showing first 2000 characters. Total: {len(cleaned_jd)} characters")

    st.success("Completed: Documents uploaded, parsed, cleaned and previewed successfully.")

    # -------------------------------------------------------------------------
    # Milestone - 2

    st.markdown(
        """
        <div style="background-color:#470047;padding:20px;border-radius:10px;margin-top:30px;">
            <h2 style="color:white;">
                Skill Extraction using NLP Module
            </h2>
            <p style="color:white;">
                Module: Skill Extraction using NLP <br>
                • spaCy and BERT-based pipelines  <br>
                • Technical and soft skills identification <br>
                • Structured skill display <br>
                • Optimized for multi-page documents
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.write("")

    st.subheader("Experience Analysis")

    col1, col2 = st.columns(2)
    with col1:
        if resume_experience:
            st.info(f"Resume Experience: {resume_experience['min_exp']}" + 
                   (f"-{resume_experience['max_exp']}" if resume_experience['max_exp'] else "+") + " years")
        else:
            st.warning("No experience information found in resume")
    with col2:
        if jd_experience:
            st.info(f"JD Required Experience: {jd_experience['min_exp']}" + 
                   (f"-{jd_experience['max_exp']}" if jd_experience['max_exp'] else "+") + " years")
        else:
            st.warning("No experience requirement found in JD")

    @st.cache_resource
    def load_spacy_model():
        try:
            return spacy.load("en_core_web_trf")
        except OSError:
            st.error("spaCy model 'en_core_web_trf' not found. Please install it using: python -m spacy download en_core_web_trf")
            return None
    
    nlp = load_spacy_model()
    
    if nlp is None:
        st.stop()

    master_technical_skills = {
        "python", "java", "c", "c++", "c#", "go", "rust", "kotlin", "swift",
        "javascript", "typescript", "php", "ruby", "r", "matlab", "scala",
        "perl", "bash", "shell scripting", "sql",
        "html", "css", "sass", "bootstrap", "tailwind css",
        "react", "angular", "vue.js", "next.js", "nuxt.js",
        "node.js", "express.js", "django", "flask", "fastapi",
        "spring", "spring boot", "laravel", "asp.net",
        "mysql", "postgresql", "sqlite", "oracle", "sql server",
        "mongodb", "cassandra", "dynamodb", "redis", "firebase",
        "elasticsearch", "neo4j",
        "numpy", "pandas", "scipy", "matplotlib", "seaborn",
        "plotly", "power bi", "tableau", "excel", "statistics",
        "data analysis", "data visualization",
        "machine learning", "deep learning", "artificial intelligence",
        "tensorflow", "keras", "pytorch", "scikit-learn", "xgboost",
        "opencv", "nlp", "computer vision", "speech recognition",
        "transformers", "huggingface", "langchain", "llm",
        "hadoop", "spark", "pyspark", "kafka", "hive", "pig",
        "hbase", "flink", "airflow", "databricks",
        "aws", "azure", "google cloud", "gcp",
        "docker", "kubernetes", "terraform", "ansible",
        "jenkins", "github actions", "ci/cd",
        "linux", "unix",
        "rest api", "graphql", "grpc", "soap",
        "microservices", "jwt", "oauth",
        "unit testing", "integration testing", "system testing",
        "pytest", "unittest", "junit", "selenium", "cypress",
        "postman",
        "android", "ios", "flutter", "react native",
        "swiftui", "kotlin multiplatform",
        "network security", "application security", "penetration testing",
        "ethical hacking", "cryptography", "owasp",
        "siem", "firewall",
        "windows", "linux", "macos",
        "git", "github", "gitlab", "bitbucket",
        "jira", "confluence", "trello",
        "data structures", "algorithms", "object oriented programming",
        "design patterns", "system design",
        "software development lifecycle", "agile", "scrum"
    }

    master_soft_skills = {
        "communication", "verbal communication", "written communication",
        "public speaking", "presentation skills", "active listening",
        "business communication", "storytelling",
        "teamwork", "collaboration", "interpersonal skills",
        "relationship building", "empathy", "emotional intelligence",
        "conflict resolution", "negotiation",
        "leadership", "people management", "team leadership",
        "decision making", "delegation", "mentoring", "coaching",
        "influencing", "strategic thinking",
        "problem solving", "critical thinking", "analytical thinking",
        "logical reasoning", "creative thinking", "innovation",
        "root cause analysis", "troubleshooting",
        "time management", "prioritization", "multitasking",
        "work ethic", "self discipline", "accountability",
        "goal setting", "organizational skills",
        "adaptability", "flexibility", "resilience",
        "learning agility", "continuous learning",
        "open mindedness", "growth mindset",
        "professionalism", "integrity", "ethical behavior",
        "reliability", "punctuality", "confidentiality",
        "creativity", "idea generation", "design thinking",
        "innovation mindset", "curiosity",
        "stress management", "self awareness", "self motivation",
        "confidence", "positive attitude", "emotional control",
        "customer focus", "customer service",
        "client management", "stakeholder management",
        "user empathy", "service mindset",
        "cross functional collaboration", "cultural awareness",
        "diversity and inclusion", "remote collaboration",
        "team alignment",
        "conflict management", "crisis management",
        "handling pressure", "decision making under pressure",
        "independent work", "collaborative work",
        "attention to detail", "quality focus",
        "result oriented", "ownership",
        "honesty", "trustworthiness", "respect",
        "fairness", "social responsibility"
    }

    def extract_soft_skills(text, skill_set):
        """Extract soft skills using simple text matching - optimized for longer text"""
        if not text:
            return []
        text = text.lower()
        found_skills = set()
        for skill in skill_set:
            if skill in text:
                found_skills.add(skill)
        return sorted(found_skills)

    def extract_technical_skills(text, skill_set):
        """Extract technical skills using spaCy PhraseMatcher - optimized for longer text"""
        if not text:
            return []
        
        # For very long documents, process in chunks to avoid memory issues
        max_length = 1000000  # spaCy's default max length
        
        if len(text) > max_length:
            # Process in chunks
            chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
            all_skills = set()
            
            for chunk in chunks:
                matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
                patterns = [nlp.make_doc(skill) for skill in skill_set]
                matcher.add("TECH_SKILLS", patterns)
                doc = nlp(chunk)
                matches = matcher(doc)
                for _, start, end in matches:
                    all_skills.add(doc[start:end].text.lower())
            
            return sorted(list(all_skills))
        else:
            # Process normally
            matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
            patterns = [nlp.make_doc(skill) for skill in skill_set]
            matcher.add("TECH_SKILLS", patterns)
            doc = nlp(text)
            matches = matcher(doc)
            skills_found = set()
            for _, start, end in matches:
                skills_found.add(doc[start:end].text.lower())
            return sorted(skills_found)

    def extract_github_username(url):
        """Extract username from GitHub URL"""
        match = re.search(r'github\.com/([^/]+)/?', url)
        return match.group(1) if match else None

    def detect_github_in_resume(text):
        """Detect GitHub profile URL or username in resume"""
        if not text:
            return None
        
        # Look for GitHub URLs
        url_pattern = r'(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9](?:[a-zA-Z0-9-]{0,38}[a-zA-Z0-9])?)'
        url_match = re.search(url_pattern, text, re.IGNORECASE)
        
        if url_match:
            return f"https://github.com/{url_match.group(1)}"
        
        # Look for standalone GitHub username patterns
        username_pattern = r'(?:github|gh)[:\s]+([a-zA-Z0-9](?:[a-zA-Z0-9-]{0,38}[a-zA-Z0-9])?)'
        username_match = re.search(username_pattern, text, re.IGNORECASE)
        
        if username_match:
            return f"https://github.com/{username_match.group(1)}"
        
        return None

    def fetch_github_repos(username):
        """Fetch public repositories of a GitHub user"""
        api_url = f"https://api.github.com/users/{username}/repos"
        try:
            response = requests.get(api_url, timeout=10)
            if response.status_code == 200:
                return response.json()
        except:
            pass
        return []

    def extract_github_skills(repos, skill_set):
        """Extract skills from repo names, descriptions, and languages"""
        found_skills = set()
        for repo in repos:
            text = f"{repo.get('name','')} {repo.get('description','')}".lower()
            for skill in skill_set:
                if skill in text:
                    found_skills.add(skill)
            if repo.get("language"):
                lang = repo["language"].lower()
                if lang in skill_set:
                    found_skills.add(lang)
        return sorted(found_skills)

    def analyze_github_profile(repos):
        """Analyze GitHub profile and provide suggestions"""
        if not repos:
            return None
        
        total_repos = len(repos)
        languages_used = {}
        topics_used = set()
        has_readme_projects = 0
        recent_activity = 0
        
        for repo in repos:
            # Count languages
            if repo.get("language"):
                lang = repo["language"]
                languages_used[lang] = languages_used.get(lang, 0) + 1
            
            # Count topics
            if repo.get("topics"):
                topics_used.update(repo["topics"])
            
            # Check for README
            if repo.get("description"):
                has_readme_projects += 1
            
            # Check recent activity (updated in last 6 months)
            if repo.get("updated_at"):
                from datetime import datetime, timedelta
                try:
                    updated = datetime.strptime(repo["updated_at"], "%Y-%m-%dT%H:%M:%SZ")
                    if updated > datetime.now() - timedelta(days=180):
                        recent_activity += 1
                except:
                    pass
        
        # Generate insights
        insights = {
            "total_repos": total_repos,
            "languages": languages_used,
            "top_language": max(languages_used.items(), key=lambda x: x[1])[0] if languages_used else None,
            "topics": list(topics_used),
            "documented_projects": has_readme_projects,
            "active_repos": recent_activity,
            "activity_rate": round((recent_activity / total_repos) * 100, 1) if total_repos > 0 else 0
        }
        
        # Generate suggestions
        suggestions = []
        
        if has_readme_projects < total_repos * 0.5:
            suggestions.append("Add detailed README files to more projects to showcase your work better")
        
        if recent_activity < 3:
            suggestions.append("Increase recent GitHub activity - consider contributing to open source or creating new projects")
        
        if len(languages_used) < 3:
            suggestions.append("Diversify your tech stack by learning new programming languages")
        
        if not topics_used:
            suggestions.append("Add topics/tags to your repositories for better discoverability")
        
        if total_repos < 5:
            suggestions.append("Build more projects to demonstrate your skills and experience")
        
        return {
            "insights": insights,
            "suggestions": suggestions
        }

    with st.spinner("Extracting skills using NLP... (This may take a moment for longer documents)"):
        resume_technical = extract_technical_skills(cleaned_resume, master_technical_skills)
        resume_soft = extract_soft_skills(cleaned_resume, master_soft_skills)
        
        # Initialize GitHub skills
        github_skills = []
        github_analysis = None
        
        # Auto-detect GitHub profile from resume
        detected_github_url = detect_github_in_resume(resume_text)
        
        # GitHub Profile Analyzer Section
        st.subheader("GitHub Profile Analyzer")
        
        if detected_github_url:
            st.success(f"GitHub profile detected in resume: {detected_github_url}")
            auto_analyze = st.checkbox("Automatically analyze detected GitHub profile", value=True)
            
            if auto_analyze:
                github_url = detected_github_url
            else:
                github_url = st.text_input(
                    "Or enter different GitHub Profile URL",
                    value=detected_github_url,
                    placeholder="https://github.com/username",
                    help="Edit or use a different GitHub profile"
                )
        else:
            st.caption("No GitHub profile detected in resume. You can manually add one below.")
            github_url = st.text_input(
                "Enter GitHub Profile URL (Optional)",
                placeholder="https://github.com/username",
                help="We'll extract technical skills from your public repositories"
            )
        
        if github_url:
            username = extract_github_username(github_url)
            if username:
                with st.spinner(f"🔍 Analyzing GitHub profile for @{username}..."):
                    repos = fetch_github_repos(username)
                    if repos:
                        # Extract skills
                        github_skills = extract_github_skills(repos, master_technical_skills)
                        
                        # Analyze profile
                        github_analysis = analyze_github_profile(repos)
                        
                        if github_skills:
                            st.success(f"Found {len(github_skills)} technical skills from {len(repos)} repositories!")
                            
                            # Display GitHub Skills
                            with st.expander("View Extracted GitHub Skills", expanded=True):
                                st.write(", ".join(github_skills))
                            
                            # Display GitHub Insights
                            if github_analysis:
                                with st.expander("GitHub Profile Insights", expanded=True):
                                    insights = github_analysis["insights"]
                                    
                                    col_i1, col_i2, col_i3, col_i4 = st.columns(4)
                                    col_i1.metric("Total Repos", insights["total_repos"])
                                    col_i2.metric("Languages", len(insights["languages"]))
                                    col_i3.metric("Documented", f"{insights['documented_projects']}/{insights['total_repos']}")
                                    col_i4.metric("Activity Rate", f"{insights['activity_rate']}%")
                                    
                                    if insights["top_language"]:
                                        st.info(f"🏆 Most used language: **{insights['top_language']}** ({insights['languages'][insights['top_language']]} repos)")
                                    
                                    if insights["languages"]:
                                        st.write("**Language Distribution:**")
                                        lang_text = ", ".join([f"{lang} ({count})" for lang, count in sorted(insights["languages"].items(), key=lambda x: x[1], reverse=True)])
                                        st.caption(lang_text)
                                
                                # Display Suggestions
                                if github_analysis["suggestions"]:
                                    with st.expander("GitHub Profile Improvement Suggestions", expanded=True):
                                        st.markdown("**Recommendations to enhance your GitHub presence:**")
                                        for suggestion in github_analysis["suggestions"]:
                                            st.warning(f"💡 {suggestion}")
                        else:
                            st.warning("No recognizable technical skills found in your repositories")
                            st.info("Tip: Add more descriptive repository names and detailed descriptions")
                    else:
                        st.error("Unable to fetch repositories. Please check the username or try again later.")
            else:
                st.error("Invalid GitHub URL format")
        
        # Merge GitHub skills with resume skills
        combined_technical_skills = sorted(set(resume_technical) | set(github_skills))
        
        resume_skills = {
            "technical": combined_technical_skills,
            "soft": resume_soft
        }
        
        # Show combined stats if GitHub was used
        if github_skills:
            st.markdown("---")
            st.subheader("Skill Integration Summary")
            col_a, col_b, col_c, col_d = st.columns(4)
            col_a.metric("Resume Skills", len(resume_technical))
            col_b.metric("GitHub Skills", len(github_skills))
            col_c.metric("Total Unique", len(combined_technical_skills))
            col_d.metric("Skills Added", len(combined_technical_skills) - len(resume_technical))

        jd_skills = {
            "technical": extract_technical_skills(cleaned_jd, master_technical_skills),
            "soft": extract_soft_skills(cleaned_jd, master_soft_skills)
        }

    st.subheader("Extracted Skills")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### Resume Skills")
        with st.expander("Technical Skills", expanded=True):
            if resume_skills["technical"]:
                st.write(", ".join(resume_skills["technical"]))
                st.caption(f"Total: {len(resume_skills['technical'])} skills")
            else:
                st.info("No technical skills found")
        with st.expander("Soft Skills", expanded=True):
            if resume_skills["soft"]:
                st.write(", ".join(resume_skills["soft"]))
                st.caption(f"Total: {len(resume_skills['soft'])} skills")
            else:
                st.info("No soft skills found")

    with col2:
        st.markdown("### JD Skills")
        with st.expander("Technical Skills", expanded=True):
            if jd_skills["technical"]:
                st.write(", ".join(jd_skills["technical"]))
                st.caption(f"Total: {len(jd_skills['technical'])} skills")
            else:
                st.info("No technical skills found")
        with st.expander("Soft Skills", expanded=True):
            if jd_skills["soft"]:
                st.write(", ".join(jd_skills["soft"]))
                st.caption(f"Total: {len(jd_skills['soft'])} skills")
            else:
                st.info("No soft skills found")

    def donut_chart(tech_count, soft_count, title):
        """Create donut chart for skill distribution"""
        labels = ["Technical Skills", "Soft Skills"]
        sizes = [tech_count, soft_count]
        if sum(sizes) == 0:
            sizes = [1, 1]
        fig = plt.figure(figsize=(4, 4))
        ax = fig.add_axes([0.15, 0.15, 0.7, 0.7]) 
        wedges, _, _ = ax.pie(
            sizes,
            startangle=90,
            autopct="%1.0f%%",
            radius=1,
            wedgeprops=dict(width=0.4, edgecolor="white"),
            colors=['#470047', '#fd7e14']
        )
        ax.set(aspect="equal")
        ax.set_title(title, pad=10, fontsize=12, fontweight='bold')
        ax.legend(
            wedges,
            labels,
            loc="lower center",
            bbox_to_anchor=(0.5, -0.25),
            ncol=2,
            frameon=False
        )
        return fig

    st.subheader("Skill Distribution Analysis")

    col1, col2 = st.columns(2)

    with col1:
        fig1 = donut_chart(
            len(resume_skills["technical"]),
            len(resume_skills["soft"]),
            "Resume Skill Distribution"
        )
        st.pyplot(fig1, use_container_width=False)
        plt.close(fig1)

    with col2:
        fig2 = donut_chart(
            len(jd_skills["technical"]),
            len(jd_skills["soft"]),
            "JD Skill Distribution"
        )
        st.pyplot(fig2, use_container_width=False)
        plt.close(fig2)

    def compute_metrics(resume_skills, jd_skills):
        """Compute skill matching metrics"""
        resume_tech = set(resume_skills["technical"])
        resume_soft = set(resume_skills["soft"])
        jd_tech = set(jd_skills["technical"])
        jd_soft = set(jd_skills["soft"])
        matched_skills = (resume_tech & jd_tech) | (resume_soft & jd_soft)
        total_jd_skills = len(jd_tech) + len(jd_soft)
        match_percentage = 0
        if total_jd_skills > 0:
            match_percentage = round((len(matched_skills) / total_jd_skills) * 100, 1)
        return {
            "resume_tech": len(resume_tech),
            "resume_soft": len(resume_soft),
            "jd_tech": len(jd_tech),
            "jd_soft": len(jd_soft),
            "total_resume": len(resume_tech) + len(resume_soft),
            "total_jd": total_jd_skills,
            "match_percent": match_percentage
        }
    
    metrics = compute_metrics(resume_skills, jd_skills)

    st.subheader("Skill Extraction Metrics")

    col1, col2, col3, col4 = st.columns(4)
    col1.metric(label="JD Technical Skills", value=metrics["jd_tech"])
    col2.metric(label="JD Soft Skills", value=metrics["jd_soft"])
    col3.metric(label="Total JD Skills", value=metrics["total_jd"])
    col4.metric(label="Basic Match %", value=f'{metrics["match_percent"]}%')

    st.success("Completed: Skills extracted successfully using NLP.")

    # --------------------------------------------------------------------------
    # Milestone - 3

    def categorize_skills(skills):
        """Categorize skills into meaningful groups"""
        categories = {
            "Programming Languages": ["python", "java", "c", "c++", "c#", "go", "rust", "kotlin", "swift",
                                     "javascript", "typescript", "php", "ruby", "r", "matlab", "scala", "perl"],
            "Web Technologies": ["html", "css", "sass", "bootstrap", "tailwind css", "react", "angular", 
                                "vue.js", "next.js", "nuxt.js", "node.js", "express.js", "django", "flask", 
                                "fastapi", "spring", "spring boot", "laravel", "asp.net"],
            "Databases": ["mysql", "postgresql", "sqlite", "oracle", "sql server", "mongodb", "cassandra", 
                         "dynamodb", "redis", "firebase", "elasticsearch", "neo4j", "sql"],
            "Data Science & ML": ["numpy", "pandas", "scipy", "matplotlib", "seaborn", "plotly", "power bi", 
                                 "tableau", "excel", "machine learning", "deep learning", "tensorflow", "keras", 
                                 "pytorch", "scikit-learn", "xgboost", "opencv", "nlp", "computer vision", 
                                 "transformers", "huggingface", "langchain", "llm", "artificial intelligence",
                                 "statistics", "data analysis", "data visualization"],
            "Cloud & DevOps": ["aws", "azure", "google cloud", "gcp", "docker", "kubernetes", "terraform", 
                              "ansible", "jenkins", "github actions", "ci/cd", "linux", "unix"],
            "Big Data": ["hadoop", "spark", "pyspark", "kafka", "hive", "pig", "hbase", "flink", 
                        "airflow", "databricks"],
            "Mobile Dev": ["android", "ios", "flutter", "react native", "swiftui", "kotlin multiplatform"],
            "Testing & QA": ["unit testing", "integration testing", "system testing", "pytest", "unittest", 
                            "junit", "selenium", "cypress", "postman"],
            "APIs & Backend": ["rest api", "graphql", "grpc", "soap", "microservices", "jwt", "oauth"],
            "Security": ["network security", "application security", "penetration testing",
                        "ethical hacking", "cryptography", "owasp", "siem", "firewall"],
            "Leadership & Management": ["leadership", "people management", "team leadership", "decision making", 
                                       "delegation", "mentoring", "coaching", "strategic thinking"],
            "Communication": ["communication", "verbal communication", "written communication", "public speaking", 
                             "presentation skills", "active listening", "storytelling", "business communication"],
            "Problem Solving": ["problem solving", "critical thinking", "analytical thinking", "logical reasoning", 
                               "creative thinking", "innovation", "troubleshooting", "root cause analysis"],
            "Work Management": ["time management", "prioritization", "multitasking", "work ethic", 
                               "self discipline", "accountability", "goal setting", "organizational skills"],
            "Adaptability": ["adaptability", "flexibility", "resilience", "learning agility", 
                            "continuous learning", "open mindedness", "growth mindset"],
            "Collaboration": ["teamwork", "collaboration", "interpersonal skills", "relationship building",
                             "cross functional collaboration", "cultural awareness", "diversity and inclusion",
                             "remote collaboration", "team alignment"],
            "Soft Skills": ["empathy", "emotional intelligence", "conflict resolution", "negotiation",
                           "professionalism", "integrity", "customer focus", "customer service"]
        }
        
        categorized = {}
        uncategorized = []
        
        for skill in skills:
            skill_lower = skill.lower()
            found = False
            for category, category_skills in categories.items():
                if skill_lower in category_skills:
                    if category not in categorized:
                        categorized[category] = []
                    categorized[category].append(skill)
                    found = True
                    break
            if not found:
                uncategorized.append(skill)
        
        if uncategorized:
            categorized["Other Skills"] = uncategorized
        
        return categorized

    def build_similarity_matrix(resume_skills, jd_skills):
        """Build similarity matrix between resume and JD skills"""
        if not resume_skills or not jd_skills:
            return np.array([])
        matrix = np.zeros((len(resume_skills), len(jd_skills)))
        for i, r_skill in enumerate(resume_skills):
            for j, j_skill in enumerate(jd_skills):
                if r_skill == j_skill:
                    matrix[i][j] = 1.0
                elif r_skill in j_skill or j_skill in r_skill:
                    matrix[i][j] = 0.5
                else:
                    matrix[i][j] = 0.0
        return matrix

    def plot_category_match_heatmap(resume_skills, jd_skills):
        """Plot single consolidated heatmap showing category-wise skill match"""
        jd_categorized = categorize_skills(jd_skills)
        resume_categorized = categorize_skills(resume_skills)
        
        common_categories = set(jd_categorized.keys()) & set(resume_categorized.keys())
        
        if not common_categories:
            st.warning("No common skill categories found between resume and JD")
            return None
        
        # Build category-wise match matrix
        categories = sorted(common_categories)
        match_data = []
        
        for category in categories:
            jd_cat_skills = jd_categorized[category]
            resume_cat_skills = set(resume_categorized[category])
            
            matched = 0
            partial = 0
            missing = 0
            
            for jd_skill in jd_cat_skills:
                if jd_skill in resume_cat_skills:
                    matched += 1
                else:
                    # Check for partial match
                    found_partial = False
                    for res_skill in resume_cat_skills:
                        if jd_skill in res_skill or res_skill in jd_skill:
                            partial += 1
                            found_partial = True
                            break
                    if not found_partial:
                        missing += 1
            
            total = len(jd_cat_skills)
            
            match_data.append({
                'Category': category,
                'Matched': matched,
                'Partial': partial,
                'Missing': missing,
                'Total': total
            })
        
        # Create stacked bar chart
        fig = go.Figure()
        
        categories_list = [d['Category'] for d in match_data]
        
        fig.add_trace(go.Bar(
            name='Matched',
            x=categories_list,
            y=[d['Matched'] for d in match_data],
            marker_color='#28a745',
            text=[d['Matched'] for d in match_data],
            textposition='inside',
            hovertemplate='<b>%{x}</b><br>Matched: %{y}<extra></extra>'
        ))
        
        fig.add_trace(go.Bar(
            name='Partial Match',
            x=categories_list,
            y=[d['Partial'] for d in match_data],
            marker_color='#ffc107',
            text=[d['Partial'] for d in match_data],
            textposition='inside',
            hovertemplate='<b>%{x}</b><br>Partial: %{y}<extra></extra>'
        ))
        
        fig.add_trace(go.Bar(
            name='Missing',
            x=categories_list,
            y=[d['Missing'] for d in match_data],
            marker_color='#dc3545',
            text=[d['Missing'] for d in match_data],
            textposition='inside',
            hovertemplate='<b>%{x}</b><br>Missing: %{y}<extra></extra>'
        ))
        
        fig.update_layout(
            title="Category-wise Skill Match Analysis (Stacked)",
            barmode='stack',
            height=500,
            xaxis_title="Skill Categories",
            yaxis_title="Number of Skills",
            hovermode='x unified',
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1
            ),
            showlegend=True
        )
        
        return fig

    resume_all_skills = resume_skills["technical"] + resume_skills["soft"]
    jd_all_skills = jd_skills["technical"] + jd_skills["soft"]

    st.markdown("""
        <div style="background-color:#470047;padding:20px;border-radius:10px;margin-top:30px;">
            <h2 style="color:white;">Skill Gap Analysis and Similarity Matching Module</h2>
            <p style="color:white;">• Skill similarity matrix visualization <br>• Resume vs JD skill comparison <br>• Missing skill identification <br>• Multi-page document analysis</p>
        </div><br>""", unsafe_allow_html=True)

    if resume_all_skills and jd_all_skills:
        with st.spinner("Building skill gap analysis..."):
            similarity_matrix = build_similarity_matrix(resume_all_skills, jd_all_skills)
            st.subheader("Category-wise Skill Match Heatmap")
            fig_heatmap = plot_category_match_heatmap(resume_all_skills, jd_all_skills)
            if fig_heatmap:
                st.plotly_chart(fig_heatmap, use_container_width=True)
    else:
        st.warning("Insufficient skills to build skill gap analysis")

    def classify_skill_matches(similarity_matrix, resume_skills, jd_skills):
        """Classify skills as matched, partial, or missing"""
        matched = set()
        partial = set()
        for j, jd_skill in enumerate(jd_skills):
            column = similarity_matrix[:, j]
            if 1.0 in column:
                matched.add(jd_skill)
            elif 0.5 in column:
                partial.add(jd_skill)
        missing = set(jd_skills) - matched - partial
        return {
            "matched": sorted(list(matched)),
            "partial": sorted(list(partial)),
            "missing": sorted(list(missing))
        }

    skill_match_result = classify_skill_matches(similarity_matrix, resume_all_skills, jd_all_skills)

    def calculate_skill_match(resume_skills, jd_skills):
        """Calculate detailed skill match statistics"""
        resume_set = set(resume_skills)
        jd_set = set(jd_skills)
        matched = resume_set & jd_set
        missing = jd_set - resume_set
        partial = set()
        for jd_skill in missing.copy():
            for res_skill in resume_set:
                if jd_skill in res_skill or res_skill in jd_skill:
                    partial.add(jd_skill)
        missing = missing - partial
        avg_match = 0
        if len(jd_set) > 0:
            avg_match = round(((len(matched) + 0.5 * len(partial)) / len(jd_set)) * 100, 1)
        return {
            "matched": len(matched),
            "partial": len(partial),
            "missing": len(missing),
            "avg_match": avg_match
        }

    all_resume_skills = set(resume_skills["technical"]) | set(resume_skills["soft"])
    all_jd_skills = set(jd_skills["technical"]) | set(jd_skills["soft"])
    match_counts = calculate_skill_match(all_resume_skills, all_jd_skills)

    st.subheader("Skill Match Overview")

    left, right = st.columns([1.3, 1])

    with left:
        labels = ["Matched", "Partially Matched", "Missing"]
        sizes = [match_counts["matched"], match_counts["partial"], match_counts["missing"]]
        colors = ["#28a745", "#fd7e14", "#dc3545"]
        if sum(sizes) == 0:
            sizes = [1, 1, 1]
        fig = plt.figure(figsize=(3.6, 3.6))
        ax = fig.add_axes([0.1, 0.15, 0.8, 0.75])
        wedges, _, autotexts = ax.pie(sizes, startangle=90, autopct="%1.0f%%", radius=1, colors=colors, wedgeprops=dict(width=0.4, edgecolor="white"))
        for autotext in autotexts:
            autotext.set_color('black')
            autotext.set_fontweight('bold')
            autotext.set_fontsize(11)
        ax.set(aspect="equal")
        ax.set_title("Skill Match Distribution", pad=8, fontsize=12, fontweight='bold')
        ax.legend(wedges, labels, loc="lower center", bbox_to_anchor=(0.5, -0.25), ncol=3, frameon=False)
        st.pyplot(fig, use_container_width=False)
        plt.close(fig)

    with right:
        with st.container():
            r1, r2 = st.columns(2)
            r1.metric("Matched Skills", match_counts["matched"])
            r2.metric("Partially Matched", match_counts["partial"])
            r3, r4 = st.columns(2)
            r3.metric("Missing Skills", match_counts["missing"])
            r4.metric("Avg Match %", f'{match_counts["avg_match"]}%')

    st.subheader("Skill Gap Details")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("### Matched Skills")
        if skill_match_result["matched"]:
            for skill in skill_match_result["matched"]:
                st.success(f"✓ {skill}")
        else:
            st.info("No perfectly matched skills")

    with col2:
        st.markdown("### Partial Matches")
        if skill_match_result["partial"]:
            for skill in skill_match_result["partial"]:
                st.warning(f"≈ {skill}")
        else:
            st.info("No partially matched skills")

    with col3:
        st.markdown("### Missing Skills")
        if skill_match_result["missing"]:
            for skill in skill_match_result["missing"]:
                st.error(f"✗ {skill}")
        else:
            st.success("No missing skills!")

    st.success("Completed: Skill gap analysis completed successfully.")

    # -----------------------------------------------------------------------------
    # Milestone - 4

    st.markdown(
        """
        <div style="background-color:#470047;padding:20px;border-radius:10px;margin-top:30px;">
            <h2 style="color:white;">
                Dashboard and Report Export Module
            </h2>
            <p style="color:white;">
                Interactive dashboard • Graphs • Multi-format report export • Optimized for multi-page analysis
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.write("")

    skills = list(all_jd_skills)
    resume_scores = [100 if skill in all_resume_skills else 0 for skill in skills]
    job_scores = [100 for _ in skills]
    overall_match = match_counts["avg_match"]
    matched_skills_count = match_counts["matched"]
    missing_skills_count = match_counts["missing"]

    c1, c2, c3 = st.columns(3)
    c1.metric("Overall Match", f"{overall_match}%")
    c2.metric("Matched Skills", str(matched_skills_count))
    c3.metric("Missing Skills", str(missing_skills_count))

    st.subheader("Categorized Skill Match Overview")
    
    jd_categorized = categorize_skills(list(all_jd_skills))
    categories_to_plot = []
    resume_category_scores = []
    jd_category_scores = []
    
    for category, cat_skills in jd_categorized.items():
        matched_in_category = sum(1 for skill in cat_skills if skill in all_resume_skills)
        total_in_category = len(cat_skills)
        if total_in_category > 0:
            categories_to_plot.append(category)
            resume_category_scores.append((matched_in_category / total_in_category) * 100)
            jd_category_scores.append(100)
    
    bar_fig = go.Figure()
    bar_fig.add_trace(go.Bar(name='Your Skills', x=categories_to_plot, y=resume_category_scores, marker_color='#470047', text=[f'{score:.0f}%' for score in resume_category_scores], textposition='outside'))
    bar_fig.add_trace(go.Bar(name='Job Requirements', x=categories_to_plot, y=jd_category_scores, marker_color='#28a745', text=['100%'] * len(categories_to_plot), textposition='outside'))
    bar_fig.update_layout(title="Skill Match by Category", barmode="group", height=500, xaxis_title="Skill Categories", yaxis_title="Coverage (%)", hovermode='x unified', yaxis=dict(range=[0, 120]), legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1))
    st.plotly_chart(bar_fig, use_container_width=True)

    radar_categories = categories_to_plot[:6] if len(categories_to_plot) > 6 else categories_to_plot
    radar_resume_scores = resume_category_scores[:6] if len(resume_category_scores) > 6 else resume_category_scores
    radar_jd_scores = [100] * len(radar_categories)
    
    radar_fig = go.Figure()
    radar_fig.add_trace(go.Scatterpolar(r=radar_resume_scores, theta=radar_categories, fill='toself', name='Current Profile', line_color='#470047'))
    radar_fig.add_trace(go.Scatterpolar(r=radar_jd_scores, theta=radar_categories, fill='toself', name='Job Requirements', line_color='#28a745'))
    radar_fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), title="Top Skill Categories Comparison", height=500, showlegend=True)
    st.plotly_chart(radar_fig, use_container_width=True)

    st.subheader("Sample Skill Proficiency")
    st.caption("Note: These are sample scores for demonstration. For actual proficiency, use skill assessment tools.")
    
    sample_skills = [("Python", 92), ("Machine Learning", 88), ("SQL", 65)]
    for skill, score in sample_skills:
        col_skill, col_progress = st.columns([1, 4])
        with col_skill:
            st.write(f"**{skill}**")
        with col_progress:
            st.progress(score / 100)
            st.caption(f"{score}%")

    st.subheader("Category-wise Similarity Score Distribution")
    
    if resume_all_skills and jd_all_skills and len(similarity_matrix) > 0:
        # Categorize skills and calculate average similarity per category
        jd_categorized = categorize_skills(jd_all_skills)
        
        category_similarities = {}
        
        for category, cat_skills in jd_categorized.items():
            cat_scores = []
            for skill in cat_skills:
                if skill in jd_all_skills:
                    j = jd_all_skills.index(skill)
                    column_scores = similarity_matrix[:, j]
                    avg_score = np.mean(column_scores) * 100
                    cat_scores.append(avg_score)
            
            if cat_scores:
                category_similarities[category] = np.mean(cat_scores)
        
        if category_similarities:
            categories = list(category_similarities.keys())
            scores = list(category_similarities.values())
            
            # Create area chart using plotly
            area_fig = go.Figure()
            area_fig.add_trace(go.Scatter(
                x=categories,
                y=scores,
                fill='tozeroy',
                name='Avg Similarity Score',
                line_color='#470047',
                fillcolor='rgba(124, 58, 237, 0.3)',
                mode='lines+markers'
            ))
            
            area_fig.update_layout(
                title="Average Similarity Score by Category",
                xaxis_title="Skill Categories",
                yaxis_title="Similarity Score (%)",
                height=400,
                hovermode='x unified',
                yaxis=dict(range=[0, 100])
            )
            
            st.plotly_chart(area_fig, use_container_width=True)
        else:
            st.info("No category data available for similarity analysis")
    else:
        st.info("Upload documents to see similarity score distribution")

    st.subheader("Upskilling Recommendations")
    st.caption("Based on missing and partially matched skills from job description")
    
    if skill_match_result["missing"]:
        st.markdown("**🔴 Priority Skills to Learn (Missing):**")
        for i, skill in enumerate(skill_match_result["missing"][:5], 1):
            st.error(f"{i}. **{skill.title()}** - Not found in your resume")
    
    if skill_match_result["partial"]:
        st.markdown("**🟡 Skills to Strengthen (Partial Match):**")
        for i, skill in enumerate(skill_match_result["partial"][:5], 1):
            st.warning(f"{i}. **{skill.title()}** - Improve proficiency in this area")
    
    if not skill_match_result["missing"] and not skill_match_result["partial"]:
        st.success("Excellent! You have all the required skills for this job!")

    st.subheader("Report Download")

    col1, col2, col3 = st.columns(3)

    with col1:
        # Separate technical and soft skills
        tech_matched = [s for s in skill_match_result['matched'] if s in jd_skills["technical"]]
        tech_partial = [s for s in skill_match_result['partial'] if s in jd_skills["technical"]]
        tech_missing = [s for s in skill_match_result['missing'] if s in jd_skills["technical"]]
        
        soft_matched = [s for s in skill_match_result['matched'] if s in jd_skills["soft"]]
        soft_partial = [s for s in skill_match_result['partial'] if s in jd_skills["soft"]]
        soft_missing = [s for s in skill_match_result['missing'] if s in jd_skills["soft"]]
        
        # Pad lists to same length within each category
        max_tech_len = max(len(tech_matched), len(tech_partial), len(tech_missing), 1)
        tech_matched += [''] * (max_tech_len - len(tech_matched))
        tech_partial += [''] * (max_tech_len - len(tech_partial))
        tech_missing += [''] * (max_tech_len - len(tech_missing))
        
        max_soft_len = max(len(soft_matched), len(soft_partial), len(soft_missing), 1)
        soft_matched += [''] * (max_soft_len - len(soft_matched))
        soft_partial += [''] * (max_soft_len - len(soft_partial))
        soft_missing += [''] * (max_soft_len - len(soft_missing))
        
        # Create dataframe with technical skills first, then gap, then soft skills
        tech_df = pd.DataFrame({
            "Technical - Matched": tech_matched,
            "Technical - Partially Matched": tech_partial,
            "Technical - Missing": tech_missing
        })
        
        soft_df = pd.DataFrame({
            "Soft Skills - Matched": soft_matched,
            "Soft Skills - Partially Matched": soft_partial,
            "Soft Skills - Missing": soft_missing
        })
        
        # Combine: technical, separator, soft
        csv_output = tech_df.to_csv(index=False)
        csv_output += "\n"  # Add blank line
        csv_output += soft_df.to_csv(index=False)
        
        st.download_button(
            "⬇Download CSV Report",
            csv_output.encode("utf-8"),
            "skill_gap_report.csv",
            "text/csv",
            help="Download skill comparison separated by Technical and Soft Skills"
        )

    with col2:
        def generate_word_report():
            doc = Document()
            title = doc.add_heading('Skill Gap Analysis Report', 0)
            title.alignment = 1
            doc.add_heading('Executive Summary', 1)
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Light Grid Accent 1'
            summary_data = [
                ('Overall Match Rate', f'{overall_match}%'),
                ('Matched Skills', str(matched_skills_count)),
                ('Partially Matched Skills', str(match_counts["partial"])),
                ('Missing Skills', str(missing_skills_count))
            ]
            for i, (label, value) in enumerate(summary_data):
                summary_table.rows[i].cells[0].text = label
                summary_table.rows[i].cells[1].text = value
            doc.add_paragraph()
            
            # Technical Skills Section
            doc.add_heading('Technical Skills Breakdown', 1)
            
            tech_matched = [s for s in skill_match_result['matched'] if s in jd_skills["technical"]]
            tech_partial = [s for s in skill_match_result['partial'] if s in jd_skills["technical"]]
            tech_missing = [s for s in skill_match_result['missing'] if s in jd_skills["technical"]]
            
            doc.add_heading('Matched Technical Skills', 2)
            if tech_matched:
                for skill in tech_matched:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No perfectly matched technical skills found.')
            
            doc.add_heading('Partially Matched Technical Skills', 2)
            if tech_partial:
                for skill in tech_partial:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No partially matched technical skills found.')
            
            doc.add_heading('Missing Technical Skills', 2)
            if tech_missing:
                for skill in tech_missing:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No missing technical skills!')
            
            # Soft Skills Section
            doc.add_heading('Soft Skills Breakdown', 1)
            
            soft_matched = [s for s in skill_match_result['matched'] if s in jd_skills["soft"]]
            soft_partial = [s for s in skill_match_result['partial'] if s in jd_skills["soft"]]
            soft_missing = [s for s in skill_match_result['missing'] if s in jd_skills["soft"]]
            
            doc.add_heading('Matched Soft Skills', 2)
            if soft_matched:
                for skill in soft_matched:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No perfectly matched soft skills found.')
            
            doc.add_heading('Partially Matched Soft Skills', 2)
            if soft_partial:
                for skill in soft_partial:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No partially matched soft skills found.')
            
            doc.add_heading('Missing Soft Skills', 2)
            if soft_missing:
                for skill in soft_missing:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No missing soft skills!')
            
            # Recommendations
            doc.add_heading('Recommendations', 1)
            doc.add_paragraph(
                f'Your current match rate is {overall_match}%. '
                f'To improve your candidacy, consider focusing on the missing skills listed above.'
            )
            
            if overall_match >= 70:
                doc.add_paragraph('You have a strong skill match for this position!', style='List Bullet')
            else:
                doc.add_paragraph('Focus on upskilling in the missing areas to increase your match rate.', style='List Bullet')
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph(f'Generated on: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
            footer.alignment = 1
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
        doc_file = generate_word_report()
        st.download_button("⬇Download DOCX Report", doc_file, "skill_gap_analysis_report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", help="Download formatted Word report with complete analysis")
    
    with col3:
        def generate_pdf_report():
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", 'B', 20)
            pdf.cell(0, 15, "Skill Gap Analysis Report", ln=True, align='C')
            pdf.ln(5)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Executive Summary", ln=True)
            pdf.set_font("Arial", '', 11)
            pdf.ln(2)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(90, 8, "Overall Match Rate:", 1, 0, fill=True)
            pdf.cell(90, 8, f"{overall_match}%", 1, 1)
            pdf.cell(90, 8, "Matched Skills:", 1, 0, fill=True)
            pdf.cell(90, 8, str(matched_skills_count), 1, 1)
            pdf.cell(90, 8, "Partially Matched Skills:", 1, 0, fill=True)
            pdf.cell(90, 8, str(match_counts['partial']), 1, 1)
            pdf.cell(90, 8, "Missing Skills:", 1, 0, fill=True)
            pdf.cell(90, 8, str(missing_skills_count), 1, 1)
            pdf.ln(10)
            
            # Match Rate Visualization
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "Match Rate Visualization", ln=True)
            pdf.set_font("Arial", '', 10)
            bar_width = 170
            bar_height = 15
            x_start = 20
            y_start = pdf.get_y()
            pdf.set_fill_color(220, 220, 220)
            pdf.rect(x_start, y_start, bar_width, bar_height, 'F')
            filled_width = (overall_match / 100) * bar_width
            if overall_match >= 70:
                pdf.set_fill_color(40, 167, 69)
            elif overall_match >= 50:
                pdf.set_fill_color(253, 126, 20)
            else:
                pdf.set_fill_color(220, 53, 69)
            pdf.rect(x_start, y_start, filled_width, bar_height, 'F')
            pdf.set_draw_color(100, 100, 100)
            pdf.rect(x_start, y_start, bar_width, bar_height, 'D')
            pdf.set_xy(x_start + bar_width/2 - 10, y_start + 3)
            pdf.set_font("Arial", 'B', 11)
            if overall_match > 30:
                pdf.set_text_color(255, 255, 255)
            else:
                pdf.set_text_color(0, 0, 0)
            pdf.cell(20, 8, f"{overall_match}%", align='C')
            pdf.set_text_color(0, 0, 0)
            pdf.ln(20)
            
            # Skills Distribution Chart
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "Skills Distribution", ln=True)
            pdf.ln(2)
            matched = match_counts['matched']
            partial = match_counts['partial']
            missing = match_counts['missing']
            total_skills = matched + partial + missing
            if total_skills > 0:
                chart_width = 170
                chart_height = 40
                x_chart = 20
                y_chart = pdf.get_y()
                matched_width = (matched / total_skills) * chart_width
                partial_width = (partial / total_skills) * chart_width
                missing_width = (missing / total_skills) * chart_width
                if matched > 0:
                    pdf.set_fill_color(40, 167, 69)
                    pdf.rect(x_chart, y_chart, matched_width, chart_height, 'F')
                if partial > 0:
                    pdf.set_fill_color(253, 126, 20)
                    pdf.rect(x_chart + matched_width, y_chart, partial_width, chart_height, 'F')
                if missing > 0:
                    pdf.set_fill_color(220, 53, 69)
                    pdf.rect(x_chart + matched_width + partial_width, y_chart, missing_width, chart_height, 'F')
                pdf.set_draw_color(100, 100, 100)
                pdf.rect(x_chart, y_chart, chart_width, chart_height, 'D')
                pdf.ln(chart_height + 5)
                pdf.set_font("Arial", '', 9)
                legend_x = 20
                legend_y = pdf.get_y()
                pdf.set_fill_color(40, 167, 69)
                pdf.rect(legend_x, legend_y, 5, 5, 'F')
                pdf.set_xy(legend_x + 7, legend_y - 1)
                pdf.cell(40, 6, f"Matched ({matched})")
                pdf.set_fill_color(253, 126, 20)
                pdf.rect(legend_x + 50, legend_y, 5, 5, 'F')
                pdf.set_xy(legend_x + 57, legend_y - 1)
                pdf.cell(40, 6, f"Partial ({partial})")
                pdf.set_fill_color(220, 53, 69)
                pdf.rect(legend_x + 100, legend_y, 5, 5, 'F')
                pdf.set_xy(legend_x + 107, legend_y - 1)
                pdf.cell(40, 6, f"Missing ({missing})")
                pdf.ln(12)
            
            # Technical Skills Section
            tech_matched = [s for s in skill_match_result['matched'] if s in jd_skills["technical"]]
            tech_partial = [s for s in skill_match_result['partial'] if s in jd_skills["technical"]]
            tech_missing = [s for s in skill_match_result['missing'] if s in jd_skills["technical"]]
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(40, 167, 69)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Technical Skills - Matched", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if tech_matched:
                for skill in tech_matched:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No technical skills matched.", 0, 1)
            pdf.ln(3)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(253, 126, 20)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Technical Skills - Partially Matched", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if tech_partial:
                for skill in tech_partial:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No technical skills partially matched.", 0, 1)
            pdf.ln(3)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(220, 53, 69)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Technical Skills - Missing (Priority)", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if tech_missing:
                for skill in tech_missing:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No missing technical skills!", 0, 1)
            pdf.ln(5)
            
            # Soft Skills Section
            soft_matched = [s for s in skill_match_result['matched'] if s in jd_skills["soft"]]
            soft_partial = [s for s in skill_match_result['partial'] if s in jd_skills["soft"]]
            soft_missing = [s for s in skill_match_result['missing'] if s in jd_skills["soft"]]
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(40, 167, 69)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Soft Skills - Matched", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if soft_matched:
                for skill in soft_matched:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No soft skills matched.", 0, 1)
            pdf.ln(3)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(253, 126, 20)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Soft Skills - Partially Matched", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if soft_partial:
                for skill in soft_partial:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No soft skills partially matched.", 0, 1)
            pdf.ln(3)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(220, 53, 69)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Soft Skills - Missing", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            if soft_missing:
                for skill in soft_missing:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, clean_skill, 0, 1)
            else:
                pdf.cell(0, 6, "No missing soft skills!", 0, 1)
            pdf.ln(5)
            
            # Recommendations
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "Recommendations", 0, 1)
            pdf.set_font("Arial", '', 10)
            recommendation = f"Your current match rate is {overall_match}%. "
            if overall_match >= 70:
                recommendation += "You have a strong skill match for this position! Continue to maintain and update these skills."
            elif overall_match >= 50:
                recommendation += "You have a moderate skill match. Focus on developing the missing skills and strengthening partially matched skills to improve your candidacy."
            else:
                recommendation += "There is significant room for improvement. Prioritize learning the missing skills listed above to increase your match rate to 70% or above."
            pdf.multi_cell(0, 6, recommendation)
            pdf.ln(3)
            
            # Action Items
            pdf.set_font("Arial", 'B', 10)
            pdf.cell(0, 6, "Action Items:", ln=True)
            pdf.set_font("Arial", '', 10)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "1. Focus on acquiring the missing skills through courses and certifications", ln=True)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "2. Strengthen partially matched skills through practical projects", ln=True)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "3. Update your resume once you acquire new skills", ln=True)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "4. Target positions with 70%+ match rate for better success", ln=True)
            pdf.ln(5)
            
            # Footer
            pdf.set_font("Arial", 'I', 9)
            pdf.cell(0, 10, f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 0, 'C')
            
            return pdf.output(dest="S").encode("latin-1")

        try:
            pdf_bytes = generate_pdf_report()
            st.download_button(label="⬇Download PDF Report", data=pdf_bytes, file_name="skill_gap_analysis_report.pdf", mime="application/pdf", help="Download formatted PDF report with visualizations")
        except Exception as e:
            st.error(f"Error generating PDF: {e}")

    st.success("Completed: Dashboard loaded successfully with export options.")
    
    st.markdown("---")
    st.markdown("""
    ### Analysis Complete!
    
    **Next Steps:**
    1. Review the missing skills identified above
    2. Consider upskilling in those areas
    3. Update your resume with newly acquired skills
    4. Download the report for your records
    
    **Pro Tip:** Aim for at least 70% match rate for better job prospects!
    
    **Multi-page Support:** This tool fully supports multi-page resumes and job descriptions in PDF, DOCX, and TXT formats.
    """)

else:
    st.info("Please upload both Resume and Job Description files, then click 'Analyze Documents' to begin.")
    st.markdown("""
    ### Supported File Formats:
    - **PDF** (.pdf) - Multi-page supported
    - **Word Document** (.docx) - Multi-page supported
    - **Text File** (.txt) - Multi-page supported
    
    ### Tips:
    - Multi-page resumes are fully supported - upload documents of any length
    - Ensure your files are properly formatted
    - Include clear skill listings in both documents
    - Experience information should be mentioned as "X years" or "X-Y years"
    - The tool will automatically detect and process all pages
    - **Click the "Analyze Documents" button after uploading both files**
    """)