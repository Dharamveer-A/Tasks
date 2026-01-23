import streamlit as st
import pdfplumber
import docx
import re
import spacy
from spacy.matcher import PhraseMatcher
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
import io

# Milestone - 1

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="Skill Gap Analysis Dashboard",
    layout="wide"
)

# ---------------- HEADER ----------------
st.markdown(
    """
    <div style="background-color:#4a5bdc;padding:20px;border-radius:10px">
        <h2 style="color:white;">
            Milestone 1: Data Ingestion, Parsing and Cleaning Module 
        </h2>
        <p style="color:white;">
            This module allows uploading resumes and job descriptions,
            extracting text from multiple document formats, and displaying
            clean, normalized content for further processing.
        </p>
    </div>
    """,
    unsafe_allow_html=True
)

st.write("")

# ---------------- INITIALIZE ----------------
resume_text = None
jd_text = None
cleaned_resume = ""
cleaned_jd = ""
resume_experience = None
jd_experience = None

# ---------------- FUNCTIONS ----------------
def extract_text(file):
    """Extract text from PDF, DOCX, or TXT files"""
    try:
        if file.type == "application/pdf":
            text = ""
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    if page.extract_text():
                        text += page.extract_text() + "\n"
            return text

        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Reset file pointer to beginning
            file.seek(0)
            document = docx.Document(file)
            return "\n".join([para.text for para in document.paragraphs])

        elif file.type == "text/plain":
            # Reset file pointer to beginning
            file.seek(0)
            return file.read().decode("utf-8")
        
        return ""
    except Exception as e:
        st.error(f"Error extracting text from {file.name}: {str(e)}")
        st.info("💡 Try re-uploading the file or use a different format (PDF/TXT)")
        return ""

def clean(text):
    """Clean and normalize text for NLP processing"""
    if not text:
        return ""

    text = text.lower()
    # Preserve experience patterns like "3-5 years" and "5+ years"
    text = re.sub(r'(\d+)\s*(?:-|–|to)\s*(\d+)', r'\1_\2', text)
    text = re.sub(r'(\d+)\s*\+', r'\1_plus', text)
    # Remove newlines and tabs
    text = re.sub(r'\n|\t', ' ', text)
    # Remove special characters but keep alphanumeric and underscores
    text = re.sub(r'[^a-z0-9_\s]', ' ', text)
    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text).strip()

    return text

def extract_experience(text):
    """Extract years of experience from text"""
    if not text:
        return None

    text = text.lower()

    # Experience range: 1-3 years, 2 to 5 yrs
    range_match = re.search(
        r'(\d+)\s*(?:-|–|to)\s*(\d+)\s*(?:years?|yrs?|year experience|years experience)',
        text
    )
    if range_match:
        return {
            "min_exp": int(range_match.group(1)),
            "max_exp": int(range_match.group(2))
        }

    # Experience plus: 3+ years
    plus_match = re.search(
        r'(\d+)\s*\+\s*(?:years?|yrs?|year experience|years experience)',
        text
    )
    if plus_match:
        return {
            "min_exp": int(plus_match.group(1)),
            "max_exp": None
        }

    return None


# ---------------- UPLOAD ----------------
st.subheader("Upload Documents")

left_col, right_col = st.columns(2)

with left_col:
    resume_file = st.file_uploader(
        "Upload Resume",
        type=["pdf", "docx", "txt"],
        help="Upload your resume in PDF, DOCX, or TXT format"
    )

with right_col:
    jd_file = st.file_uploader(
        "Upload Job Description",
        type=["pdf", "docx", "txt"],
        help="Upload job description in PDF, DOCX, or TXT format"
    )

if jd_file and resume_file:

    # ---------------- PREVIEW ----------------
    st.subheader("Parsed Document Preview")

    col1, col2 = st.columns(2)

    with col1:
        if resume_file:
            resume_text = extract_text(resume_file)
            st.text_area("Resume Preview", resume_text, height=300, key="resume_preview")

    with col2:
        if jd_file:
            jd_text = extract_text(jd_file)
            st.text_area("Job Description Preview", jd_text, height=300, key="jd_preview")

    # ---------------- PROCESS ----------------
    if resume_text:
        cleaned_resume = clean(resume_text)
        resume_experience = extract_experience(resume_text)

    if jd_text:
        cleaned_jd = clean(jd_text)
        jd_experience = extract_experience(jd_text)

    # ---------------- CLEANED OUTPUT ----------------
    st.subheader("Cleaned Files")

    col1, col2 = st.columns(2)

    with col1:
        if cleaned_resume:
            st.text_area("Cleaned Resume", cleaned_resume, height=200, key="cleaned_resume")
    with col2:
        if cleaned_jd:
            st.text_area("Cleaned Job Description", cleaned_jd, height=200, key="cleaned_jd")

    st.success(
        "✅ Milestone 1 Completed: Documents uploaded, parsed, cleaned and previewed successfully."
    )

    # -------------------------------------------------------------------------
    # Milestone - 2

    st.markdown(
        """
        <div style="background-color:#4a5bdc;padding:20px;border-radius:10px;margin-top:30px;">
            <h2 style="color:white;">
                Milestone 2: Skill Extraction using NLP Module
            </h2>
            <p style="color:white;">
                Module: Skill Extraction using NLP <br>
                • spaCy and BERT-based pipelines  <br>
                • Technical and soft skills identification <br>
                • Structured skill display
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.write("")

    st.subheader("Experience Analysis")

    col1, col2 = st.columns(2)
    with col1:
        if resume_experience:
            st.info(f"Resume Experience: {resume_experience['min_exp']}" + 
                   (f"-{resume_experience['max_exp']}" if resume_experience['max_exp'] else "+") + " years")
        else:
            st.warning("No experience information found in resume")
    with col2:
        if jd_experience:
            st.info(f"JD Required Experience: {jd_experience['min_exp']}" + 
                   (f"-{jd_experience['max_exp']}" if jd_experience['max_exp'] else "+") + " years")
        else:
            st.warning("No experience requirement found in JD")

    # Load spaCy model with error handling
    @st.cache_resource
    def load_spacy_model():
        try:
            return spacy.load("en_core_web_trf")
        except OSError:
            st.error("spaCy model 'en_core_web_trf' not found. Please install it using: python -m spacy download en_core_web_trf")
            return None
    
    nlp = load_spacy_model()
    
    if nlp is None:
        st.stop()

    master_technical_skills = {
        # Programming Languages
        "python", "java", "c", "c++", "c#", "go", "rust", "kotlin", "swift",
        "javascript", "typescript", "php", "ruby", "r", "matlab", "scala",
        "perl", "bash", "shell scripting", "sql",

        # Web Technologies
        "html", "css", "sass", "bootstrap", "tailwind css",
        "react", "angular", "vue.js", "next.js", "nuxt.js",
        "node.js", "express.js", "django", "flask", "fastapi",
        "spring", "spring boot", "laravel", "asp.net",

        # Databases
        "mysql", "postgresql", "sqlite", "oracle", "sql server",
        "mongodb", "cassandra", "dynamodb", "redis", "firebase",
        "elasticsearch", "neo4j",

        # Data Science & Analytics
        "numpy", "pandas", "scipy", "matplotlib", "seaborn",
        "plotly", "power bi", "tableau", "excel", "statistics",
        "data analysis", "data visualization",

        # Machine Learning & AI
        "machine learning", "deep learning", "artificial intelligence",
        "tensorflow", "keras", "pytorch", "scikit-learn", "xgboost",
        "opencv", "nlp", "computer vision", "speech recognition",
        "transformers", "huggingface", "langchain", "llm",

        # Big Data
        "hadoop", "spark", "pyspark", "kafka", "hive", "pig",
        "hbase", "flink", "airflow", "databricks",

        # Cloud & DevOps
        "aws", "azure", "google cloud", "gcp",
        "docker", "kubernetes", "terraform", "ansible",
        "jenkins", "github actions", "ci/cd",
        "linux", "unix",

        # APIs & Backend
        "rest api", "graphql", "grpc", "soap",
        "microservices", "jwt", "oauth",

        # Testing
        "unit testing", "integration testing", "system testing",
        "pytest", "unittest", "junit", "selenium", "cypress",
        "postman",

        # Mobile Development
        "android", "ios", "flutter", "react native",
        "swiftui", "kotlin multiplatform",

        # Cybersecurity
        "network security", "application security", "penetration testing",
        "ethical hacking", "cryptography", "owasp",
        "siem", "firewall",

        # Operating Systems
        "windows", "linux", "macos",

        # Version Control & Tools
        "git", "github", "gitlab", "bitbucket",
        "jira", "confluence", "trello",

        # Other Technical Skills
        "data structures", "algorithms", "object oriented programming",
        "design patterns", "system design",
        "software development lifecycle", "agile", "scrum"
    }

    master_soft_skills = {
        # Communication
        "communication", "verbal communication", "written communication",
        "public speaking", "presentation skills", "active listening",
        "business communication", "storytelling",

        # Interpersonal Skills
        "teamwork", "collaboration", "interpersonal skills",
        "relationship building", "empathy", "emotional intelligence",
        "conflict resolution", "negotiation",

        # Leadership
        "leadership", "people management", "team leadership",
        "decision making", "delegation", "mentoring", "coaching",
        "influencing", "strategic thinking",

        # Problem Solving & Thinking
        "problem solving", "critical thinking", "analytical thinking",
        "logical reasoning", "creative thinking", "innovation",
        "root cause analysis", "troubleshooting",

        # Time & Work Management
        "time management", "prioritization", "multitasking",
        "work ethic", "self discipline", "accountability",
        "goal setting", "organizational skills",

        # Adaptability & Learning
        "adaptability", "flexibility", "resilience",
        "learning agility", "continuous learning",
        "open mindedness", "growth mindset",

        # Professionalism
        "professionalism", "integrity", "ethical behavior",
        "reliability", "punctuality", "confidentiality",

        # Creativity & Innovation
        "creativity", "idea generation", "design thinking",
        "innovation mindset", "curiosity",

        # Emotional & Personal Skills
        "stress management", "self awareness", "self motivation",
        "confidence", "positive attitude", "emotional control",

        # Customer & Service Orientation
        "customer focus", "customer service",
        "client management", "stakeholder management",
        "user empathy", "service mindset",

        # Collaboration & Culture
        "cross functional collaboration", "cultural awareness",
        "diversity and inclusion", "remote collaboration",
        "team alignment",

        # Conflict & Crisis Handling
        "conflict management", "crisis management",
        "handling pressure", "decision making under pressure",

        # Work Style
        "independent work", "collaborative work",
        "attention to detail", "quality focus",
        "result oriented", "ownership",

        # Ethics & Values
        "honesty", "trustworthiness", "respect",
        "fairness", "social responsibility"
    }


    def extract_soft_skills(text, skill_set):
        """Extract soft skills using simple text matching"""
        if not text:
            return []

        text = text.lower()
        found_skills = set()

        for skill in skill_set:
            if skill in text:
                found_skills.add(skill)

        return sorted(found_skills)

    def extract_technical_skills(text, skill_set):
        """Extract technical skills using spaCy PhraseMatcher"""
        if not text:
            return []

        matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
        patterns = [nlp.make_doc(skill) for skill in skill_set]
        matcher.add("TECH_SKILLS", patterns)

        doc = nlp(text)
        matches = matcher(doc)

        skills_found = set()
        for _, start, end in matches:
            skills_found.add(doc[start:end].text.lower())

        return sorted(skills_found)

    # Extract skills
    with st.spinner("Extracting skills using NLP..."):
        resume_skills = {
            "technical": extract_technical_skills(cleaned_resume, master_technical_skills),
            "soft": extract_soft_skills(cleaned_resume, master_soft_skills)
        }

        jd_skills = {
            "technical": extract_technical_skills(cleaned_jd, master_technical_skills),
            "soft": extract_soft_skills(cleaned_jd, master_soft_skills)
        }

    st.subheader("Extracted Skills")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### Resume Skills")
        with st.expander("Technical Skills", expanded=True):
            if resume_skills["technical"]:
                st.write(", ".join(resume_skills["technical"]))
            else:
                st.info("No technical skills found")
        with st.expander("Soft Skills"):
            if resume_skills["soft"]:
                st.write(", ".join(resume_skills["soft"]))
            else:
                st.info("No soft skills found")

    with col2:
        st.markdown("### JD Skills")
        with st.expander("Technical Skills", expanded=True):
            if jd_skills["technical"]:
                st.write(", ".join(jd_skills["technical"]))
            else:
                st.info("No technical skills found")
        with st.expander("Soft Skills"):
            if jd_skills["soft"]:
                st.write(", ".join(jd_skills["soft"]))
            else:
                st.info("No soft skills found")

    def donut_chart(tech_count, soft_count, title):
        """Create donut chart for skill distribution"""
        labels = ["Technical Skills", "Soft Skills"]
        sizes = [tech_count, soft_count]
        if sum(sizes) == 0:
            sizes = [1, 1]

        fig = plt.figure(figsize=(4, 4))
        ax = fig.add_axes([0.15, 0.15, 0.7, 0.7]) 

        wedges, _, _ = ax.pie(
            sizes,
            startangle=90,
            autopct="%1.0f%%",
            radius=1,
            wedgeprops=dict(width=0.4, edgecolor="white"),
            colors=['#4a5bdc', '#fd7e14']
        )

        ax.set(aspect="equal")
        ax.set_title(title, pad=10, fontsize=12, fontweight='bold')
        ax.legend(
            wedges,
            labels,
            loc="lower center",
            bbox_to_anchor=(0.5, -0.25),
            ncol=2,
            frameon=False
        )

        return fig

    st.subheader("Skill Distribution Analysis")

    col1, col2 = st.columns(2)

    with col1:
        fig1 = donut_chart(
            len(resume_skills["technical"]),
            len(resume_skills["soft"]),
            "Resume Skill Distribution"
        )
        st.pyplot(fig1, use_container_width=False)
        plt.close(fig1)

    with col2:
        fig2 = donut_chart(
            len(jd_skills["technical"]),
            len(jd_skills["soft"]),
            "JD Skill Distribution"
        )
        st.pyplot(fig2, use_container_width=False)
        plt.close(fig2)

    def compute_metrics(resume_skills, jd_skills):
        """Compute skill matching metrics"""
        resume_tech = set(resume_skills["technical"])
        resume_soft = set(resume_skills["soft"])
        jd_tech = set(jd_skills["technical"])
        jd_soft = set(jd_skills["soft"])

        matched_skills = (resume_tech & jd_tech) | (resume_soft & jd_soft)
        total_jd_skills = len(jd_tech) + len(jd_soft)

        match_percentage = 0
        if total_jd_skills > 0:
            match_percentage = round((len(matched_skills) / total_jd_skills) * 100, 1)

        return {
            "resume_tech": len(resume_tech),
            "resume_soft": len(resume_soft),
            "jd_tech": len(jd_tech),
            "jd_soft": len(jd_soft),
            "total_resume": len(resume_tech) + len(resume_soft),
            "total_jd": total_jd_skills,
            "match_percent": match_percentage
        }
    
    metrics = compute_metrics(resume_skills, jd_skills)

    st.subheader("Skill Extraction Metrics")

    col1, col2, col3, col4 = st.columns(4)

    col1.metric(
        label="JD Technical Skills",
        value=metrics["jd_tech"]
    )

    col2.metric(
        label="JD Soft Skills",
        value=metrics["jd_soft"]
    )

    col3.metric(
        label="Total JD Skills",
        value=metrics["total_jd"]
    )

    col4.metric(
        label="Basic Match %",
        value=f'{metrics["match_percent"]}%'
    )

    st.success("✅ Milestone 2 Completed: Skills extracted successfully using NLP.")

    # --------------------------------------------------------------------------
    # Milestone - 3

    def build_similarity_matrix(resume_skills, jd_skills):
        """Build similarity matrix between resume and JD skills"""
        if not resume_skills or not jd_skills:
            return np.array([])
        
        matrix = np.zeros((len(resume_skills), len(jd_skills)))

        for i, r_skill in enumerate(resume_skills):
            for j, j_skill in enumerate(jd_skills):
                if r_skill == j_skill:
                    matrix[i][j] = 1.0
                elif r_skill in j_skill or j_skill in r_skill:
                    matrix[i][j] = 0.5
                else:
                    matrix[i][j] = 0.0

        return matrix

    def plot_similarity_heatmap(matrix, resume_labels, jd_labels):
        """Plot skill similarity heatmap"""
        fig, ax = plt.subplots(figsize=(10, 6))

        im = ax.imshow(matrix, cmap="Purples", aspect='auto')

        ax.set_xticks(range(len(jd_labels)))
        ax.set_yticks(range(len(resume_labels)))

        ax.set_xticklabels(jd_labels, rotation=45, ha="right", fontsize=8)
        ax.set_yticklabels(resume_labels, fontsize=8)

        ax.set_xlabel("Job Description Skills", fontweight='bold')
        ax.set_ylabel("Resume Skills", fontweight='bold')
        ax.set_title("Skill Similarity Matrix", fontsize=14, fontweight='bold', pad=15)

        cbar = fig.colorbar(im, ax=ax)
        cbar.set_label("Similarity Score", fontweight='bold')

        plt.tight_layout()
        return fig

    resume_all_skills = resume_skills["technical"] + resume_skills["soft"]
    jd_all_skills = jd_skills["technical"] + jd_skills["soft"]

    st.markdown(
        """
        <div style="background-color:#4a5bdc;padding:20px;border-radius:10px;margin-top:30px;">
            <h2 style="color:white;">
                Milestone 3: Skill Gap Analysis and Similarity Matching Module
            </h2>
            <p style="color:white;">
                • Skill similarity matrix visualization <br>
                • Resume vs JD skill comparison <br>
                • Missing skill identification
            </p>
        </div>
        <br>
        """,
        unsafe_allow_html=True
    )

    if resume_all_skills and jd_all_skills:
        similarity_matrix = build_similarity_matrix(
            resume_all_skills,
            jd_all_skills
        )

        fig_heatmap = plot_similarity_heatmap(
            similarity_matrix,
            resume_all_skills,
            jd_all_skills
        )
        st.pyplot(fig_heatmap)
        plt.close(fig_heatmap)
    else:
        st.warning("⚠️ Insufficient skills to build similarity matrix")

    def classify_skill_matches(similarity_matrix, resume_skills, jd_skills):
        """Classify skills as matched, partial, or missing"""
        matched = set()
        partial = set()

        for j, jd_skill in enumerate(jd_skills):
            column = similarity_matrix[:, j]

            if 1.0 in column:
                matched.add(jd_skill)
            elif 0.5 in column:
                partial.add(jd_skill)

        missing = set(jd_skills) - matched - partial

        return {
            "matched": sorted(list(matched)),
            "partial": sorted(list(partial)),
            "missing": sorted(list(missing))
        }

    skill_match_result = classify_skill_matches(
        similarity_matrix,
        resume_all_skills,
        jd_all_skills
    )

    def calculate_skill_match(resume_skills, jd_skills):
        """Calculate detailed skill match statistics"""
        resume_set = set(resume_skills)
        jd_set = set(jd_skills)

        matched = resume_set & jd_set
        missing = jd_set - resume_set

        # Partial = same base word (simple heuristic)
        partial = set()
        for jd_skill in missing.copy():
            for res_skill in resume_set:
                if jd_skill in res_skill or res_skill in jd_skill:
                    partial.add(jd_skill)

        missing = missing - partial

        avg_match = 0
        if len(jd_set) > 0:
            avg_match = round(
                ((len(matched) + 0.5 * len(partial)) / len(jd_set)) * 100,
                1
            )

        return {
            "matched": len(matched),
            "partial": len(partial),
            "missing": len(missing),
            "avg_match": avg_match
        }

    all_resume_skills = (
        set(resume_skills["technical"]) |
        set(resume_skills["soft"])
    )

    all_jd_skills = (
        set(jd_skills["technical"]) |
        set(jd_skills["soft"])
    )

    match_counts = calculate_skill_match(
        all_resume_skills,
        all_jd_skills
    )

    st.subheader("Skill Match Overview")

    left, right = st.columns([1.3, 1])

    # ---------------- LEFT: DONUT ----------------
    with left:
        labels = ["Matched", "Partially Matched", "Missing"]
        sizes = [match_counts["matched"], match_counts["partial"], match_counts["missing"]]
        colors = ["#28a745", "#fd7e14", "#dc3545"]

        if sum(sizes) == 0:
            sizes = [1, 1, 1]

        fig = plt.figure(figsize=(3.6, 3.6))
        ax = fig.add_axes([0.1, 0.15, 0.8, 0.75])

        wedges, _, autotexts = ax.pie(
            sizes,
            startangle=90,
            autopct="%1.0f%%",
            radius=1,
            colors=colors,
            wedgeprops=dict(width=0.4, edgecolor="white")
        )

        # Make percentage text black and bold
        for autotext in autotexts:
            autotext.set_color('black')
            autotext.set_fontweight('bold')
            autotext.set_fontsize(11)

        ax.set(aspect="equal")
        ax.set_title("Skill Match Distribution", pad=8, fontsize=12, fontweight='bold')
        ax.legend(
            wedges,
            labels,
            loc="lower center",
            bbox_to_anchor=(0.5, -0.25),
            ncol=3,
            frameon=False
        )

        st.pyplot(fig, use_container_width=False)
        plt.close(fig)

    # ---------------- RIGHT: METRICS ----------------
    with right:
        with st.container():
            r1, r2 = st.columns(2)
            r1.metric("✅ Matched Skills", match_counts["matched"])
            r2.metric("⚠️ Partially Matched", match_counts["partial"])

            r3, r4 = st.columns(2)
            r3.metric("❌ Missing Skills", match_counts["missing"])
            r4.metric("📊 Avg Match %", f'{match_counts["avg_match"]}%')

    st.subheader("Skill Gap Details")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("### ✅ Matched Skills")
        if skill_match_result["matched"]:
            for skill in skill_match_result["matched"]:
                st.success(f"✓ {skill}")
        else:
            st.info("No perfectly matched skills")

    with col2:
        st.markdown("### ⚠️ Partial Matches")
        if skill_match_result["partial"]:
            for skill in skill_match_result["partial"]:
                st.warning(f"≈ {skill}")
        else:
            st.info("No partially matched skills")

    with col3:
        st.markdown("### ❌ Missing Skills")
        if skill_match_result["missing"]:
            for skill in skill_match_result["missing"]:
                st.error(f"✗ {skill}")
        else:
            st.success("No missing skills!")

    st.success("✅ Milestone 3 Completed: Skill gap analysis completed successfully.")

    # -----------------------------------------------------------------------------
    # Milestone - 4

    st.markdown("""
    <div style="background-color:#4a90e2;padding:15px;border-radius:8px;margin-top:30px;">
    <h2 style="color:white;text-align:center;">
    Milestone 4: Dashboard and Report Export Module
    </h2>
    <p style="color:white;text-align:center;">
    Interactive dashboard • Graphs • Multi-format report export
    </p>
    </div>
    """, unsafe_allow_html=True)

    st.write("")

    # ---------------- DATA ----------------
    skills = list(all_jd_skills)
    resume_scores = [
        100 if skill in all_resume_skills else 0
        for skill in skills
    ]

    job_scores = [100 for _ in skills]

    overall_match = match_counts["avg_match"]
    matched_skills_count = match_counts["matched"]
    missing_skills_count = match_counts["missing"]

    # ---------------- METRICS ----------------
    c1, c2, c3 = st.columns(3)
    c1.metric("📊 Overall Match", f"{overall_match}%")
    c2.metric("✅ Matched Skills", str(matched_skills_count))
    c3.metric("❌ Missing Skills", str(missing_skills_count))

    # ---------------- BAR CHART ----------------
    bar_fig = go.Figure()
    bar_fig.add_bar(x=skills, y=resume_scores, name="Resume Skills", marker_color='#4a5bdc')
    bar_fig.add_bar(x=skills, y=job_scores, name="Job Requirements", marker_color='#28a745')
    bar_fig.update_layout(
        title="Skill Match Overview",
        barmode="group",
        height=400,
        xaxis_title="Skills",
        yaxis_title="Score (%)",
        hovermode='x unified'
    )

    # ---------------- RADAR CHART ----------------
    radar_skills = skills[:5] if len(skills) >= 5 else skills
    radar_resume = resume_scores[:5] if len(resume_scores) >= 5 else resume_scores
    radar_job = job_scores[:5] if len(job_scores) >= 5 else job_scores
    
    radar_fig = go.Figure()
    radar_fig.add_trace(go.Scatterpolar(
        r=radar_resume,
        theta=radar_skills,
        fill='toself',
        name='Current Profile',
        line_color='#4a5bdc'
    ))
    radar_fig.add_trace(go.Scatterpolar(
        r=radar_job,
        theta=radar_skills,
        fill='toself',
        name='Job Requirements',
        line_color='#28a745'
    ))
    radar_fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
        title="Top Skills Comparison",
        height=400,
        showlegend=True
    )

    left, right = st.columns([2, 1])
    left.plotly_chart(bar_fig, use_container_width=True)
    right.plotly_chart(radar_fig, use_container_width=True)

    # ---------------- SKILL PROGRESS ----------------
    st.subheader("📈 Sample Skill Proficiency")
    st.caption("Note: These are sample scores for demonstration. For actual proficiency, use skill assessment tools.")
    
    sample_skills = [("Python", 92), ("Machine Learning", 88), ("SQL", 65)]
    for skill, score in sample_skills:
        col_skill, col_progress = st.columns([1, 4])
        with col_skill:
            st.write(f"**{skill}**")
        with col_progress:
            st.progress(score / 100)
            st.caption(f"{score}%")

    # ---------------- AREA CHART ----------------
    st.subheader("📊 Skill Similarity Score Distribution")
    
    if resume_all_skills and jd_all_skills and len(similarity_matrix) > 0:
        # Calculate average similarity score for each JD skill
        avg_similarity_scores = []
        for j in range(len(jd_all_skills)):
            column_scores = similarity_matrix[:, j]
            avg_score = np.mean(column_scores) * 100  # Convert to percentage
            avg_similarity_scores.append(avg_score)
        
        # Create area chart data
        area_df = pd.DataFrame({
            'Skill': jd_all_skills,
            'Similarity Score (%)': avg_similarity_scores
        })
        
        # Create area chart using plotly
        area_fig = go.Figure()
        area_fig.add_trace(go.Scatter(
            x=area_df['Skill'],
            y=area_df['Similarity Score (%)'],
            fill='tozeroy',
            name='Similarity Score',
            line_color='#7c3aed',
            fillcolor='rgba(124, 58, 237, 0.3)'
        ))
        
        area_fig.update_layout(
            title="Average Similarity Score per Job Requirement",
            xaxis_title="Skills",
            yaxis_title="Similarity Score (%)",
            height=400,
            hovermode='x unified',
            yaxis=dict(range=[0, 100])
        )
        
        st.plotly_chart(area_fig, use_container_width=True)
    else:
        st.info("Upload documents to see similarity score distribution")

    # ---------------- UPSKILLING ----------------
    st.subheader("💡 Upskilling Recommendations")
    st.caption("Based on missing and partially matched skills from job description")
    
    # Show missing skills
    if skill_match_result["missing"]:
        st.markdown("**🔴 Priority Skills to Learn (Missing):**")
        for i, skill in enumerate(skill_match_result["missing"][:5], 1):
            st.error(f"{i}. **{skill.title()}** - Not found in your resume")
    
    # Show partially matched skills
    if skill_match_result["partial"]:
        st.markdown("**🟡 Skills to Strengthen (Partial Match):**")
        for i, skill in enumerate(skill_match_result["partial"][:5], 1):
            st.warning(f"{i}. **{skill.title()}** - Improve proficiency in this area")
    
    if not skill_match_result["missing"] and not skill_match_result["partial"]:
        st.success("🎉 Excellent! You have all the required skills for this job!")

    # ---------------- REPORT CONTENT ----------------
    report_text = f"""
SKILL GAP ANALYSIS REPORT
{'='*50}

SUMMARY
-------
Overall Match: {overall_match}%
Matched Skills: {matched_skills_count}
Partially Matched: {match_counts['partial']}
Missing Skills: {missing_skills_count}
Total JD Requirements: {len(all_jd_skills)}

MATCHED SKILLS
--------------
{chr(10).join(skill_match_result['matched']) if skill_match_result['matched'] else 'None'}

PARTIALLY MATCHED SKILLS
------------------------
{chr(10).join(skill_match_result['partial']) if skill_match_result['partial'] else 'None'}

MISSING SKILLS
--------------
{chr(10).join(skill_match_result['missing']) if skill_match_result['missing'] else 'None'}

DETAILED SKILL BREAKDOWN
------------------------
"""

    for i in range(len(skills)):
        report_text += f"""
{skills[i]}
  Resume Score: {resume_scores[i]}%
  Job Requirement: {job_scores[i]}%
  Status: {'✓ Matched' if resume_scores[i] == 100 else '✗ Missing'}
"""

    report_text += f"""

RECOMMENDATIONS
---------------
Based on this analysis, focus on developing the missing skills to improve your match rate.
Current match rate: {overall_match}%
Target match rate: 80%+

Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}
"""

    st.subheader("📥 Report Download")

    col1, col2, col3 = st.columns(3)

    with col1:
        # ---------------- CSV ----------------
        # Create three separate lists for the three columns
        matched_list = skill_match_result['matched'] if skill_match_result['matched'] else []
        partial_list = skill_match_result['partial'] if skill_match_result['partial'] else []
        missing_list = skill_match_result['missing'] if skill_match_result['missing'] else []
        
        # Pad lists to same length
        max_len = max(len(matched_list), len(partial_list), len(missing_list), 1)
        matched_list += [''] * (max_len - len(matched_list))
        partial_list += [''] * (max_len - len(partial_list))
        missing_list += [''] * (max_len - len(missing_list))
        
        df_export = pd.DataFrame({
            "Matched Skills": matched_list,
            "Partially Matched Skills": partial_list,
            "Missing Skills": missing_list
        })

        st.download_button(
            "⬇️ CSV Report",
            df_export.to_csv(index=False).encode("utf-8"),
            "skill_gap_report.csv",
            "text/csv",
            help="Download skill comparison in 3 columns: Matched, Partially Matched, Missing"
        )

    with col2:
        # ---------------- DOCX WITH VISUALIZATIONS ----------------
        def generate_word_report():
            doc = Document()
            
            # Title
            title = doc.add_heading('Skill Gap Analysis Report', 0)
            title.alignment = 1  # Center
            
            # Executive Summary
            doc.add_heading('Executive Summary', 1)
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Light Grid Accent 1'
            
            summary_data = [
                ('Overall Match Rate', f'{overall_match}%'),
                ('Matched Skills', str(matched_skills_count)),
                ('Partially Matched Skills', str(match_counts["partial"])),
                ('Missing Skills', str(missing_skills_count))
            ]
            
            for i, (label, value) in enumerate(summary_data):
                summary_table.rows[i].cells[0].text = label
                summary_table.rows[i].cells[1].text = value
            
            doc.add_paragraph()
            
            # Skills Breakdown
            doc.add_heading('Skills Breakdown', 1)
            
            # Matched Skills
            doc.add_heading('Matched Skills', 2)
            if skill_match_result['matched']:
                for skill in skill_match_result['matched']:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No perfectly matched skills found.')
            
            # Partially Matched Skills
            doc.add_heading('Partially Matched Skills', 2)
            if skill_match_result['partial']:
                for skill in skill_match_result['partial']:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No partially matched skills found.')
            
            # Missing Skills
            doc.add_heading('Missing Skills', 2)
            if skill_match_result['missing']:
                for skill in skill_match_result['missing']:
                    doc.add_paragraph(skill, style='List Bullet')
            else:
                doc.add_paragraph('No missing skills - Great job!')
            
            # Recommendations
            doc.add_heading('Recommendations', 1)
            doc.add_paragraph(
                f'Your current match rate is {overall_match}%. '
                f'To improve your candidacy, consider focusing on the missing skills listed above.'
            )
            
            if overall_match >= 70:
                doc.add_paragraph('You have a strong skill match for this position!', style='List Bullet')
            else:
                doc.add_paragraph('Focus on upskilling in the missing areas to increase your match rate.', style='List Bullet')
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph(f'Generated on: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
            footer.alignment = 1
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        doc_file = generate_word_report()

        st.download_button(
            "⬇️ DOCX Report",
            doc_file,
            "skill_gap_analysis_report.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Download formatted Word report with complete analysis"
        )
    
    with col3:

        def generate_pdf_report():
            pdf = FPDF()
            pdf.add_page()
            
            # Title
            pdf.set_font("Arial", 'B', 20)
            pdf.cell(0, 15, "Skill Gap Analysis Report", ln=True, align='C')
            pdf.ln(5)
            
            # Executive Summary
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, "Executive Summary", ln=True)
            pdf.set_font("Arial", '', 11)
            pdf.ln(2)
            
            # Summary box
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(90, 8, "Overall Match Rate:", 1, 0, fill=True)
            pdf.cell(90, 8, f"{overall_match}%", 1, 1)
            pdf.cell(90, 8, "Matched Skills:", 1, 0, fill=True)
            pdf.cell(90, 8, str(matched_skills_count), 1, 1)
            pdf.cell(90, 8, "Partially Matched Skills:", 1, 0, fill=True)
            pdf.cell(90, 8, str(match_counts['partial']), 1, 1)
            pdf.cell(90, 8, "Missing Skills:", 1, 0, fill=True)
            pdf.cell(90, 8, str(missing_skills_count), 1, 1)
            pdf.ln(10)
            
            # Visual Progress Bar for Match Rate
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "Match Rate Visualization", ln=True)
            pdf.set_font("Arial", '', 10)
            
            # Draw progress bar
            bar_width = 170
            bar_height = 15
            x_start = 20
            y_start = pdf.get_y()
            
            # Background (gray)
            pdf.set_fill_color(220, 220, 220)
            pdf.rect(x_start, y_start, bar_width, bar_height, 'F')
            
            # Filled portion (green/yellow/red based on score)
            filled_width = (overall_match / 100) * bar_width
            if overall_match >= 70:
                pdf.set_fill_color(40, 167, 69)  # Green
            elif overall_match >= 50:
                pdf.set_fill_color(253, 126, 20)  # Orange
            else:
                pdf.set_fill_color(220, 53, 69)  # Red
            pdf.rect(x_start, y_start, filled_width, bar_height, 'F')
            
            # Border
            pdf.set_draw_color(100, 100, 100)
            pdf.rect(x_start, y_start, bar_width, bar_height, 'D')
            
            # Percentage text
            pdf.set_xy(x_start + bar_width/2 - 10, y_start + 3)
            pdf.set_font("Arial", 'B', 11)
            pdf.set_text_color(255, 255, 255) if overall_match > 30 else pdf.set_text_color(0, 0, 0)
            pdf.cell(20, 8, f"{overall_match}%", align='C')
            pdf.set_text_color(0, 0, 0)
            pdf.ln(20)
            
            # Skills Distribution Chart
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "Skills Distribution", ln=True)
            pdf.ln(2)
            
            # Draw bar chart
            matched = match_counts['matched']
            partial = match_counts['partial']
            missing = match_counts['missing']
            total_skills = matched + partial + missing
            
            if total_skills > 0:
                chart_width = 170
                chart_height = 40
                x_chart = 20
                y_chart = pdf.get_y()
                
                # Calculate widths
                matched_width = (matched / total_skills) * chart_width
                partial_width = (partial / total_skills) * chart_width
                missing_width = (missing / total_skills) * chart_width
                
                # Draw matched (green)
                if matched > 0:
                    pdf.set_fill_color(40, 167, 69)
                    pdf.rect(x_chart, y_chart, matched_width, chart_height, 'F')
                
                # Draw partial (orange)
                if partial > 0:
                    pdf.set_fill_color(253, 126, 20)
                    pdf.rect(x_chart + matched_width, y_chart, partial_width, chart_height, 'F')
                
                # Draw missing (red)
                if missing > 0:
                    pdf.set_fill_color(220, 53, 69)
                    pdf.rect(x_chart + matched_width + partial_width, y_chart, missing_width, chart_height, 'F')
                
                # Border
                pdf.set_draw_color(100, 100, 100)
                pdf.rect(x_chart, y_chart, chart_width, chart_height, 'D')
                
                pdf.ln(chart_height + 5)
                
                # Legend
                pdf.set_font("Arial", '', 9)
                legend_x = 20
                legend_y = pdf.get_y()
                
                # Matched legend
                pdf.set_fill_color(40, 167, 69)
                pdf.rect(legend_x, legend_y, 5, 5, 'F')
                pdf.set_xy(legend_x + 7, legend_y - 1)
                pdf.cell(40, 6, f"Matched ({matched})")
                
                # Partial legend
                pdf.set_fill_color(253, 126, 20)
                pdf.rect(legend_x + 50, legend_y, 5, 5, 'F')
                pdf.set_xy(legend_x + 57, legend_y - 1)
                pdf.cell(40, 6, f"Partial ({partial})")
                
                # Missing legend
                pdf.set_fill_color(220, 53, 69)
                pdf.rect(legend_x + 100, legend_y, 5, 5, 'F')
                pdf.set_xy(legend_x + 107, legend_y - 1)
                pdf.cell(40, 6, f"Missing ({missing})")
                
                pdf.ln(12)
            
            # Matched Skills
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(40, 167, 69)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Matched Skills", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            
            if skill_match_result['matched']:
                for skill in skill_match_result['matched']:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, f"{clean_skill}", 0, 1)
            else:
                pdf.cell(0, 6, "No perfectly matched skills found.", 0, 1)
            pdf.ln(3)
            
            # Partially Matched Skills
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(253, 126, 20)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Partially Matched Skills", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            
            if skill_match_result['partial']:
                for skill in skill_match_result['partial']:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, f"{clean_skill}", 0, 1)
            else:
                pdf.cell(0, 6, "No partially matched skills found.", 0, 1)
            pdf.ln(3)
            
            # Missing Skills
            pdf.set_font("Arial", 'B', 12)
            pdf.set_fill_color(220, 53, 69)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(0, 8, "Missing Skills - Priority for Development", 0, 1, fill=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", '', 10)
            
            if skill_match_result['missing']:
                for skill in skill_match_result['missing']:
                    clean_skill = skill.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(5, 6, '', 0, 0)
                    pdf.cell(0, 6, f"{clean_skill}", 0, 1)
            else:
                pdf.cell(0, 6, "No missing skills - Great job!", 0, 1)
            pdf.ln(5)
            
            # Recommendations
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "Recommendations", 0, 1)
            pdf.set_font("Arial", '', 10)
            
            recommendation = f"Your current match rate is {overall_match}%. "
            if overall_match >= 70:
                recommendation += "You have a strong skill match for this position! Continue to maintain and update these skills."
            elif overall_match >= 50:
                recommendation += "You have a moderate skill match. Focus on developing the missing skills and strengthening partially matched skills to improve your candidacy."
            else:
                recommendation += "There is significant room for improvement. Prioritize learning the missing skills listed above to increase your match rate to 70% or above."
            
            pdf.multi_cell(0, 6, recommendation)
            pdf.ln(3)
            
            # Action Items
            pdf.set_font("Arial", 'B', 10)
            pdf.cell(0, 6, "Action Items:", ln=True)
            pdf.set_font("Arial", '', 10)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "1. Focus on acquiring the missing skills through courses and certifications", ln=True)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "2. Strengthen partially matched skills through practical projects", ln=True)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "3. Update your resume once you acquire new skills", ln=True)
            pdf.cell(5, 6, '', 0, 0)
            pdf.cell(0, 6, "4. Target positions with 70%+ match rate for better success", ln=True)
            pdf.ln(5)
            
            # Footer
            pdf.set_font("Arial", 'I', 9)
            pdf.cell(0, 10, f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 0, 'C')
            
            return pdf.output(dest="S").encode("latin-1")


        try:
            pdf_bytes = generate_pdf_report()

            _ = st.download_button(
                label="⬇️ PDF Report",
                data=pdf_bytes,
                file_name="skill_gap_analysis_report.pdf",
                mime="application/pdf",
                help="Download formatted PDF report with visualizations"
            )

        except Exception as e:
            st.error(f"Error generating PDF: {e}")


    st.success("✅ Milestone 4 Completed: Dashboard loaded successfully with export options.")
    
    # Final Summary
    st.markdown("---")
    st.markdown("""
    ### 🎯 Analysis Complete!
    
    **Next Steps:**
    1. Review the missing skills identified above
    2. Consider upskilling in those areas
    3. Update your resume with newly acquired skills
    4. Download the report for your records
    
    **Pro Tip:** Aim for at least 70% match rate for better job prospects!
    """)

else:
    st.info("👆 Please upload both Resume and Job Description files to begin the analysis.")
    st.markdown("""
    ### 📋 Supported File Formats:
    - **PDF** (.pdf)
    - **Word Document** (.docx)
    - **Text File** (.txt)
    
    ### 💡 Tips:
    - Ensure your files are properly formatted
    - Include clear skill listings in both documents
    - Experience information should be mentioned as "X years" or "X-Y years"
    """)