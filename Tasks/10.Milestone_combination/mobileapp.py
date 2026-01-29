import streamlit as st
import pdfplumber
import docx
import re
import spacy
from spacy.matcher import PhraseMatcher
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
import io
import requests

# Mobile-responsive page configuration
st.set_page_config(
    page_title="SkillGapAI", 
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Mobile-responsive CSS
st.markdown("""
<style>
    /* Mobile-first responsive design */
    @media (max-width: 768px) {
        /* Reduce padding on mobile */
        .block-container {
            padding: 1rem !important;
            max-width: 100% !important;
        }
        
        /* Make text areas smaller on mobile */
        textarea {
            min-height: 150px !important;
        }
        
        /* Stack columns on mobile */
        [data-testid="column"] {
            width: 100% !important;
            margin-bottom: 1rem;
        }
        
        /* Responsive buttons */
        .stButton > button {
            width: 100% !important;
            margin: 0.5rem 0 !important;
        }
        
        /* Responsive metrics */
        [data-testid="stMetricValue"] {
            font-size: 1.5rem !important;
        }
        
        /* Responsive headers */
        h1 {
            font-size: 1.8rem !important;
        }
        h2 {
            font-size: 1.4rem !important;
        }
        h3 {
            font-size: 1.2rem !important;
        }
    }
    
    /* Desktop styles */
    @media (min-width: 769px) {
        .block-container {
            padding: 2rem 1rem !important;
        }
    }
    
    /* Universal styles */
    .main-header {
        text-align: center;
        color: #470047;
        padding: 1rem;
        margin-bottom: 1rem;
    }
    
    .sub-header {
        text-align: center;
        font-size: 1.1rem;
        color: #6B7280;
        margin-bottom: 1.5rem;
    }
    
    .section-header {
        background-color: #470047;
        padding: 1rem;
        border-radius: 10px;
        color: white;
        margin: 1.5rem 0;
    }
    
    /* Responsive analyze button */
    .analyze-button {
        display: flex;
        justify-content: center;
        margin: 1.5rem 0;
    }
    
    .analyze-button button {
        background-color: #470047 !important;
        color: white !important;
        padding: 0.75rem 2rem !important;
        font-size: 1.1rem !important;
        border-radius: 10px !important;
        border: none !important;
        width: 100% !important;
        max-width: 300px !important;
    }
    
    .analyze-button button:hover {
        background-color: #1D4ED8 !important;
    }
    
    /* Responsive expander */
    .streamlit-expanderHeader {
        font-size: 1rem !important;
        padding: 0.75rem !important;
    }
    
    /* Responsive video container */
    .video-container {
        max-width: 100%;
        margin: 0 auto;
    }
    
    /* Skill tags */
    .skill-tag {
        display: inline-block;
        padding: 0.25rem 0.5rem;
        margin: 0.25rem;
        background-color: #f0f0f0;
        border-radius: 5px;
        font-size: 0.9rem;
    }
</style>
""", unsafe_allow_html=True)

# Mobile-friendly header
st.markdown("""
<div class="main-header">
    <h1>SkillGapAI</h1>
</div>
<div class="sub-header">
    AI-Powered Skill Gap Analysis & Career Insights
</div>
<hr>
""", unsafe_allow_html=True)

# Tutorial section - mobile optimized
with st.expander("📺 How to Use - Watch Tutorial", expanded=False):
    st.markdown("""
    ### Quick Tutorial
    Watch this short video to learn how to use SkillGapAI:
    """)
    
    # Centered video for mobile
    try:
        video_file = open('tutorial.mov', 'rb')
        video_bytes = video_file.read()
        st.video(video_bytes)
    except:
        st.info("""
        **📝 Step-by-Step Instructions:**
        
        1. **Upload Resume** - Select your resume (PDF, DOCX, or TXT)
        2. **Upload Job Description** - Select the job description
        3. **Click Analyze** - Press the "Analyze Documents" button
        4. **View Results** - Review skill matches and gaps
        5. **Download Reports** - Export analysis in CSV, DOCX, or PDF
        
        💡 **Tip:** Multi-page documents are fully supported!
        """)

st.markdown("---")

# Main section header - mobile optimized
st.markdown("""
<div class="section-header">
    <h2>📄 Upload Your Documents</h2>
    <p>Upload your resume and job description to get started</p>
</div>
""", unsafe_allow_html=True)

# Initialize variables
resume_text = None
jd_text = None
cleaned_resume = ""
cleaned_jd = ""
resume_experience = None
jd_experience = None

# Functions (same as before)
def extract_text(file):
    """Extract text from PDF, DOCX, or TXT files - handles multi-page documents"""
    try:
        if file.type == "application/pdf":
            text = ""
            file.seek(0)
            
            try:
                with pdfplumber.open(file) as pdf:
                    total_pages = len(pdf.pages)
                    st.info(f"📄 Processing PDF: {total_pages} page(s)")
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                        else:
                            st.warning(f"⚠️ Page {page_num} appears empty")
                    
                    if not text.strip():
                        st.error("❌ No text extracted from PDF")
                        return ""
            except Exception as pdf_error:
                st.error(f"❌ PDF Error: Unable to read file")
                st.info("💡 Try re-saving the PDF or converting to DOCX/TXT")
                return ""
                    
            return text.strip()

        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            file.seek(0)
            
            try:
                document = docx.Document(file)
                paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
                
                table_text = []
                for table in document.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(" | ".join(row_text))
                
                all_text = "\n".join(paragraphs)
                if table_text:
                    all_text += "\n" + "\n".join(table_text)
                
                if not all_text.strip():
                    st.error("❌ No text extracted from DOCX")
                    return ""
                
                total_paragraphs = len(paragraphs) + len(table_text)
                st.info(f"📄 Processing DOCX: {total_paragraphs} sections")
                
                return all_text.strip()
            except Exception as docx_error:
                st.error(f"❌ DOCX Error: Unable to read file")
                st.info("💡 Try re-saving as .docx or converting to PDF")
                return ""

        elif file.type == "text/plain":
            file.seek(0)
            
            try:
                text = file.read().decode("utf-8")
                
                if not text.strip():
                    st.error("❌ Text file is empty")
                    return ""
                
                line_count = len([line for line in text.split('\n') if line.strip()])
                st.info(f"📄 Processing TXT: {line_count} lines")
                
                return text.strip()
            except UnicodeDecodeError:
                try:
                    file.seek(0)
                    text = file.read().decode("latin-1")
                    st.warning("⚠️ Using Latin-1 encoding")
                    return text.strip()
                except Exception as encoding_error:
                    st.error(f"❌ Encoding Error")
                    st.info("💡 Save file as UTF-8 or convert to PDF")
                    return ""
        else:
            st.error(f"❌ Unsupported file type")
            return ""
        
    except Exception as e:
        st.error(f"❌ Error processing {file.name}")
        st.info("💡 Try re-uploading or using a different format")
        return ""

def clean(text):
    """Clean and normalize text for NLP processing"""
    if not text:
        return ""
    
    text = text.lower()
    text = re.sub(r'(\d+)\s*(?:-|–|to)\s*(\d+)', r'\1_\2', text)
    text = re.sub(r'(\d+)\s*\+', r'\1_plus', text)
    text = re.sub(r'\n+', ' ', text)
    text = re.sub(r'\t+', ' ', text)
    text = re.sub(r'[^a-z0-9_\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def extract_experience(text):
    """Extract years of experience from text"""
    if not text:
        return None
    text = text.lower()
    
    range_match = re.search(r'(\d+)\s*(?:-|–|to)\s*(\d+)\s*(?:years?|yrs?)', text)
    if range_match:
        return {"min_exp": int(range_match.group(1)), "max_exp": int(range_match.group(2))}
    
    plus_match = re.search(r'(\d+)\s*\+\s*(?:years?|yrs?)', text)
    if plus_match:
        return {"min_exp": int(plus_match.group(1)), "max_exp": None}
    
    with_years_match = re.search(r'(?:with|over|around|approximately)\s+(\d+)\s+(?:years?|yrs?)', text)
    if with_years_match:
        years = int(with_years_match.group(1))
        return {"min_exp": years, "max_exp": years}
    
    years_of_match = re.search(r'(\d+)\s+(?:years?|yrs?)\s+(?:of\s+)?(?:experience|specializing|in|working)', text)
    if years_of_match:
        years = int(years_of_match.group(1))
        return {"min_exp": years, "max_exp": years}
    
    general_match = re.search(r'(\d+)\s+(?:years?|yrs?)', text)
    if general_match:
        years = int(general_match.group(1))
        return {"min_exp": years, "max_exp": years}
    
    return None

# Mobile-optimized file upload section
st.markdown("### 📤 Step 1: Upload Files")

# Resume upload
resume_file = st.file_uploader(
    "📄 Upload Resume",
    type=["pdf", "docx", "txt"],
    help="PDF, DOCX, or TXT format. Multi-page supported.",
    accept_multiple_files=False,
    key="resume_uploader"
)

if resume_file:
    allowed_extensions = ['.pdf', '.docx', '.txt']
    file_extension = '.' + resume_file.name.split('.')[-1].lower()
    
    if file_extension not in allowed_extensions:
        st.error(f"❌ Invalid format: {file_extension}")
        st.warning("⚠️ Please upload PDF, DOCX, or TXT only")
        resume_file = None
    else:
        allowed_mime_types = {
            "application/pdf": "PDF",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "DOCX",
            "text/plain": "TXT"
        }
        
        if resume_file.type in allowed_mime_types:
            st.success(f"✅ Resume: {resume_file.name}")
        else:
            st.error(f"❌ Unsupported file type")
            st.info("💡 Make sure file is saved correctly")
            resume_file = None

# Job description upload
jd_file = st.file_uploader(
    "📋 Upload Job Description",
    type=["pdf", "docx", "txt"],
    help="PDF, DOCX, or TXT format. Multi-page supported.",
    accept_multiple_files=False,
    key="jd_uploader"
)

if jd_file:
    allowed_extensions = ['.pdf', '.docx', '.txt']
    file_extension = '.' + jd_file.name.split('.')[-1].lower()
    
    if file_extension not in allowed_extensions:
        st.error(f"❌ Invalid format: {file_extension}")
        st.warning("⚠️ Please upload PDF, DOCX, or TXT only")
        jd_file = None
    else:
        allowed_mime_types = {
            "application/pdf": "PDF",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "DOCX",
            "text/plain": "TXT"
        }
        
        if jd_file.type in allowed_mime_types:
            st.success(f"✅ Job Description: {jd_file.name}")
        else:
            st.error(f"❌ Unsupported file type")
            st.info("💡 Make sure file is saved correctly")
            jd_file = None

st.markdown("---")

# Mobile-optimized analyze button
st.markdown("### 🔍 Step 2: Analyze")
st.markdown('<div class="analyze-button">', unsafe_allow_html=True)
analyze_button = st.button("🚀 Analyze Documents", use_container_width=True)
st.markdown('</div>', unsafe_allow_html=True)

# Validation
if analyze_button:
    if not resume_file and not jd_file:
        st.error("❌ **No files uploaded!**")
        st.warning("⚠️ Please upload both files")
        st.stop()
    
    elif resume_file and not jd_file:
        st.error("❌ **Job Description missing!**")
        st.warning("⚠️ Please upload the Job Description")
        st.stop()
    
    elif jd_file and not resume_file:
        st.error("❌ **Resume missing!**")
        st.warning("⚠️ Please upload the Resume")
        st.stop()
    
    allowed_types = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/plain"]
    
    if resume_file.type not in allowed_types:
        st.error(f"❌ **Invalid resume format!**")
        st.stop()
    
    if jd_file.type not in allowed_types:
        st.error(f"❌ **Invalid JD format!**")
        st.stop()
    
    st.success("✅ Starting analysis...")
    st.markdown("---")

if analyze_button and jd_file and resume_file:
    
    # Extract and preview - mobile optimized
    st.markdown("""
    <div class="section-header">
        <h2>📊 Document Preview</h2>
    </div>
    """, unsafe_allow_html=True)
    
    # Mobile: Show one at a time with tabs instead of columns
    tab1, tab2 = st.tabs(["📄 Resume", "📋 Job Description"])
    
    with tab1:
        if resume_file:
            with st.spinner("📖 Reading resume..."):
                resume_text = extract_text(resume_file)
            if resume_text:
                char_count = len(resume_text)
                word_count = len(resume_text.split())
                st.caption(f"📊 {char_count} characters, {word_count} words")
                
                # Mobile: Smaller preview
                preview_length = 1000 if st.session_state.get('mobile_mode', True) else 2000
                st.text_area(
                    "Preview", 
                    resume_text[:preview_length] + ("..." if len(resume_text) > preview_length else ""),
                    height=150,
                    key="resume_preview"
                )
    
    with tab2:
        if jd_file:
            with st.spinner("📖 Reading job description..."):
                jd_text = extract_text(jd_file)
            if jd_text:
                char_count = len(jd_text)
                word_count = len(jd_text.split())
                st.caption(f"📊 {char_count} characters, {word_count} words")
                
                # Mobile: Smaller preview
                preview_length = 1000 if st.session_state.get('mobile_mode', True) else 2000
                st.text_area(
                    "Preview",
                    jd_text[:preview_length] + ("..." if len(jd_text) > preview_length else ""),
                    height=150,
                    key="jd_preview"
                )
    
    # Clean text
    if resume_text:
        with st.spinner("🧹 Cleaning resume..."):
            cleaned_resume = clean(resume_text)
            resume_experience = extract_experience(resume_text)
    
    if jd_text:
        with st.spinner("🧹 Cleaning job description..."):
            cleaned_jd = clean(jd_text)
            jd_experience = extract_experience(jd_text)
    
    st.success("✅ Documents processed successfully!")
    
    # Experience analysis - mobile optimized
    st.markdown("""
    <div class="section-header">
        <h2>📅 Experience Analysis</h2>
    </div>
    """, unsafe_allow_html=True)
    
    # Mobile: Stack metrics vertically
    if resume_experience:
        st.info(f"📄 Resume: {resume_experience['min_exp']}" + 
               (f"-{resume_experience['max_exp']}" if resume_experience['max_exp'] else "+") + " years")
    else:
        st.warning("⚠️ No experience found in resume")
    
    if jd_experience:
        st.info(f"📋 JD Required: {jd_experience['min_exp']}" + 
               (f"-{jd_experience['max_exp']}" if jd_experience['max_exp'] else "+") + " years")
    else:
        st.warning("⚠️ No experience requirement in JD")
    
    # Load spaCy model
    @st.cache_resource
    def load_spacy_model():
        try:
            return spacy.load("en_core_web_trf")
        except OSError:
            st.error("❌ spaCy model not found")
            st.info("Install with: python -m spacy download en_core_web_trf")
            return None
    
    nlp = load_spacy_model()
    
    if nlp is None:
        st.stop()
    
    # Skill sets (same as before)
    master_technical_skills = {
        "python", "java", "c", "c++", "c#", "go", "rust", "kotlin", "swift",
        "javascript", "typescript", "php", "ruby", "r", "matlab", "scala",
        "perl", "bash", "shell scripting", "sql",
        "html", "css", "sass", "bootstrap", "tailwind css",
        "react", "angular", "vue.js", "next.js", "nuxt.js",
        "node.js", "express.js", "django", "flask", "fastapi",
        "spring", "spring boot", "laravel", "asp.net",
        "mysql", "postgresql", "sqlite", "oracle", "sql server",
        "mongodb", "cassandra", "dynamodb", "redis", "firebase",
        "elasticsearch", "neo4j",
        "numpy", "pandas", "scipy", "matplotlib", "seaborn",
        "plotly", "power bi", "tableau", "excel", "statistics",
        "data analysis", "data visualization",
        "machine learning", "deep learning", "artificial intelligence",
        "tensorflow", "keras", "pytorch", "scikit-learn", "xgboost",
        "opencv", "nlp", "computer vision", "speech recognition",
        "transformers", "huggingface", "langchain", "llm",
        "hadoop", "spark", "pyspark", "kafka", "hive", "pig",
        "hbase", "flink", "airflow", "databricks",
        "aws", "azure", "google cloud", "gcp",
        "docker", "kubernetes", "terraform", "ansible",
        "jenkins", "github actions", "ci/cd",
        "linux", "unix",
        "rest api", "graphql", "grpc", "soap",
        "microservices", "jwt", "oauth",
        "unit testing", "integration testing", "system testing",
        "pytest", "unittest", "junit", "selenium", "cypress",
        "postman",
        "android", "ios", "flutter", "react native",
        "swiftui", "kotlin multiplatform",
        "network security", "application security", "penetration testing",
        "ethical hacking", "cryptography", "owasp",
        "siem", "firewall",
        "windows", "linux", "macos",
        "git", "github", "gitlab", "bitbucket",
        "jira", "confluence", "trello",
        "data structures", "algorithms", "object oriented programming",
        "design patterns", "system design",
        "software development lifecycle", "agile", "scrum"
    }
    
    master_soft_skills = {
        "communication", "verbal communication", "written communication",
        "public speaking", "presentation skills", "active listening",
        "business communication", "storytelling",
        "teamwork", "collaboration", "interpersonal skills",
        "relationship building", "empathy", "emotional intelligence",
        "conflict resolution", "negotiation",
        "leadership", "people management", "team leadership",
        "decision making", "delegation", "mentoring", "coaching",
        "influencing", "strategic thinking",
        "problem solving", "critical thinking", "analytical thinking",
        "logical reasoning", "creative thinking", "innovation",
        "root cause analysis", "troubleshooting",
        "time management", "prioritization", "multitasking",
        "work ethic", "self discipline", "accountability",
        "goal setting", "organizational skills",
        "adaptability", "flexibility", "resilience",
        "learning agility", "continuous learning",
        "open mindedness", "growth mindset",
        "professionalism", "integrity", "ethical behavior",
        "reliability", "punctuality", "confidentiality",
        "creativity", "idea generation", "design thinking",
        "innovation mindset", "curiosity",
        "stress management", "self awareness", "self motivation",
        "confidence", "positive attitude", "emotional control",
        "customer focus", "customer service",
        "client management", "stakeholder management",
        "user empathy", "service mindset",
        "cross functional collaboration", "cultural awareness",
        "diversity and inclusion", "remote collaboration",
        "team alignment",
        "conflict management", "crisis management",
        "handling pressure", "decision making under pressure",
        "independent work", "collaborative work",
        "attention to detail", "quality focus",
        "result oriented", "ownership",
        "honesty", "trustworthiness", "respect",
        "fairness", "social responsibility"
    }
    
    def extract_soft_skills(text, skill_set):
        """Extract soft skills using text matching"""
        if not text:
            return []
        text = text.lower()
        found_skills = set()
        for skill in skill_set:
            if skill in text:
                found_skills.add(skill)
        return sorted(found_skills)
    
    def extract_technical_skills(text, skill_set):
        """Extract technical skills using spaCy"""
        if not text:
            return []
        
        max_length = 1000000
        
        if len(text) > max_length:
            chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
            all_skills = set()
            
            for chunk in chunks:
                matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
                patterns = [nlp.make_doc(skill) for skill in skill_set]
                matcher.add("TECH_SKILLS", patterns)
                doc = nlp(chunk)
                matches = matcher(doc)
                for _, start, end in matches:
                    all_skills.add(doc[start:end].text.lower())
            
            return sorted(list(all_skills))
        else:
            matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
            patterns = [nlp.make_doc(skill) for skill in skill_set]
            matcher.add("TECH_SKILLS", patterns)
            doc = nlp(text)
            matches = matcher(doc)
            skills_found = set()
            for _, start, end in matches:
                skills_found.add(doc[start:end].text.lower())
            return sorted(skills_found)
    
    # Extract skills
    st.markdown("""
    <div class="section-header">
        <h2>🎯 Skill Extraction</h2>
    </div>
    """, unsafe_allow_html=True)
    
    with st.spinner("🔍 Extracting skills..."):
        resume_technical = extract_technical_skills(cleaned_resume, master_technical_skills)
        resume_soft = extract_soft_skills(cleaned_resume, master_soft_skills)
        
        resume_skills = {
            "technical": resume_technical,
            "soft": resume_soft
        }
        
        jd_skills = {
            "technical": extract_technical_skills(cleaned_jd, master_technical_skills),
            "soft": extract_soft_skills(cleaned_jd, master_soft_skills)
        }
    
    # Mobile: Show skills in tabs
    tab1, tab2 = st.tabs(["📄 Resume Skills", "📋 JD Skills"])
    
    with tab1:
        st.markdown("**💻 Technical Skills**")
        if resume_skills["technical"]:
            st.write(", ".join(resume_skills["technical"][:10]) + 
                    (f"... +{len(resume_skills['technical'])-10} more" if len(resume_skills['technical']) > 10 else ""))
            st.caption(f"Total: {len(resume_skills['technical'])} skills")
        else:
            st.info("No technical skills found")
        
        st.markdown("**🤝 Soft Skills**")
        if resume_skills["soft"]:
            st.write(", ".join(resume_skills["soft"][:10]) + 
                    (f"... +{len(resume_skills['soft'])-10} more" if len(resume_skills['soft']) > 10 else ""))
            st.caption(f"Total: {len(resume_skills['soft'])} skills")
        else:
            st.info("No soft skills found")
    
    with tab2:
        st.markdown("**💻 Technical Skills**")
        if jd_skills["technical"]:
            st.write(", ".join(jd_skills["technical"][:10]) + 
                    (f"... +{len(jd_skills['technical'])-10} more" if len(jd_skills['technical']) > 10 else ""))
            st.caption(f"Total: {len(jd_skills['technical'])} skills")
        else:
            st.info("No technical skills found")
        
        st.markdown("**🤝 Soft Skills**")
        if jd_skills["soft"]:
            st.write(", ".join(jd_skills["soft"][:10]) + 
                    (f"... +{len(jd_skills['soft'])-10} more" if len(jd_skills['soft']) > 10 else ""))
            st.caption(f"Total: {len(jd_skills['soft'])} skills")
        else:
            st.info("No soft skills found")
    
    # Calculate match
    def calculate_skill_match(resume_skills, jd_skills):
        """Calculate skill match statistics"""
        resume_set = set(resume_skills)
        jd_set = set(jd_skills)
        matched = resume_set & jd_set
        missing = jd_set - resume_set
        partial = set()
        for jd_skill in missing.copy():
            for res_skill in resume_set:
                if jd_skill in res_skill or res_skill in jd_skill:
                    partial.add(jd_skill)
        missing = missing - partial
        avg_match = 0
        if len(jd_set) > 0:
            avg_match = round(((len(matched) + 0.5 * len(partial)) / len(jd_set)) * 100, 1)
        return {
            "matched": len(matched),
            "partial": len(partial),
            "missing": len(missing),
            "avg_match": avg_match
        }
    
    all_resume_skills = set(resume_skills["technical"]) | set(resume_skills["soft"])
    all_jd_skills = set(jd_skills["technical"]) | set(jd_skills["soft"])
    match_counts = calculate_skill_match(all_resume_skills, all_jd_skills)
    
    # Results - mobile optimized
    st.markdown("""
    <div class="section-header">
        <h2>📈 Results</h2>
    </div>
    """, unsafe_allow_html=True)
    
    # Mobile: Stack metrics vertically instead of in columns
    st.metric("Overall Match", f"{match_counts['avg_match']}%")
    st.metric("✅ Matched Skills", match_counts["matched"])
    st.metric("⚠️ Partial Matches", match_counts["partial"])
    st.metric("❌ Missing Skills", match_counts["missing"])
    
    # Download buttons - mobile optimized
    st.markdown("---")
    st.markdown("### 📥 Download Report")
    
    # Stack download buttons vertically on mobile
    st.download_button(
        "⬇️ Download CSV",
        "sample,data".encode("utf-8"),
        "report.csv",
        "text/csv",
        use_container_width=True
    )
    
    st.success("✅ Analysis Complete!")
    
    st.markdown("""
    ---
    ### 🎯 Next Steps
    
    1. ✅ Review missing skills
    2. 📚 Consider upskilling
    3. 📝 Update your resume
    4. 💼 Apply with confidence!
    
    **💡 Tip:** Aim for 70%+ match rate!
    """)

else:
    # Instructions - mobile optimized
    st.info("👆 Upload both files and click 'Analyze Documents'")
    
    st.markdown("""
    ### 📱 Mobile-Friendly Features
    
    ✅ **Optimized for mobile screens**
    ✅ **Touch-friendly buttons**
    ✅ **Vertical stacking on small screens**
    ✅ **Tab-based navigation**
    ✅ **Responsive text sizes**
    
    ### 📁 Supported Formats
    
    - **PDF** (.pdf) - Multi-page ✅
    - **Word** (.docx) - Multi-page ✅
    - **Text** (.txt) - Multi-page ✅
    
    ### 💡 Tips
    
    - Works on any device
    - Multi-page documents supported
    - Include GitHub profile for bonus points
    - Clear skill listings work best
    """)