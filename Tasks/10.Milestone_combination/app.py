import streamlit as st
import pdfplumber
import docx
import re
import spacy
from spacy.matcher import PhraseMatcher
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
import io

# Milestone - 1

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="Skill Gap Analysis Dashboard",
    layout="wide"
)

# ---------------- HEADER ----------------
st.markdown(
    """
    <div style="background-color:#4a5bdc;padding:20px;border-radius:10px">
        <h2 style="color:white;">
            Milestone 1: Data Ingestion, Parsing and Cleaning Module 
        </h2>
        <p style="color:white;">
            This module allows uploading resumes and job descriptions,
            extracting text from multiple document formats, and displaying
            clean, normalized content for further processing.
        </p>
    </div>
    """,
    unsafe_allow_html=True
)

st.write("")

# ---------------- INITIALIZE ----------------
resume_text = None
jd_text = None
cleaned_resume = ""
cleaned_jd = ""
resume_experience = None
jd_experience = None

# ---------------- FUNCTIONS ----------------
def extract_text(file):
    if file.type == "application/pdf":
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                if page.extract_text():
                    text += page.extract_text() + "\n"
        return text

    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        document = docx.Document(file)
        return "\n".join([para.text for para in document.paragraphs])

    elif file.type == "text/plain":
        return file.read().decode("utf-8")

    return ""

def clean(text):
    if not text:
        return ""

    text = text.lower()
    text = re.sub(r'(\d+)\s*(?:-|–|to)\s*(\d+)', r'\1_\2', text)
    text = re.sub(r'(\d+)\s*\+', r'\1_plus', text)
    text = re.sub(r'\n|\t', ' ', text)
    text = re.sub(r'[^a-z0-9_\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()

    return text

def extract_experience(text):
    if not text:
        return None

    text = text.lower()

    # experience range: 1-3 years, 2 to 5 yrs
    range_match = re.search(
        r'(\d+)\s*(?:-|–|to)\s*(\d+)\s*(?:years?|yrs?|year experience|years experience)',
        text
    )
    if range_match:
        return {
            "min_exp": int(range_match.group(1)),
            "max_exp": int(range_match.group(2))
        }

    # experience plus: 3+ years
    plus_match = re.search(
        r'(\d+)\s*\+\s*(?:years?|yrs?|year experience|years experience)',
        text
    )
    if plus_match:
        return {
            "min_exp": int(plus_match.group(1)),
            "max_exp": None
        }

    return None


# ---------------- UPLOAD ----------------
st.subheader("Upload Documents")

left_col, right_col = st.columns(2)

with left_col:
    resume_file = st.file_uploader(
        "Upload Resume",
        type=["pdf", "docx", "txt"]
    )

with right_col:
    jd_file = st.file_uploader(
        "Upload Job Description",
        type=["pdf", "docx", "txt"]
    )

if jd_file and resume_file:

    # ---------------- PREVIEW ----------------
    st.subheader("Parsed Document Preview")

    col1, col2 = st.columns(2)

    with col1:
        if resume_file:
            resume_text = extract_text(resume_file)
            st.text_area("Resume Preview", resume_text, height=300)

    with col2:
        if jd_file:
            jd_text = extract_text(jd_file)
            st.text_area("Job Description Preview", jd_text, height=300)

    # ---------------- PROCESS ----------------
    if resume_text:
        cleaned_resume = clean(resume_text)
        resume_experience = extract_experience(resume_text)

    if jd_text:
        cleaned_jd = clean(jd_text)
        jd_experience = extract_experience(jd_text)

    # ---------------- CLEANED OUTPUT ----------------
    st.subheader("Cleaned Files")

    col1, col2 = st.columns(2)

    with col1:
        if cleaned_resume:
            st.text_area("Cleaned Resume", cleaned_resume, height=200)
    with col2:
        if cleaned_jd:
            st.text_area("Cleaned Job Description", cleaned_jd, height=200)

    st.success(
        "Milestone 1 Completed: Documents uploaded, parsed, cleaned and previewed successfully."
    )


    # Milestone - 2

    st.markdown(
        """
        <div style="background-color:#4a5bdc;padding:20px;border-radius:10px">
            <h2 style="color:white;">
                Milestone 2: Skill Extraction using NLP Module
            </h2>
            <p style="color:white;">
                Module: Skill Extraction using NLP <br>
                • spaCy and BERT-based pipelines  <br>
                • Technical and soft skills identification <br>
                • Structured skill display
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.subheader("Experience :")

    col1, col2 = st.columns(2)
    with col1:
        st.write("Resume Experience found:", resume_experience)
    with col2:
        st.write("JD Experience found:", jd_experience)

    nlp = spacy.load("en_core_web_trf")

    master_technical_skills = {
        # Programming Languages
        "python", "java", "c", "c++", "c#", "go", "rust", "kotlin", "swift",
        "javascript", "typescript", "php", "ruby", "r", "matlab", "scala",
        "perl", "bash", "shell scripting",

        # Web Technologies
        "html", "css", "sass", "bootstrap", "tailwind css",
        "react", "angular", "vue.js", "next.js", "nuxt.js",
        "node.js", "express.js", "django", "flask", "fastapi",
        "spring", "spring boot", "laravel", "asp.net",

        # Databases
        "mysql", "postgresql", "sqlite", "oracle", "sql server",
        "mongodb", "cassandra", "dynamodb", "redis", "firebase",
        "elasticsearch", "neo4j",

        # Data Science & Analytics
        "numpy", "pandas", "scipy", "matplotlib", "seaborn",
        "plotly", "power bi", "tableau", "excel", "statistics",
        "data analysis", "data visualization",

        # Machine Learning & AI
        "machine learning", "deep learning", "artificial intelligence",
        "tensorflow", "keras", "pytorch", "scikit-learn", "xgboost",
        "opencv", "nlp", "computer vision", "speech recognition",
        "transformers", "huggingface", "langchain", "llm",

        # Big Data
        "hadoop", "spark", "pyspark", "kafka", "hive", "pig",
        "hbase", "flink", "airflow", "databricks",

        # Cloud & DevOps
        "aws", "azure", "google cloud", "gcp",
        "docker", "kubernetes", "terraform", "ansible",
        "jenkins", "github actions", "ci/cd",
        "linux", "unix",

        # APIs & Backend
        "rest api", "graphql", "grpc", "soap",
        "microservices", "jwt", "oauth",

        # Testing
        "unit testing", "integration testing", "system testing",
        "pytest", "unittest", "junit", "selenium", "cypress",
        "postman",

        # Mobile Development
        "android", "ios", "flutter", "react native",
        "swiftui", "kotlin multiplatform",

        # Cybersecurity
        "network security", "application security", "penetration testing",
        "ethical hacking", "cryptography", "owasp",
        "siem", "firewall",

        # Operating Systems
        "windows", "linux", "macos",

        # Version Control & Tools
        "git", "github", "gitlab", "bitbucket",
        "jira", "confluence", "trello",

        # Other Technical Skills
        "data structures", "algorithms", "object oriented programming",
        "design patterns", "system design",
        "software development lifecycle", "agile", "scrum"
    }

    master_soft_skills = {
        # Communication
        "communication", "verbal communication", "written communication",
        "public speaking", "presentation skills", "active listening",
        "business communication", "storytelling",

        # Interpersonal Skills
        "teamwork", "collaboration", "interpersonal skills",
        "relationship building", "empathy", "emotional intelligence",
        "conflict resolution", "negotiation",

        # Leadership
        "leadership", "people management", "team leadership",
        "decision making", "delegation", "mentoring", "coaching",
        "influencing", "strategic thinking",

        # Problem Solving & Thinking
        "problem solving", "critical thinking", "analytical thinking",
        "logical reasoning", "creative thinking", "innovation",
        "root cause analysis", "troubleshooting",

        # Time & Work Management
        "time management", "prioritization", "multitasking",
        "work ethic", "self discipline", "accountability",
        "goal setting", "organizational skills",

        # Adaptability & Learning
        "adaptability", "flexibility", "resilience",
        "learning agility", "continuous learning",
        "open mindedness", "growth mindset",

        # Professionalism
        "professionalism", "integrity", "ethical behavior",
        "reliability", "punctuality", "confidentiality",

        # Creativity & Innovation
        "creativity", "idea generation", "design thinking",
        "innovation mindset", "curiosity",

        # Emotional & Personal Skills
        "stress management", "self awareness", "self motivation",
        "confidence", "positive attitude", "emotional control",

        # Customer & Service Orientation
        "customer focus", "customer service",
        "client management", "stakeholder management",
        "user empathy", "service mindset",

        # Collaboration & Culture
        "cross functional collaboration", "cultural awareness",
        "diversity and inclusion", "remote collaboration",
        "team alignment",

        # Conflict & Crisis Handling
        "conflict management", "crisis management",
        "handling pressure", "decision making under pressure",

        # Work Style
        "independent work", "collaborative work",
        "attention to detail", "quality focus",
        "result oriented", "ownership",

        # Ethics & Values
        "honesty", "trustworthiness", "respect",
        "fairness", "social responsibility"
    }


    def extract_soft_skills(text, skill_set):
        if not text:
            return []

        text = text.lower()
        found_skills = set()

        for skill in skill_set:
            if skill in text:
                found_skills.add(skill)

        return sorted(found_skills)

    def extract_technical_skills(text, skill_set):
        if not text:
            return []

        matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
        patterns = [nlp.make_doc(skill) for skill in skill_set]
        matcher.add("TECH_SKILLS", patterns)

        doc = nlp(text)
        matches = matcher(doc)

        skills_found = set()
        for _, start, end in matches:
            skills_found.add(doc[start:end].text.lower())

        return sorted(skills_found)

    resume_skills = {
        "technical": extract_technical_skills(cleaned_resume, master_technical_skills),
        "soft": extract_soft_skills(cleaned_resume, master_soft_skills)
    }

    jd_skills = {
        "technical": extract_technical_skills(cleaned_jd, master_technical_skills),
        "soft": extract_soft_skills(cleaned_jd, master_soft_skills)
    }

    st.subheader("Extracted Skills")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### Resume Skills")
        st.write("Technical:", resume_skills["technical"])
        st.write("Soft:", resume_skills["soft"])

    with col2:
        st.markdown("### JD Skills")
        st.write("Technical:", jd_skills["technical"])
        st.write("Soft:", jd_skills["soft"])

    def donut_chart(tech_count, soft_count, title):
        labels = ["Technical Skills", "Soft Skills"]
        sizes = [tech_count, soft_count]
        if sum(sizes) == 0:
            sizes = [1, 1]

        fig = plt.figure(figsize=(4, 4))
        ax = fig.add_axes([0.15, 0.15, 0.7, 0.7]) 

        wedges, _, _ = ax.pie(
            sizes,
            startangle=90,
            autopct="%1.0f%%",
            radius=1,
            wedgeprops=dict(width=0.4, edgecolor="white")
        )

        ax.set(aspect="equal")
        ax.set_title(title, pad=10)
        ax.legend(
            wedges,
            labels,
            loc="lower center",
            bbox_to_anchor=(0.5, -0.25),
            ncol=2,
            frameon=False
        )

        st.pyplot(fig, use_container_width=False)




    st.subheader("Skill Distribution Analysis")

    col1, col2 = st.columns(2)

    with col1:
        donut_chart(
            len(resume_skills["technical"]),
            len(resume_skills["soft"]),
            "Resume Skill Distribution"
        )

    with col2:
        donut_chart(
            len(jd_skills["technical"]),
            len(jd_skills["soft"]),
            "JD Skill Distribution"
        )

    def compute_metrics(resume_skills, jd_skills):
        resume_tech = set(resume_skills["technical"])
        resume_soft = set(resume_skills["soft"])
        jd_tech = set(jd_skills["technical"])
        jd_soft = set(jd_skills["soft"])

        matched_skills = (resume_tech & jd_tech) | (resume_soft & jd_soft)
        total_jd_skills = len(jd_tech) + len(jd_soft)

        match_percentage = 0
        if total_jd_skills > 0:
            match_percentage = round((len(matched_skills) / total_jd_skills) * 100, 1)

        return {
            "resume_tech": len(resume_tech),
            "resume_soft": len(resume_soft),
            "jd_tech": len(jd_tech),
            "jd_soft": len(jd_soft),
            "total_resume": len(resume_tech) + len(resume_soft),
            "total_jd": total_jd_skills,
            "match_percent": match_percentage
        }
    metrics = compute_metrics(resume_skills, jd_skills)

    st.subheader("Skill Extraction Metrics")

    col1, col2, col3, col4 = st.columns(4)

    col1.metric(
        label="Technical Skills",
        value=metrics["jd_tech"]
    )

    col2.metric(
        label="Soft Skills",
        value=metrics["jd_soft"]
    )

    col3.metric(
        label="Total Skills",
        value=metrics["total_jd"]
    )

    col4.metric(
        label="Avg Skill Match %",
        value=f'{metrics["match_percent"]}%'
    )


    # Milestone - 3

    def build_similarity_matrix(resume_skills, jd_skills):
        matrix = np.zeros((len(resume_skills), len(jd_skills)))

        for i, r_skill in enumerate(resume_skills):
            for j, j_skill in enumerate(jd_skills):
                if r_skill == j_skill:
                    matrix[i][j] = 1.0
                elif r_skill in j_skill or j_skill in r_skill:
                    matrix[i][j] = 0.5
                else:
                    matrix[i][j] = 0.0

        return matrix

    def plot_similarity_heatmap(matrix, resume_labels, jd_labels):
        fig, ax = plt.subplots(figsize=(8, 5))

        im = ax.imshow(matrix, cmap="YlGn")

        ax.set_xticks(range(len(jd_labels)))
        ax.set_yticks(range(len(resume_labels)))

        ax.set_xticklabels(jd_labels, rotation=45, ha="right")
        ax.set_yticklabels(resume_labels)

        ax.set_title("Skill Similarity Matrix")

        cbar = fig.colorbar(im, ax=ax)
        cbar.set_label("Similarity Score")

        st.pyplot(fig)

    resume_all_skills = resume_skills["technical"] + resume_skills["soft"]
    jd_all_skills = jd_skills["technical"] + jd_skills["soft"]

    st.markdown(
        """
        <div style="background-color:#4a5bdc;padding:20px;border-radius:10px">
            <h2 style="color:white;">
                Milestone 3: Skill Gap Analysis and Similarity Matching Module
            </h2>
            <p style="color:white;">
                • Skill similarity matrix visualization <br>
                • Resume vs JD skill comparison <br>
                • Missing skill identification
            </p>
        </div>
        <br>
        """,
        unsafe_allow_html=True
    )

    if resume_all_skills and jd_all_skills:
        similarity_matrix = build_similarity_matrix(
            resume_all_skills,
            jd_all_skills
        )

        plot_similarity_heatmap(
            similarity_matrix,
            resume_all_skills,
            jd_all_skills
        )
    else:
        st.warning("Insufficient skills to build similarity matrix")

    missing_skills = list(set(jd_all_skills) - set(resume_all_skills))

    def classify_skill_matches(similarity_matrix, resume_skills, jd_skills):
        matched = set()
        partial = set()

        for j, jd_skill in enumerate(jd_skills):
            column = similarity_matrix[:, j]

            if 1.0 in column:
                matched.add(jd_skill)
            elif 0.5 in column:
                partial.add(jd_skill)

        missing = set(jd_skills) - matched - partial

        return {
            "matched": list(matched),
            "partial": list(partial),
            "missing": list(missing)
        }

    skill_match_result = classify_skill_matches(
        similarity_matrix,
        resume_all_skills,
        jd_all_skills
    )

    def calculate_skill_match(resume_skills, jd_skills):
        resume_set = set(resume_skills)
        jd_set = set(jd_skills)

        matched = resume_set & jd_set
        missing = jd_set - resume_set

        # Partial = same base word (simple heuristic)
        partial = set()
        for jd_skill in missing.copy():
            for res_skill in resume_set:
                if jd_skill in res_skill or res_skill in jd_skill:
                    partial.add(jd_skill)

        missing = missing - partial

        avg_match = 0
        if len(jd_set) > 0:
            avg_match = round(
                ((len(matched) + 0.5 * len(partial)) / len(jd_set)) * 100,
                1
            )

        return {
            "matched": len(matched),
            "partial": len(partial),
            "missing": len(missing),
            "avg_match": avg_match
        }

    all_resume_skills = (
        set(resume_skills["technical"]) |
        set(resume_skills["soft"])
    )

    all_jd_skills = (
        set(jd_skills["technical"]) |
        set(jd_skills["soft"])
    )

    match_counts = calculate_skill_match(
        all_resume_skills,
        all_jd_skills
    )


    st.subheader("Skill Match Overview")

    left, right = st.columns([1.3, 1])

    # ---------------- LEFT: DONUT ----------------
    with left:
        labels = ["Matched", "Partially Matched", "Missing"]
        sizes = [match_counts["matched"], match_counts["partial"], match_counts["missing"]]
        colors = ["#28a745", "#fd7e14", "#dc3545"]

        if sum(sizes) == 0:
            sizes = [1, 1, 1]

        fig = plt.figure(figsize=(3.6, 3.6))
        ax = fig.add_axes([0.1, 0.15, 0.8, 0.75])

        wedges, _, _ = ax.pie(
            sizes,
            startangle=90,
            autopct="%1.0f%%",
            radius=1,
            colors=colors,
            wedgeprops=dict(width=0.4, edgecolor="white")
        )

        ax.set(aspect="equal")
        ax.set_title("Skill Match Distribution", pad=8)
        ax.legend(
            wedges,
            labels,
            loc="lower center",
            bbox_to_anchor=(0.5, -0.25),
            ncol=3,
            frameon=False
        )

        st.pyplot(fig, use_container_width=False)

    # ---------------- RIGHT: METRICS ----------------
    with right:
        with st.container():
            r1, r2 = st.columns(2)
            r1.metric("Matched Skills", match_counts["matched"])
            r2.metric("Partially Matched", match_counts["partial"])

            r3, r4 = st.columns(2)
            r3.metric("Missing Skills", match_counts["missing"])
            r4.metric("Avg Match %", f'{match_counts["avg_match"]}%')



    st.subheader("Skill Gap Details")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("### Matched Skills")
        st.write(skill_match_result["matched"])

    with col2:
        st.markdown("### Partial Matches")
        st.write(skill_match_result["partial"])

    with col3:
        st.markdown("### Missing Skills")
        st.write(skill_match_result["missing"])


    # Milestone - 4

    # ---------------- HEADER ----------------
    st.markdown("""
    <div style="background-color:#4a90e2;padding:15px;border-radius:8px">
    <h2 style="color:white;text-align:center;">
    Milestone 4: Dashboard and Report Export Module
    </h2>
    <p style="color:white;text-align:center;">
    Interactive dashboard • Graphs • Multi-format report export
    </p>
    </div>
    """, unsafe_allow_html=True)

    st.write("")

    # ---------------- DATA ----------------
    skills = list(all_jd_skills)
    resume_scores = [
        100 if skill in all_resume_skills else 0
        for skill in skills
    ]

    job_scores = [100 for _ in skills]

    overall_match = match_counts["avg_match"]
    matched_skills = match_counts["matched"]
    missing_skills_count = match_counts["missing"]

    # ---------------- METRICS ----------------
    c1, c2, c3 = st.columns(3)
    c1.metric("Overall Match", f"{overall_match}%")
    c2.metric("Matched Skills", str(matched_skills))
    c3.metric("Missing Skills", str(missing_skills_count))

    # ---------------- BAR CHART ----------------
    bar_fig = go.Figure()
    bar_fig.add_bar(x=skills, y=resume_scores, name="Resume Skills")
    bar_fig.add_bar(x=skills, y=job_scores, name="Job Requirements")
    bar_fig.update_layout(
        title="Skill Match Overview",
        barmode="group",
        height=400
    )

    # ---------------- RADAR CHART ----------------
    radar_skills = skills[:5] if len(skills) >= 5 else skills
    radar_resume = resume_scores[:5] if len(resume_scores) >= 5 else resume_scores
    radar_job = job_scores[:5] if len(job_scores) >= 5 else job_scores
    
    radar_fig = go.Figure()
    radar_fig.add_trace(go.Scatterpolar(
        r=radar_resume,
        theta=radar_skills,
        fill='toself',
        name='Current Profile'
    ))
    radar_fig.add_trace(go.Scatterpolar(
        r=radar_job,
        theta=radar_skills,
        fill='toself',
        name='Job Requirements'
    ))
    radar_fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
        title="Technical Skill Comparison",
        height=400
    )

    left, right = st.columns([2, 1])
    left.plotly_chart(bar_fig, use_container_width=True)
    right.plotly_chart(radar_fig, use_container_width=True)

    # ---------------- SKILL PROGRESS ----------------
    st.subheader("Skill Comparison")
    sample_skills = [("Python", 92), ("Machine Learning", 88), ("SQL", 65)]
    for skill, score in sample_skills:
        st.write(skill)
        st.progress(score / 100)

    # ---------------- UPSKILLING ----------------
    st.subheader("Upskilling Recommendations")
    st.info("☁️ AWS Cloud Services\nComplete AWS Certified Solutions Architect course")
    st.info("📊 Advanced Statistics\nEnroll in Advanced Statistics for Data Science")
    st.info("📈 Project Management\nConsider PMP certification")

    # ---------------- REPORT CONTENT ----------------
    report_text = f"""
SKILL GAP ANALYSIS REPORT

Overall Match: {overall_match}%
Matched Skills: {matched_skills}
Missing Skills: {missing_skills_count}

SKILL DETAILS
"""

    for i in range(len(skills)):
        report_text += f"""
{skills[i]}
Resume Score: {resume_scores[i]}%
Job Requirement: {job_scores[i]}%
"""

    st.subheader("Report Download")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        # ---------------- CSV ----------------
        df = pd.DataFrame({
            "Skill": skills,
            "Resume Score (%)": resume_scores,
            "Job Requirement (%)": job_scores
        })

        st.download_button(
            "⬇️ CSV Report",
            df.to_csv(index=False).encode("utf-8"),
            "skill_gap_report.csv",
            "text/csv"
        )

    with col2:
        # ---------------- TXT ----------------
        st.download_button(
            "⬇️ TXT Report",
            report_text.encode("utf-8"),
            "skill_gap_report.txt",
            "text/plain"
        )

    with col3:
        # ---------------- DOC ----------------
        def generate_doc(text):
            doc = Document()
            doc.add_heading("Skill Gap Analysis Report", level=1)
            for line in text.split("\n"):
                doc.add_paragraph(line)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        doc_file = generate_doc(report_text)

        st.download_button(
            "⬇️ DOCX Report",
            doc_file,
            "skill_gap_report.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    with col4:
        # ---------------- PDF ----------------
        def generate_pdf(text):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in text.split("\n"):
                pdf.multi_cell(0, 8, line)
            return pdf.output(dest="S").encode("latin-1")

        pdf_bytes = generate_pdf(report_text)

        st.download_button(
            "⬇️ PDF Report",
            pdf_bytes,
            "skill_gap_report.pdf",
            "application/pdf"
        )

    st.success("Dashboard loaded successfully ✔️")